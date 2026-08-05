#!/usr/bin/env python3
"""14라운드 — 두 문서 빌더(build_html·build_docx)에서 고친 것을 지키는 항목.

`selftest.py` 가 옆의 `selftest_*.py` 를 글로브로 끌어가므로 이 파일은 따로 등록할
곳이 없다. 항목마다 어느 결함을 지키는지 적는다.

여기 담긴 것의 공통점 하나: **두 문서가 같은 재료를 다르게 읽고 있었다.** 표지
정보표의 행 폭(docx 4칸 · HTML 2칸 언팩), 문서가 세는 테이블 집합(본문은 영역 ·
개수는 스키마 전체), 역할명이 빈 경우(HTML 은 안 그리고 docx 는 ` · ` 만 남겼다).
그래서 한쪽만 보는 시험으로는 못 잡는다 — 아래 항목들은 둘을 나란히 세운다.

15라운드가 덧붙인 것은 성격이 하나 더 있다. 14라운드는 이 파일들을 고쳐 놓고 "회귀
시험 7개, 전부 mutation 확인" 이라고 적었는데, 독립 검증자가 제 뮤테이션을 돌리니
**고친 자리 다섯이 되돌려도 전부 초록**이었다. 고친 것은 맞았지만 지키는 시험이
없었다. 그래서 아래 항목의 합격 기준은 '고쳤다' 가 아니라 **'되돌리면 빨강이다'** 다.
어느 뮤턴트를 잡는 항목인지 각 케이스 주석에 적는다.
"""
import json
import os
import re
import sys
from html import escape

from selftest_kit import (HERE, Fail, case, col, drawn_names, eq, has, run, table,
                          write_schema)

EXPECT_CASES = 26       # 등록 개수를 파일이 스스로 못박는다 (selftest_kit.load_extras)

EXAMPLES = HERE.parent / 'examples'


def schema_for(spec):
    """spec 이 이름을 대는 테이블만 담은 최소 스키마.

    배포 예제는 '이 문서를 이렇게 적으라' 는 뼈대일 뿐 스키마를 담지 않는다. 예제가
    실제로 문서를 내는지 보려면 그 뼈대가 가리키는 테이블이 있어야 한다.
    """
    sch = {}
    for a in spec.get('areas', []):
        for t in a[3]:
            sch[t] = table(t.split('.')[-1],
                           [col('id'), col('name', 'text', comment='a name')],
                           schema=a[2] or 'public', pk=['id'])
    return sch


def area_keys(html):
    """<h4 id=...> 가 말하는 테이블 키 — 문서가 실제로 절을 내준 것."""
    return re.findall(r'<h4 id="[^"]*">([^<]*)', html)


# ── 1. 함께 배포하는 예제가 세 산출물을 다 내는가 ────────────────────────────

@case('artifacts: every shipped example spec builds all three outputs')
def _(work):
    # 14라운드. examples/minimal*.spec.json 의 `doc.meta` 는 4칸짜리 한 행인데
    # build_html.py 만 `for k, v in meta` 로 2칸을 언팩해서 ValueError 로 죽었다.
    # 같은 spec 으로 build_docx.py 는 멀쩡히 저장하므로, **둘 중 하나만 나왔다** —
    # HTML 정의서가 한 글자도 안 나오는데 docx 는 나오니 실패로 보이지도 않았다.
    # 6라운드가 같은 버그를 docx 쪽에서 고치면서 시험도 docx 쪽만 남긴 자리다.
    specs = sorted(EXAMPLES.glob('*.spec.json'))
    if len(specs) < 2:
        raise Fail(f'no example specs found under {EXAMPLES} — this case would '
                   f'measure nothing')
    for i, sp in enumerate(specs):
        spec = json.loads(sp.read_text(encoding='utf-8'))
        w = work.parent / f'ex{i}'
        write_schema(w, schema_for(spec))
        (w / 'erd.spec.json').write_text(sp.read_text(encoding='utf-8'), encoding='utf-8')
        for script in ('build_erd.py', 'build_docx.py', 'build_html.py'):
            r = run(script, w)
            if 'Traceback' in r.stderr:
                raise Fail(f'{sp.name} → {script}:\n{r.stderr[-500:]}')
        for out in ('T.html', 'T.docx', 'T.graphml'):
            if not (w / out).exists():
                raise Fail(f'{sp.name} produced no {out} — a shipped example must not '
                           f'leave one of the three artifacts behind')
        # 표지가 실제로 그 값을 실었는지도 본다. 죽지만 않고 정보표를 통째로
        # 버리는 것도 같은 결함의 다른 얼굴이다.
        #
        # 15R — **뮤턴트 K**: 여기는 '있어야 할 것이 있는가' 만 물었다. 그래서
        # `meta_pairs` 에서 빈 쌍을 거르는 조건을 지워도 전부 초록이었다 — 표지가
        # `Author: Jane Doe &nbsp;|&nbsp; : ` 로 나가도 아무도 말하지 않았다. 함께
        # 배포하는 minimal*.spec.json 은 `["Date", …, "Author", ""]` 라 그 자리를
        # 실제로 밟는다. 있는 것만 보지 말고 **표지 줄 전체가 그것뿐인지**를 본다.
        html = (w / 'T.html').read_text(encoding='utf-8')
        want = []
        for row in spec.get('doc', {}).get('meta', []):
            cells = [('' if c is None else c) for c in list(row)] + [''] * 4
            for k, v in ((cells[0], cells[1]), (cells[2], cells[3])):
                if str(v).strip():                 # 내용이 빈 쌍은 표지에 안 실린다
                    want.append(f'{escape(str(k))}: {escape(str(v))}')
        m = re.search(r'<div class="sub">(.*?)</div>', html, re.S)
        if want and not m:
            raise Fail(f'{sp.name}: the cover carries no meta line at all')
        if m:
            eq(m.group(1).split(' &nbsp;|&nbsp; ') if m.group(1) else [], want,
               f'{sp.name}: the cover line is exactly the pairs that have content — '
               f'an empty cell must not leave a bare ": " behind')


# ── 2. 문서가 말하는 수와 문서가 싣는 것 ─────────────────────────────────────

def multi_db(work):
    """두 DB 를 합쳐 영역은 **기능별**로 적은 스키마 — SKILL.md 가 권하는 흐름."""
    write_schema(work, {
        'shop.orders': table('orders', [col('id'), col('user_id')], pk=['id'], db='shop'),
        'shop.order_items': table('order_items', [col('id'), col('order_id')],
                                  pk=['id'], db='shop'),
        'shop.users': table('users', [col('id'), col('email', 'text')],
                            pk=['id'], db='shop'),
        'mart.agg_daily': table('agg_daily', [col('day', 'date'), col('n')], db='mart'),
        'mart.agg_user': table('agg_user', [col('user_id'), col('n')], db='mart')})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', 'Order flow', 'public',
                   ['shop.orders', 'shop.order_items', 'mart.agg_daily']],
                  ['B', 'Members', 'public', ['shop.users', 'mart.agg_user']]],
        'doc': {'db_names': {'shop': 'shop prod', 'mart': 'mart warehouse'}}}),
        encoding='utf-8')


@case('html: the DB summary adds up to what the document actually lists')
def _(work):
    # 14라운드. 영역이 속한 DB 를 **영역의 첫 테이블**로 정하면서 개수는 테이블마다
    # 제 db 로 셌다. 한 영역이 두 DB 에 걸치면 (a) <h2>·목차가 남의 DB 테이블까지
    # 이 DB 것으로 적었고 (b) 어느 영역의 첫 테이블도 아닌 DB 는 요약표에서 **행
    # 자체가 사라졌다** — doc.db_names 로 이름까지 적어 준 DB 인데도. 요약표는
    # 3테이블·6컬럼이라 적었는데 문서가 싣는 것은 5테이블·12컬럼이었다.
    multi_db(work)
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    schema = json.loads((work / 'schema.json').read_text(encoding='utf-8'))

    rows = re.findall(r'<tr><td><b>([^<]*)</b></td><td>(\d+)</td><td>(\d+)</td>'
                      r'<td>(\d+)</td><td>(\d+)</td></tr>', html)
    if not rows:
        raise Fail('the DB summary table has no rows at all — nothing to add up')
    names = {'shop': 'shop prod', 'mart': 'mart warehouse'}
    eq(sorted(r[0] for r in rows), sorted(names.values()),
       'every db in the schema gets exactly one row in the summary — the DB that is '
       'no area first table used to lose its row entirely')

    # 본문은 <h2> 마다 한 DB 다. 그 아래 실린 <h4> 를 세어 요약표와 맞춰 본다.
    listed = {}
    for chunk in html.split('<h2>')[1:]:
        head = chunk.split('</h2>')[0]
        m = re.match(r'DB: (.*) · tables (\d+)$', head)
        if not m:
            continue
        keys = area_keys(chunk)
        eq(len(keys), int(m.group(2)),
           f'the <h2> for {m.group(1)!r} says how many tables follow it')
        listed[m.group(1)] = keys
    eq(sorted(listed), sorted(names.values()), 'one body section per db')

    total = 0
    for label, n_areas, n_tables, n_cols, _n_fks in rows:
        keys = listed[label]
        eq(int(n_tables), len(keys), f'summary row {label!r}: tables')
        eq(int(n_cols), sum(len(schema[k]['columns']) for k in keys),
           f'summary row {label!r}: columns')
        if int(n_areas) < 1:
            raise Fail(f'summary row {label!r} claims {n_areas} areas')
        total += len(keys)
    eq(total, len(area_keys(html)),
       'the summary adds up to every section the document lists')
    # 한 영역이 두 DB 에 걸친다는 사실을 문서가 말해야 한다 — 조용히 쪼개면 읽는
    # 사람은 같은 이름의 영역이 왜 두 번 나오는지 알 수 없다.
    has(html, '<span class="badge">shop prod</span>',
        'an area split across DBs says which DB this section is')

    # 반대 방향도 본다: 영역이 테이블을 하나 놓치면 요약표도 그만큼 적게 세야 한다.
    # (그 입력은 config.load_spec 이 '기타' 영역으로 받아 주므로, 여기서는 영역
    # 목록에서 직접 뺀다 — 재는 것은 '요약표가 어느 집합에서 수를 뽑는가' 다.)
    drv = work / 'drop_one.py'
    drv.write_text('import sys\n'
                   f'sys.path.insert(0, {str(HERE)!r})\n'
                   'import erd\n'
                   "erd.AREAS[0][3].remove('mart.agg_daily')\n"
                   'import build_html\n'
                   'build_html.main()\n', encoding='utf-8')
    run(str(drv), work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    rows = re.findall(r'<tr><td><b>([^<]*)</b></td><td>(\d+)</td><td>(\d+)</td>'
                      r'<td>(\d+)</td><td>(\d+)</td></tr>', html)
    eq(sum(int(r[2]) for r in rows), len(area_keys(html)),
       'the summary counts what the document lists, not what the schema holds')


@case('docx: the table count the document states matches the sections it lists')
def _(work):
    # 14라운드. 본문 4장은 AREAS 를 돌고 표지·1.2·4장 머리말은 len(SCHEMA) 를 세서,
    # 'of 11 tables' 라고 적힌 문서에 4.x.x 절이 10개만 실렸다. 5장 관계표도 SCHEMA
    # 를 돌아 **본문에 절이 없는 테이블의 관계**를 실었다.
    #
    # 그 어긋남을 만들던 입력(영역에 안 든 테이블)은 같은 라운드에 config.load_spec
    # 이 '기타' 영역으로 받게 고쳤다. 그래서 여기서는 영역 목록에서 테이블 하나를
    # 직접 빼고 문서를 만든다 — 문서가 제 숫자를 어디서 뽑는지가 이 항목이 재는
    # 것이고, 그건 영역이 어떻게 채워지든 성립해야 한다.
    write_schema(work, {
        'orders': table('orders', [col('id')], pk=['id']),
        'users': table('users', [col('id')], pk=['id']),
        'payments': table('payments', [col('id'), col('order_id')], pk=['id'],
                          fks=[{'column': 'order_id', 'ref_table': 'orders',
                                'ref_column': 'id', 'on_delete': 'NO ACTION'}])})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', 'All', 'public', ['orders', 'users', 'payments']]]}),
        encoding='utf-8')
    run('merge_desc.py', work)
    run('build_erd.py', work)
    drv = work / 'drop_one.py'
    drv.write_text('import sys\n'
                   f'sys.path.insert(0, {str(HERE)!r})\n'
                   'import erd\n'
                   "erd.AREAS[0][3].remove('payments')   # 영역에서만 뺀다\n"
                   'import build_docx\n'
                   'p = build_docx.build()\n'
                   'print(p)\n', encoding='utf-8')
    run(str(drv), work)

    from docx import Document
    doc = Document(str(work / 'T.docx'))
    heads = [p.text for p in doc.paragraphs if re.match(r'^4\.\d+\.\d+ ', p.text)]
    listed = sorted(h.split(' ', 1)[1].split(' · ')[0] for h in heads)
    eq(listed, ['orders', 'users'], 'the areas decide which tables get a section')

    text = '\n'.join(p.text for p in doc.paragraphs)
    said = [int(n) for n in re.findall(r'relationships of (\d+) tables', text)]
    said += [int(n) for n in re.findall(r'Tables (\d+) · columns', text)]
    if not said:
        raise Fail('neither chapter 1.2 nor chapter 4 states a table count — '
                   'this case would measure nothing')
    for n in said:
        eq(n, len(listed), 'the number the document states is the number it lists')
    cover = [c.text for r in doc.tables[0].rows for c in r.cells]
    if 'Tables' not in cover:
        raise Fail(f'the cover table does not state a table count at all: {cover}')
    eq(cover[cover.index('Tables') + 1], str(len(listed)),
       'the cover table counts the same set')
    # 5장은 절이 있는 테이블의 관계만 싣는다
    rel = [[c.text for c in r.cells] for t in doc.tables for r in t.rows
           if any('ON DELETE' in c.text for c in r.cells)]
    for r in rel:
        if r[0] not in listed:
            raise Fail(f'chapter 5 carries a relationship of {r[0]!r}, which has no '
                       f'section in chapter 4: {r}')


# ── 3. 앵커 ──────────────────────────────────────────────────────────────────

@case('html: every table gets its own anchor')
def _(work):
    # 14라운드. `'tb_' + re.sub(r'[^a-zA-Z0-9_]', '_', key)` 는 뭉개는 함수라 단사가
    # 아니다. `order-items` 와 `order_items` 가 같은 `tb_order_items` 가 됐고(따옴표
    # 친 식별자는 13라운드가 지원 대상으로 넣은 입력이다), 다중 DB 에서는 `a.b_c` 와
    # `a_b.c` 가 부딪혔다. 같은 id 가 둘이면 목차도 FK 상호참조도 전부 앞엣것으로
    # 가서, 문서가 조용히 남의 테이블을 가리킨다.
    write_schema(work, {
        'order-items': table('order-items', [col('id')], pk=['id']),
        'order_items': table('order_items', [col('id'), col('ref')], pk=['id'],
                             fks=[{'column': 'ref', 'ref_table': 'order-items',
                                   'ref_column': 'id', 'on_delete': 'CASCADE'}]),
        'a.b_c': table('b_c', [col('id')], pk=['id'], db='a'),
        'a_b.c': table('c', [col('id')], pk=['id'], db='a_b')})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')

    ids = re.findall(r'<h4 id="([^"]+)"', html)
    eq(len(set(ids)), 4, f'one id per table, all different — got {ids}')
    eq(len(ids), 4, 'and one section per table')
    all_ids = re.findall(r'\sid="([^"]+)"', html)
    dup = sorted({i for i in all_ids if all_ids.count(i) > 1})
    eq(dup, [], 'no id appears twice anywhere in the document')
    for href in sorted(set(re.findall(r'href="#([^"]+)"', html))):
        eq(all_ids.count(href), 1, f'href #{href} lands on exactly one id')
    # 뭉개지지 않는 이름은 제 이름을 그대로 갖는다 — 읽는 사람이 URL 로 쓰는 값이다
    has(html, '<h4 id="tb_order_items">order_items',
        'a name that needs no mangling keeps its plain anchor')


# ── 4. 빈 역할명 ─────────────────────────────────────────────────────────────

@case('docx: a table with no role gets a heading with no dangling separator')
def _(work):
    # 14라운드. `f'4.{ai}.{ti} {tname} · {ROLE.get(tname, "")}'` 는 역할명이 비어도
    # 구분자를 찍어 '4.1.1 users · ' 로 끝났다. 같은 값이 비면 HTML 은 역할 블록을
    # 아예 안 그린다 — 같은 데이터를 두 문서가 다르게 처리하던 자리다.
    write_schema(work, {
        'plain': table('plain', [col('id')], pk=['id']),
        'named': table('named', [col('id')], pk=['id'])})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', 'All', 'public', ['plain', 'named']]],
        'roles': {'named': 'Has a role'}}), encoding='utf-8')
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_docx.py', work)
    from docx import Document
    heads = [p.text for p in Document(str(work / 'T.docx')).paragraphs
             if re.match(r'^4\.\d+\.\d+ ', p.text)]
    eq(sorted(heads), sorted(['4.1.1 plain', '4.1.2 named · Has a role']),
       'an empty role leaves no separator behind, a real one is still joined')


@case('docx: a column made unique says so')
def _(work):
    # 14라운드. UNIQUE 가 docx 어디에도 없었다 — 4장 컬럼표의 '구분' 칸은 PK·FK 만
    # 알고, docx 에는 HTML 이 가진 제약·인덱스 절도 없다. 제출용 문서에서만 스키마
    # 사실 하나가 통째로 사라지고 있었다.
    write_schema(work, {
        't': table('t', [col('id'), col('email', 'text'), col('other', 'text')],
                   pk=['id'], uniques=[['email']])})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_docx.py', work)
    from docx import Document
    doc = Document(str(work / 'T.docx'))
    marks = {}
    for t in doc.tables:
        for r in t.rows:
            cells = [c.text for c in r.cells]
            if len(cells) == 4 and cells[1] in ('id', 'email', 'other'):
                marks[cells[1]] = cells[0]
    eq(marks, {'id': 'PK', 'email': 'UQ', 'other': ''},
       'the Kind column tells a unique column from a plain one')


# ── 5. 산출 경로 ─────────────────────────────────────────────────────────────

@case('errors: ERD_HTML_OUT pointed at a directory is a message, not a traceback')
def _(work):
    # 14라운드. `Path(os.environ.get('ERD_HTML_OUT', …))` 를 그대로 열어
    # `IsADirectoryError: [Errno 21]` 이 사용자에게 그대로 갔다 — 어느 변수가
    # 그랬는지 한 글자도 없는 역추적이다. 빈 문자열도 `Path('')` → `Path('.')` 라
    # 같은 자리에서 같은 모양으로 죽었다.
    write_schema(work, {'t': table('t', [col('id')], pk=['id'])})
    run('build_erd.py', work)
    d = work / 'somewhere'
    d.mkdir(parents=True, exist_ok=True)
    r = run('build_html.py', work, env={'ERD_HTML_OUT': str(d)}, expect_ok=False)
    if r.returncode == 0:
        raise Fail('writing the document into a directory reported success')
    both = r.stdout + r.stderr
    if 'Traceback' in both:
        raise Fail(f'a raw traceback reached the user:\n{both[-500:]}')
    has(both, 'ERD_HTML_OUT', 'the message names the variable that has to change')
    # 빈 값은 '설정하지 않은 것' 으로 친다 — 다른 경로 변수들과 같은 규칙이다
    r = run('build_html.py', work, env={'ERD_HTML_OUT': ''})
    if 'Traceback' in r.stderr:
        raise Fail(f'an empty ERD_HTML_OUT tracebacked:\n{r.stderr[-400:]}')
    if not (work / 'T.html').exists():
        raise Fail('an empty ERD_HTML_OUT must fall back to the default file name')


# ── 6. 수를 어디서 뽑는가 — 자리마다 ────────────────────────────────────────
# 14라운드는 이 부류를 두 케이스로 덮었다고 적었지만, 그 둘은 **일부만** 물었다.
# 표지의 컬럼 수·FK 수, 4장 머리말의 컬럼 수, 5.1 의 관계 총계, HTML 이 다 쓰고
# 찍는 줄 — 네 자리가 아무 데서도 안 재고 있었다. 자리마다 따로 반증되게 한다.

def counted_fixture(work):
    """영역에 든 둘과 들지 않은 하나. 안 든 쪽이 컬럼도 FK 도 혼자 갖고 있다."""
    write_schema(work, {
        'orders': table('orders', [col('id'), col('a'), col('b')], pk=['id']),
        'users': table('users', [col('id'), col('c')], pk=['id']),
        'payments': table('payments', [col('id'), col('order_id'), col('x'), col('y')],
                          pk=['id'],
                          fks=[{'column': 'order_id', 'ref_table': 'orders',
                                'ref_column': 'id', 'on_delete': 'NO ACTION'}])})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', 'All', 'public', ['orders', 'users', 'payments']]]}),
        encoding='utf-8')
    run('merge_desc.py', work)
    run('build_erd.py', work)


def drop_from_area(work, name, builder, tail='build.main()'):
    """영역 목록에서 테이블 하나를 빼고 문서를 만드는 드라이버.

    그 어긋남을 만들던 입력(영역에 안 든 테이블)은 14라운드에 config.load_spec 이
    '기타' 영역으로 받게 고쳤다. 재는 것은 '문서가 제 숫자를 어디서 뽑는가' 이고,
    그건 영역이 어떻게 채워지든 성립해야 하므로 여기서는 영역에서 직접 뺀다.
    """
    drv = work / f'drop_{builder}.py'
    drv.write_text('import sys\n'
                   f'sys.path.insert(0, {str(HERE)!r})\n'
                   'import erd\n'
                   f'erd.AREAS[0][3].remove({name!r})\n'
                   f'import {builder} as build\n'
                   f'{tail}\n', encoding='utf-8')
    return run(str(drv), work)


@case('html: the line the run prints counts the sections the file carries')
def _(work):
    # 15R — **뮤턴트 L**: `log.html_done` 의 `tables=` 를 `len(SCHEMA)` 로 되돌려도
    # 141개가 전부 초록이었다. 그러면 `HTML tables 3 …` 이라고 찍히는데 파일에 실린
    # 절은 2개다. 사람이 산출물을 안 열어 보고 그 줄만 믿는 것이 정상 흐름이라,
    # 이 줄은 산출물만큼이나 문서다 — 찍힌 수와 실린 절을 맞춘다.
    counted_fixture(work)
    r = drop_from_area(work, 'payments', 'build_html')
    html = (work / 'T.html').read_text(encoding='utf-8')
    m = re.search(r'tables (\d+)', r.stdout)
    if not m:
        raise Fail(f'the run printed no table count at all — this case would measure '
                   f'nothing:\n{r.stdout!r}')
    listed = area_keys(html)
    eq(listed, ['orders', 'users'], 'the areas decide which tables get a section')
    eq(int(m.group(1)), len(listed),
       'the number the finish line prints is the number of sections in the file')


@case('docx: every count the document states is the count of what it lists')
def _(work):
    # 15R — **뮤턴트 DD·FF**: 14라운드의 docx 케이스는 표지의 `Tables` 와 4장 절 수만
    # 봤다. 그래서 `doc_counts` 의 컬럼 수를 SCHEMA 전체에서 뽑게 되돌려도(DD),
    # build_docx 의 `N_FKS` 를 SCHEMA 전체에서 뽑게 되돌려도(FF) 전부 초록이었다 —
    # 표지가 `Columns 9`·`FKs 1` 이라고 적는데 문서가 싣는 것은 5컬럼·0관계다.
    # 문서가 수를 말하는 **네 자리를 전부** 같은 집합에 건다.
    counted_fixture(work)
    drop_from_area(work, 'payments', 'build_docx', tail='build.build()')

    schema = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    from docx import Document
    doc = Document(str(work / 'T.docx'))
    heads = [p.text for p in doc.paragraphs if re.match(r'^4\.\d+\.\d+ ', p.text)]
    listed = sorted(h.split(' ', 1)[1].split(' · ')[0] for h in heads)
    eq(listed, ['orders', 'users'], 'the areas decide which tables get a section')
    n_cols = sum(len(schema[k]['columns']) for k in listed)
    n_fks = sum(len(schema[k]['fks']) for k in listed)
    if n_cols == sum(len(v['columns']) for v in schema.values()):
        raise Fail('the fixture no longer tells the two sets apart — the table left out '
                   'of the areas must carry columns of its own')

    cover = [c.text for r in doc.tables[0].rows for c in r.cells]
    for word, want in (('Tables', len(listed)), ('Columns', n_cols), ('FKs', n_fks)):
        if word not in cover:
            raise Fail(f'the cover states no {word} at all: {cover}')
        eq(cover[cover.index(word) + 1], str(want),
           f'the cover counts the sections the document lists ({word})')

    text = '\n'.join(p.text for p in doc.paragraphs)
    seen = 0
    for rx, want, what in (
            (r'relationships of (\d+) tables', [len(listed)], 'chapter 1.2 scope'),
            (r'Tables (\d+) · columns (\d+)', [len(listed), n_cols], 'chapter 4 lead-in'),
            # 5.1 은 그 수로 문단을 연다. 4장 머리말도 '… in total.' 로 끝나므로
            # 줄머리에 못박지 않으면 엉뚱한 수를 집는다.
            (r'^(\d+) in total\.', [n_fks], 'chapter 5.1 relationship total')):
        got = re.findall(rx, text, re.M)
        if not got:
            raise Fail(f'{what}: the document states no number here — this case would '
                       f'measure nothing')
        for row_ in got:
            eq([int(x) for x in ((row_,) if isinstance(row_, str) else row_)], want,
               f'{what} counts what the document lists')
        seen += 1
    eq(seen, 3, 'all three prose counts were found and checked')


# ── 7. 눌러도 아무 데도 안 가는 링크 ────────────────────────────────────────

@case('html: a reference link always lands on a section this file has')
def _(work):
    # 15R — **뮤턴트 U**: `col_flags` 에서 '절이 있는 테이블에만 링크' 가드를 지워도
    # 전부 초록이었다. 14라운드의 앵커 케이스는 **모든 테이블이 실리는** 스키마를
    # 써서, 링크가 빠질 수 있는 자리를 한 번도 밟지 않았다. 여기서는 부모를 영역에서
    # 빼서 죽은 링크가 실제로 생길 수 있게 한다.
    counted_fixture(work)
    # payments 가 아니라 **부모**를 뺀다 — 자식은 남아서 링크를 걸려 한다
    drop_from_area(work, 'orders', 'build_html')
    html = (work / 'T.html').read_text(encoding='utf-8')
    listed = area_keys(html)
    eq(sorted(listed), ['payments', 'users'], 'the parent has no section in this file')

    all_ids = re.findall(r'\sid="([^"]+)"', html)
    dangling = sorted({h for h in re.findall(r'href="#([^"]+)"', html)
                       if all_ids.count(h) != 1})
    eq(dangling, [],
       'every href lands on exactly one id — a link to a table with no section is a '
       'link that goes nowhere')
    # 링크를 지우는 것으로 도망가지 못하게, 관계 자체는 여전히 적혀 있어야 한다
    has(html, 'orders.id',
        'the reference is still stated in words even when it cannot be linked')


# ── 8. 앵커가 옮겨 앉지 않는다 ──────────────────────────────────────────────

def anchors_of(work, keys):
    write_schema(work, {k: table(k.split('.')[-1], [col('id')], pk=['id']) for k in keys})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    return dict(zip(re.findall(r'<h4 id="([^"]+)"', html),
                    area_keys(html))), html


@case('html: an anchor is decided by the table key alone')
def _(work):
    # 15R — 새 버그. 꼬리 번호는 **자리를 세는 값**이라 스키마가 자라면 옮겨 앉았다:
    #   테이블 x-y · x.y      → {'tb_x_y': 'x-y',  'tb_x_y_2': 'x.y'}
    #   여기에 'x y' 를 더하면 → {'tb_x_y': 'x y',  'tb_x_y_2': 'x-y', …}
    # `#tb_x_y` 도 `#tb_x_y_2` 도 어제와 다른 테이블로 간다. 코드가 스스로 "읽는
    # 사람이 URL 로 쓰는 값" 이라고 적어 둔 값인데 안정성을 재는 것이 없었다.
    two, _ = anchors_of(work.parent / 'a2', ['x-y', 'x.y'])
    three, html = anchors_of(work.parent / 'a3', ['x y', 'x-y', 'x.y'])
    before = {k: i for i, k in two.items()}
    after = {k: i for i, k in three.items()}
    eq(len(two), 2, f'two tables, two anchors — got {two}')
    eq(len(three), 3, f'three tables, three anchors — got {three}')
    for k in ('x-y', 'x.y'):
        eq(after[k], before[k],
           f'the anchor of {k!r} must not move when another table joins the schema')
    # 단사성은 여전히 이 파일이 센다 — 결정적이기만 하고 겹치면 더 나쁘다
    all_ids = re.findall(r'\sid="([^"]+)"', html)
    eq(sorted({i for i in all_ids if all_ids.count(i) > 1}), [],
       'no id appears twice anywhere in the document')

    # 같은 부류: 영역 id 도 (영역코드, DB) 에서만 나와야 한다. 한 영역이 두 DB 에
    # 걸치면 절이 둘 생기는데, 예전엔 둘째가 `area_A_2` 라 DB 가 하나 늘거나 순서가
    # 바뀌면 두 절이 서로 자리를 바꿨다.
    def area_ids(w, dbs):
        write_schema(w, {f'{d}.t{i}': table(f't{i}', [col('id')], pk=['id'], db=d)
                         for i, d in enumerate(dbs)})
        (w / 'erd.spec.json').write_text(json.dumps({
            'areas': [['A', 'One area', 'public',
                       [f'{d}.t{i}' for i, d in enumerate(dbs)]]]}), encoding='utf-8')
        run('merge_desc.py', w)
        run('build_erd.py', w)
        run('build_html.py', w)
        h = (w / 'T.html').read_text(encoding='utf-8')
        return re.findall(r'<h3 id="([^"]+)"', h)

    two_db = area_ids(work.parent / 'd2', ['shop', 'mart'])
    three_db = area_ids(work.parent / 'd3', ['warehouse', 'shop', 'mart'])
    eq(len(set(two_db)), 2, f'one id per (area, db) slice — got {two_db}')
    eq(len(set(three_db)), 3, f'one id per (area, db) slice — got {three_db}')
    eq([i for i in three_db if i in two_db], two_db,
       'the area ids that were already published keep pointing at the same slice '
       f'when a third DB joins — was {two_db}, now {three_db}')


# ── 9. spec 이 손으로 적는 장 ───────────────────────────────────────────────

@case('html: the chapters a spec writes by hand reach both documents')
def _(work):
    # 15R — 14라운드는 `doc.derives`·`doc.mapping`·`doc.open_items` 셋이 HTML 에 없는
    # 것을 하나의 근거("매체 차이")로 함께 넘겼다. 그 근거는 derives 에만 참이다 —
    # 그쪽은 HTML 이 인라인 SVG 안에 실제로 그린다. 이 둘은 **표로만 실리는 것**이라
    # 표가 없으면 그냥 빠진 것이고, 같은 spec 을 준 사람이 docx 에는 있고 HTML 에는
    # 없는 장을 갖게 된다.
    write_schema(work, {'t': table('t', [col('id')], pk=['id'])})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', 'All', 'public', ['t']]],
        'doc': {'mapping': [['1', 'proposed_name', 't', 'yes', 'renamed on build']],
                'open_items': [['P1', 'retention', 't', 'none yet', 'decide with ops']],
                'mapping_note': 'mapping footnote', 'open_note': 'open footnote'}}),
        encoding='utf-8')
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_docx.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')

    from docx import Document
    doc = Document(str(work / 'T.docx'))
    cells = {c.text for t_ in doc.tables for r in t_.rows for c in r.cells}
    words = '\n'.join(p.text for p in doc.paragraphs)
    for v in ('proposed_name', 'renamed on build', 'retention', 'decide with ops'):
        if v not in cells:
            raise Fail(f'docx lost {v!r} — this case would measure nothing')
        has(html, v, f'HTML carries the same hand-written cell as docx: {v!r}')
    for v in ('mapping footnote', 'open footnote'):
        if v not in words:
            raise Fail(f'docx lost {v!r} — this case would measure nothing')
        has(html, v, f'HTML carries the same hand-written note as docx: {v!r}')
    # 장 제목도 같은 카탈로그 항목에서 나와야 한다 — 두 문서가 같은 장을 다른 말로
    # 부르면 그 자체가 '같은 재료를 다르게 읽는 자리' 다
    for head in [p.text for p in doc.paragraphs if re.match(r'^[67]\. ', p.text)]:
        has(html, escape(head), f'HTML calls the chapter what docx calls it: {head!r}')


# ── 10. 사람에게 나가는 줄은 전부 카탈로그를 거친다 ─────────────────────────

BUILDERS = ('build_html.py', 'build_docx.py', 'build_erd.py')

# 사람 앞에 글자를 내놓는 부름. 이름을 **점까지 이어** 본 뒤 모양으로 가른다 —
# `print` 하나만 세면 `sys.stdout.write(...)` 로 같은 구멍을 다른 철자로 다시 뚫을 수
# 있고, `import sys as _s` 한 줄이면 `sys.` 라는 글자마저 사라진다. 그래서 이름 전체를
# 놓고 '표준 출력·표준 오류에 쓰는 부름인가' 를 묻는다.
_SPEAK_NAMES = ('print', 'exit', 'quit', 'sys.exit', 'SystemExit', 'warnings.warn')
_STD = re.compile(r'(?:^|\.)(stdout|stderr)(?:\.|$)')


def _speaks(who):
    """이 이름이 사람 앞에 글자를 내놓는가.

    닿지 않는 자리가 하나 남는다 — `w = sys.stdout.write` 처럼 **부름을 이름에 담아**
    부르는 경우다. 그건 이 파일들에 없고, 생기면 그 대입 자체가 눈에 띈다. 여기서
    막는 것은 '한 줄로 슬쩍 끼워 넣는' 모양이다.
    """
    if not who:
        return False
    if who in _SPEAK_NAMES:
        return True
    head, _, last = who.rpartition('.')
    return last in ('write', 'writelines') and bool(_STD.search(head))

# config 가 한 벌만 갖는 규칙. 이 이름을 이 파일들 안에서 **어떤 방법으로든** 새로
# 만들면 두 판이 조용히 갈린다 — def 든, 람다 대입이든, import … as 든.
_CONFIG_RULES = ('as_file', 'as_dir', 'env_flag', 'env_mode')


def _dotted(node):
    """`print` · `sys.stdout.write` 처럼 점으로 이은 이름. 이름이 아니면 None."""
    import ast
    parts = []
    while isinstance(node, ast.Attribute):
        parts.append(node.attr)
        node = node.value
    if not isinstance(node, ast.Name):
        return None
    parts.append(node.id)
    return '.'.join(reversed(parts))


def _from_catalog(node):
    """이 식이 카탈로그에서 나온 말인가 — `T(...)` 이거나, 그것들로만 이어졌는가.

    **문자열 상수는 거짓이다.** 그것이 곧 영어 하드코딩이다. 공백뿐인 상수('\\n' 등)와
    f-string 의 `{값}` 자리는 말이 아니라 서식·값이므로 참으로 친다.
    """
    import ast
    if isinstance(node, ast.Call):
        return _dotted(node.func) == 'T'
    if isinstance(node, ast.BinOp):
        return _from_catalog(node.left) and _from_catalog(node.right)
    if isinstance(node, ast.IfExp):
        return _from_catalog(node.body) and _from_catalog(node.orelse)
    if isinstance(node, ast.JoinedStr):
        return all(_from_catalog(v) for v in node.values)
    if isinstance(node, ast.FormattedValue):
        return True
    if isinstance(node, ast.Constant):
        v = node.value
        if isinstance(v, bytes):        # sys.stdout.buffer.write(b'…') 도 글자다
            v = v.decode('utf-8', 'replace')
        return not isinstance(v, str) or not v.strip()
    return False


@case('build: every line the builders put in front of a person goes through the catalog')
def _(work):
    # 15R — `build_html.py` 에 `try: from config import as_file / except ImportError:`
    # 로 제 사본을 둔 폴백이 있었다. 이 트리에서는 **도달 불가능한 죽은 코드**였고,
    # 도달했다면 카탈로그를 안 거친 영어 문자열을 뱉어 `ERD_LANG=ko` 에서도 영어가
    # 나왔다 — 다국어 원칙의 유일한 구멍이었다. 죽은 코드는 시험이 실행으로는 못
    # 잡으므로 소스에서 센다.
    #
    # 16R — 그 소스 스캔이 **정규식**이라 같은 구멍을 다른 철자로 다시 뚫으면 아무
    # 말도 안 했다. 독립 검증자가 셋을 통과시켰다:
    #   M15  sys.stdout.write('note: …')   — `print` 만 찾으므로 안 걸린다
    #   M16  print ('warning: …')          — 괄호 앞 공백 하나로 정규식이 빗나간다
    #   M17  as_file = lambda …            — `def` 만 찾으므로 안 걸린다. 그리고 이건
    #        죽은 코드가 아니라 **살아 있는 우회**였다: ERD_LANG=ko 에서 실제로
    #        `ERD_HTML_OUT must name a file, not a directory: …` 라는 영어가 샜다.
    # 철자를 세는 대신 **구문을 센다** — 파이썬이 그 코드를 읽는 방식 그대로.
    import ast
    spoke = 0
    for name in BUILDERS:
        src = (HERE / name).read_text(encoding='utf-8')
        tree = ast.parse(src, filename=name)

        for node in ast.walk(tree):
            if not isinstance(node, ast.Call):
                continue
            who = _dotted(node.func)
            if not _speaks(who):
                continue
            spoke += 1
            for arg in node.args:            # sep=·end=·file= 은 말이 아니다
                if not _from_catalog(arg):
                    raise Fail(
                        f'{name}:{node.lineno}: {who}() is handed '
                        f'{ast.unparse(arg)!r}, not a catalog lookup — a line that '
                        f'skips the catalog stays English under ERD_LANG=ko')

        # 경로·플래그 규칙은 config 에 한 벌뿐이다. 사본을 두면 두 판이 조용히 갈린다.
        made = []
        for node in ast.walk(tree):
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                if node.name in _CONFIG_RULES:
                    made.append(f'def {node.name} (line {node.lineno})')
            elif isinstance(node, ast.Assign):
                for t_ in node.targets:
                    if isinstance(t_, ast.Name) and t_.id in _CONFIG_RULES:
                        made.append(f'{t_.id} = … (line {node.lineno})')
            elif isinstance(node, (ast.AnnAssign, ast.NamedExpr)):
                if isinstance(node.target, ast.Name) and node.target.id in _CONFIG_RULES:
                    made.append(f'{node.target.id} := … (line {node.lineno})')
            elif isinstance(node, (ast.Import, ast.ImportFrom)):
                mod = getattr(node, 'module', None)
                for a in node.names:
                    bound = a.asname or a.name.split('.')[0]
                    if bound in _CONFIG_RULES and mod != 'config':
                        made.append(f'import … as {bound} (line {node.lineno})')
        eq(made, [], f'{name} makes its own {"/".join(_CONFIG_RULES)}: {made}')

    if spoke < 8:                    # 지금 열 자리다 — 스캔이 눈을 감으면 여기가 먼저 운다
        raise Fail(f'only {spoke} speaking sites found across the three builders — the '
                   f'scan is not looking at what it thinks it is')


@case('build: the AST scan reads the shape, not one spelling of it')
def _(work):
    # 16R — 위 케이스가 정말 **구문**을 보는지를 잰다. 정규식 판이 통과시킨 세 우회를
    # 소스에 먹여 보고, 하나라도 통과하면 여기가 빨강이다. 재는 자를 재는 자리다 —
    # 이 파일의 기록이 세 번 이름 붙인 '재는 손이 만든 것을 재는 손이 채점한다' 를
    # 이 스캔에도 댄다. (실제 파일은 손대지 않는다. 소스 문자열만 만들어 먹인다.)
    import ast
    base = (HERE / 'build_html.py').read_text(encoding='utf-8')
    bypasses = {
        'M15 sys.stdout.write': "\nsys.stdout.write('note: something happened')\n",
        'M15 through an aliased sys':
            "\nimport sys as _s\n_s.stdout.write('note: something happened')\n",
        'M15 through sys.stdout.buffer':
            "\nimport sys\nsys.stdout.buffer.write(b'note')\n",
        'M16 print with a space': "\nprint ('warning: something happened')\n",
        'M17 lambda instead of def':
            "\nas_file = lambda path, env: path\n",
        'sys.exit with a message': "\nsys.exit('cannot write there')\n",
        'f-string with English in it':
            "\nprint(f'wrote {1} files')\n",
    }
    for what, extra in bypasses.items():
        tree = ast.parse(base + extra)
        caught = False
        for node in ast.walk(tree):
            if isinstance(node, ast.Call):
                who = _dotted(node.func)
                if _speaks(who) and not all(_from_catalog(a) for a in node.args):
                    caught = True
            if isinstance(node, ast.Assign):
                for t_ in node.targets:
                    if isinstance(t_, ast.Name) and t_.id in _CONFIG_RULES:
                        caught = True
        if not caught:
            raise Fail(f'the scan lets {what} through — the same hole spelled another '
                       f'way is the same hole')
    # 그리고 지금 소스는 이 자를 통과해야 한다 (자가 무조건 빨강이면 재는 것이 없다)
    tree = ast.parse(base)
    for node in ast.walk(tree):
        if isinstance(node, ast.Call):
            who = _dotted(node.func)
            if _speaks(who):
                for a in node.args:
                    if not _from_catalog(a):
                        raise Fail(f'build_html.py:{node.lineno}: {ast.unparse(a)!r}')


# ── 11. ERD_STALE ───────────────────────────────────────────────────────────

@case('errors: ERD_STALE follows the one yes/no rule, plus its own mode word')
def _(work):
    # 15R — `in ('warn','ok','yes','1')` 라는 제 규칙이 남아 있었다. 14라운드가 다른
    # 다섯 자리를 `env_flag` 로 묶을 때 여기만 빠져서, `ERD_STALE=true`·`on`·`y` 는
    # 켰다고 생각한 사람에게 **조용히 안 먹었다** — 문서가 안 나오고 멈춘다. 오타도
    # 조용히 멈췄다. 이건 불리언이 아니라 모드이므로 'warn' 은 계속 알아듣는다.
    write_schema(work, {'t': table('t', [col('id')], pk=['id'])})
    run('build_erd.py', work)
    # 그림을 그린 뒤 스키마를 더 새것으로 만든다 — 이제 박으려는 그림이 오래됐다
    ref = (work / 'schema.json')
    later = max(p.stat().st_mtime for p in (work / 'out').glob('*.png')) + 5
    os.utime(ref, (later, later))
    if not run('build_html.py', work, expect_ok=False).returncode:
        raise Fail('a stale diagram no longer stops the document — this case would '
                   'measure nothing')

    for value in ('warn', 'ok', 'true', 'on', 'y', '1', 'yes'):
        r = run('build_html.py', work, env={'ERD_STALE': value}, expect_ok=False)
        if r.returncode != 0:
            raise Fail(f'ERD_STALE={value!r} was meant to let the build through but it '
                       f'stopped:\n{(r.stdout + r.stderr)[-400:]}')
        if '[warn]' not in r.stdout:
            raise Fail(f'ERD_STALE={value!r} passed silently — passing is allowed, '
                       f'passing without saying so is not:\n{r.stdout!r}')
    for value in ('0', 'off', 'no', 'false'):
        r = run('build_html.py', work, env={'ERD_STALE': value}, expect_ok=False)
        if r.returncode == 0:
            raise Fail(f'ERD_STALE={value!r} reads as off but let a stale diagram in')
    # 오타는 조용히 어느 쪽으로도 가지 않는다 — 이름을 대고 기본값(멈춤)으로 간다
    r = run('build_html.py', work, env={'ERD_STALE': 'wran'}, expect_ok=False)
    if r.returncode == 0:
        raise Fail('a typo in ERD_STALE was taken as a yes')
    has(r.stdout + r.stderr, 'ERD_STALE',
        'an unreadable value names the variable that has to change')


# ── 12. 여러 DB 를 합치는 자리 ──────────────────────────────────────────────

def merge_run(work, labels, ok=True):
    """merge_schemas.py 는 라벨을 argv 로 받는다 — run() 이 인자를 붙일 자리가 없다."""
    drv = work / f'merge_{"_".join(labels) or "none"}.py'
    drv.write_text(f'import sys\nsys.path.insert(0, {str(HERE)!r})\n'
                   f'sys.argv = ["merge_schemas.py"] + {labels!r}\n'
                   'import merge_schemas\nmerge_schemas.main()\n', encoding='utf-8')
    return run(str(drv), work, expect_ok=ok)


@case('merge_schemas: two databases that share a table name stay two tables')
def _(work):
    # 16R — 이 파일은 15·16라운드를 통틀어 아무에게도 배정된 적이 없다. 모듈 docstring
    # 은 "테이블 키는 `<라벨>.<테이블명>` 이라 이름이 같아도 부딪히지 않는다" 고
    # 적는데, 라벨을 붙이는 코드는 introspect.py 에만 있었다. `merged.update(part)`
    # 한 줄이라, parse_ddl 이 만든 파일이나 라벨 없이 뽑아 둔 예전 파일을 섞으면
    # 같은 이름의 테이블이 **말 한 마디 없이** 하나로 합쳐졌다:
    #     shop  tables 1 · mart  tables 1  →  total tables 1
    # 3라운드 발견(같은 이름 테이블이 하나로 합쳐짐)이 다중 DB 경로에 그대로 있었다.
    work.mkdir(parents=True, exist_ok=True)
    (work / 'schema.shop.json').write_text(json.dumps({
        'orders': table('orders', [col('id')], pk=['id']),
        'items': table('items', [col('id'), col('order_id')], pk=['id'],
                       fks=[{'column': 'order_id', 'ref_table': 'orders',
                             'ref_column': 'id', 'on_delete': 'CASCADE'}])}),
        encoding='utf-8')
    (work / 'schema.mart.json').write_text(json.dumps({
        'orders': table('orders', [col('id'), col('n')])}), encoding='utf-8')

    r = merge_run(work, ['shop', 'mart'])
    merged = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq(sorted(merged), ['mart.orders', 'shop.items', 'shop.orders'],
       'a table name that exists in two databases gets one key per database')

    # 찍힌 줄이 파일과 같은 말을 해야 한다. 예전엔 부분 둘이 각각 1 이라고 찍힌
    # 다음 줄에 합계가 1 이라고 찍혔다 — 화면만 봐도 어긋나 있었는데 아무도 안 봤다.
    parts = [int(n) for n in re.findall(r'tables\s+(\d+) ·', r.stdout)]
    total = re.search(r'total\s+tables (\d+) · columns (\d+) · FKs (\d+)', r.stdout)
    if not total or len(parts) < 3:
        raise Fail(f'merge_schemas printed no per-part/total tally — this case would '
                   f'measure nothing:\n{r.stdout!r}')
    eq(int(total.group(1)), sum(parts[:-1]),
       'the total is the sum of the parts — a table swallowed by a name clash used to '
       'vanish between those two lines')
    eq(int(total.group(1)), len(merged), 'and the total is what the file holds')

    # FK 는 **키를 가리키는 값**이라 키를 옮기면 같이 옮겨야 한다. 안 옮기면 합친 뒤
    # '대상에 없는 FK' 로 몰려 조용히 버려진다 — 관계가 사라진 ERD 가 나온다.
    eq([fk['ref_table'] for fk in merged['shop.items']['fks']], ['shop.orders'],
       'a foreign key follows its parent to the new key')
    eq(int(total.group(3)), 1, 'and it is still counted as a relationship')
    # db 칸도 갈라야 한다 — 키만 갈라 놓으면 HTML 이 두 DB 를 한 묶음으로 싣는다
    eq({k: v['db'] for k, v in merged.items()},
       {'shop.orders': 'shop', 'shop.items': 'shop', 'mart.orders': 'mart'},
       'each table says which database it came from')

    # introspect 가 만든 파일(이미 라벨이 붙은 키)은 두 번 붙지 않는다
    w2 = work.parent / 'already'
    w2.mkdir(parents=True, exist_ok=True)
    (w2 / 'schema.shop.json').write_text(json.dumps({
        'shop.orders': table('orders', [col('id')], pk=['id'], db='shop')}),
        encoding='utf-8')
    (w2 / 'schema.mart.json').write_text(json.dumps({
        'mart.orders': table('orders', [col('id')], db='mart')}), encoding='utf-8')
    merge_run(w2, ['shop', 'mart'])
    eq(sorted(json.loads((w2 / 'schema.json').read_text(encoding='utf-8'))),
       ['mart.orders', 'shop.orders'],
       'a file introspect already labelled keeps its keys — no shop.shop.orders')

    # 같은 라벨을 두 번 적어도 읽을 파일은 하나다 (부분 줄과 합계가 어긋나지 않는다)
    r = merge_run(work, ['shop', 'shop'])
    eq(len(re.findall(r'^  shop', r.stdout, re.M)), 1,
       'the same label twice reads the same file once')


# ── 13. import 는 아무것도 안 만든다 ────────────────────────────────────────

@case('build: importing a builder does not create anything in the caller cwd')
def _(work):
    # 15R 이 config·i18n·parse_ddl·introspect 에 세운 규칙인데, 16R 독립 검증자가
    # 재 보니 **다섯 자리가 남아 있었다** (빈 디렉터리에서):
    #     merge_schemas → ./erd-build
    #     erd / build_erd / build_html / build_docx → ./erd-build  ./erd-build/out
    # 원인은 `from config import OUT/WORK/SCHEMA_JSON` 이다. 그 이름들은 PEP 562
    # 모듈 __getattr__ 로 늦춰져 있어서, **가져오는 그 순간** 값이 만들어지고 mkdir 이
    # 돈다. 15R 의 회귀 케이스(selftest_r14_config.py)는 네 파일만 봤다.
    #
    # 여기서는 두 방향으로 잰다.
    #   ① 실행: merge_schemas 는 erd 를 안 거치므로 통째로 재현할 수 있다.
    #   ② 실행(대역): 세 builder 는 `import erd` 가 그 자체로 schema.json 을 읽는
    #      모듈이라, erd 를 대역으로 세우고 **그 파일 자신이** 무엇을 만드는지 본다.
    #      erd.py 는 이 라운드의 다른 사람 몫이라 여기서 고칠 수 없다 — 남은 오염이
    #      누구 것인지 이 케이스가 갈라 준다.
    #   ③ 소스: 늦춰 둔 이름을 다시 `from config import …` 로 가져오면 빨강.
    import ast
    import subprocess
    import config as _cfg

    lazy = sorted(_cfg._LAZY)
    if not lazy:
        raise Fail('config no longer defers any value — this case would measure nothing')

    probe = ('import os, sys, types\n'
             'from pathlib import Path\n'
             f'sys.path.insert(0, {str(HERE)!r})\n'
             'if {stub}:\n'
             "    s = types.ModuleType('erd')\n"
             '    s.AREAS = []; s.AREA_NAME = {{}}; s.AREA_SCHEMA = {{}}; s.LAYERS = {{}}\n'
             '    s.ROLE = {{}}; s.SCHEMA = {{}}; s.SPEC = {{}}; s.DERIVES = []\n'
             "    s.layer = lambda k: ''\n"
             "    s.badge = lambda k: ('', '')\n"
             "    s.col_role = lambda t, c: ''\n"
             "    s.OUT = Path('/nonexistent-erd-stub')\n"
             "    sys.modules['erd'] = s\n"
             'import {mod}\n'
             "print(sorted(p for p in os.listdir('.')))\n")

    def made(mod, stub):
        d = work.parent / f'cwd_{mod}_{int(stub)}'
        d.mkdir(parents=True, exist_ok=True)
        env = {k: v for k, v in os.environ.items() if not k.startswith('ERD_')}
        r = subprocess.run([sys.executable, '-c', probe.format(mod=mod, stub=stub)],
                           capture_output=True, text=True, encoding='utf-8', cwd=str(d), env=env)
        if r.returncode != 0:
            raise Fail(f'importing {mod} (stub erd={stub}) failed — this case cannot '
                       f'measure anything until the stub carries what {mod} imports '
                       f'from erd:\n{r.stderr[-600:]}')
        return sorted(p.name for p in d.iterdir())

    eq(made('merge_schemas', False), [],
       'importing merge_schemas must not build a directory tree in the caller cwd — '
       'it used to leave ./erd-build behind')

    for mod in ('build_erd', 'build_html', 'build_docx'):
        eq(made(mod, True), [],
           f'with erd stood in for, importing {mod} creates nothing of its own — what '
           f'is left in a real run comes from erd.py reading schema.json at import')

    for name in BUILDERS + ('merge_schemas.py',):
        tree = ast.parse((HERE / name).read_text(encoding='utf-8'), filename=name)
        eager = [f'{a.name} (line {n.lineno})'
                 for n in ast.walk(tree) if isinstance(n, ast.ImportFrom)
                 and n.module == 'config' for a in n.names if a.name in lazy]
        eq(eager, [],
           f'{name} pulls a deferred config value in at import time: {eager}\n'
           f'      (deferred: {", ".join(lazy)} — ask for them as config.X where they '
           f'are used)')


# ── 14. 사람이 적은 글자가 마크업이 되지 않는다 ────────────────────────────

_INJECT = '<script>alert(1)</script>'


def injected(work):
    """사람이 손으로 적는 자리마다 태그를 심은 스키마 + spec."""
    write_schema(work, {'shop.t': table(
        f'tbl{_INJECT}',
        [col(f'c{_INJECT}', f'text{_INJECT}', default=f'd{_INJECT}',
             comment=f'cm{_INJECT}'),
         col('plain')],
        db='shop', pk=[f'c{_INJECT}'], uniques=[[f'c{_INJECT}']],
        checks=[{'name': f'ck{_INJECT}', 'def': f'cd{_INJECT}'}],
        indexes=[{'name': 'ix', 'def': f'ix{_INJECT}'}],
        note=f'note{_INJECT}')})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', f'area{_INJECT}', 'public', ['shop.t']]],
        'roles': {'shop.t': f'role{_INJECT}'},
        'doc': {'title': f'title{_INJECT}',
                'intro': '<b>markup on purpose</b>',
                'meta': [[f'mk{_INJECT}', f'mv{_INJECT}', '', '']],
                'area_desc': {'A': f'ad{_INJECT}'},
                'db_names': {'shop': f'db{_INJECT}'},
                'mapping': [[f'm{i}{_INJECT}' for i in range(5)]],
                'open_items': [[f'o{i}{_INJECT}' for i in range(5)]],
                'mapping_intro': f'mi{_INJECT}',
                'mapping_note': f'mn{_INJECT}',
                'open_note': f'on{_INJECT}'}}), encoding='utf-8')


@case('html: nothing a person wrote can turn into markup')
def _(work):
    # 16R — `escape()` 를 쓰는 자리가 스물 몇 개인데 그것을 지키는 케이스가 **0건**
    # 이었다. 독립 검증자 M21: `<h1>{escape(TITLE)}</h1>` → `<h1>{TITLE}</h1>` 로
    # 바꿔도 161개가 전부 초록이었다.
    #     PRISTINE  <h1>&lt;script&gt;alert(1)&lt;/script&gt;</h1>
    #     M21       <h1><script>alert(1)</script></h1>
    # `doc.title` 은 사용자가 손으로 쓰는 값이다. 한 자리만 재면 나머지 스물이 그대로
    # 무방비니, **사람이 적는 자리 전부**에 같은 글자를 심고 문서 전체에서 센다.
    injected(work)
    run('merge_desc.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')

    # 문서 안에서 진짜 <script> 는 확대 스크립트 하나뿐이다. 어느 자리에서 escape 가
    # 빠지든 이 수가 늘어난다 — 자리마다 케이스를 하나씩 두지 않아도 전부 걸린다.
    eq(html.count('<script>'), 1,
       'the only <script> in the document is the one the builder wrote itself — a '
       'field that skips escape() puts a second one in')
    eq(html.count('</script>'), 1, 'and it closes exactly once')
    for bad in ('<b>bold-from-a-column', 'onerror='):
        if bad in html:
            raise Fail(f'{bad!r} reached the document as markup')

    # 그리고 심은 것이 정말 실렸는지 — 안 실렸으면 위 셈은 아무것도 안 잰 것이다.
    safe = escape(_INJECT)
    # (`note` 는 뺀다 — HTML 은 역할명이 있으면 테이블 코멘트를 안 싣는다. 여기
    #  테이블은 roles 를 갖고 있으므로 그 자리는 role 이 차지한다.)
    planted = ('title', 'tbl', 'c', 'text', 'd', 'cm', 'area', 'ad', 'db', 'role',
               'mk', 'mv', 'm0', 'o0', 'mi', 'mn', 'on', 'cd', 'ix')
    missing = [p for p in planted if f'{p}{safe}' not in html]
    eq(missing, [],
       'every hand-written field must actually reach the document escaped — a field '
       'that never arrives cannot prove anything about escaping')

    # **선언된 예외 하나**: doc.intro 는 일부러 날 HTML 이다. 함께 배포하는
    # examples/full.spec.json 의 intro 가 <b> 를 쓴다. 예외를 여기 적어 두지 않으면
    # 다음 사람이 '여기가 빠졌다' 고 읽거나, 반대로 다른 자리를 열어도 같아 보인다.
    has(html, '<div class="cmt"><b>markup on purpose</b></div>',
        'doc.intro is the one field that is deliberately raw markup')


# ── 15. 문서의 낱말은 전부 카탈로그에서 나온다 ─────────────────────────────

def _labels(html):
    """문서가 스스로 붙인 말 — 표 머리칸과 (색 없는) 배지."""
    th = [re.sub(r'<[^>]+>', '', x) for x in re.findall(r'<th[^>]*>(.*?)</th>', html)]
    badge = re.findall(r'<span class="badge">([^<]*)</span>', html)
    return th, badge


@case('html: every label the document prints comes out of the catalog')
def _(work):
    # 16R — `build_html.py` 만 카탈로그를 안 거치는 자리가 셋이었고, **재는 케이스가
    # 0개**였다 (검증자 뮤턴트 X-htmlnull: `Null`→`ZZZ` → all 161 passed):
    #     build_html.py:233  badges = [f'cols {len(t["columns"])}']   → 네 언어 모두 `cols 3`
    #     build_html.py:259  "<th …>Null</th>"                        → 네 언어 모두 `Null`
    #     build_html.py:405  f'…{len(tables)} tables…'                → 네 언어 모두 `3 tables`
    # docx 는 같은 것을 `T('word.tables')`·`T('col.name')` 로 낸다 — 한 사람이 받은
    # 두 문서가 같은 칸을 다른 말로 불렀다.
    #
    # 낱말을 하나씩 못박는 대신 **원칙을 잰다**: 같은 문서를 두 말로 뽑아, 자리마다
    # 말이 실제로 바뀌는지 본다. 안 바뀌는 것은 아래 목록에 이름이 있어야 한다.
    write_schema(work, {'t': table('t', [col('id'), col('name', 'text')], pk=['id'])})
    run('merge_desc.py', work)
    run('build_html.py', work)
    en_th, en_badge = _labels((work / 'T.html').read_text(encoding='utf-8'))
    run('build_html.py', work, env={'ERD_LANG': 'ko'})
    ko_th, ko_badge = _labels((work / 'T.html').read_text(encoding='utf-8'))

    if len(en_th) < 10 or not en_badge:
        raise Fail(f'the document carries {len(en_th)} header cells and '
                   f'{len(en_badge)} badges — this case would measure nothing')
    eq(len(ko_th), len(en_th), 'the same document in another language, cell for cell')
    eq(len(ko_badge), len(en_badge), 'and badge for badge')

    # 말이 바뀌지 않아도 되는 것: 표기·SQL 낱말, 그리고 **네 언어가 같은 낱말을 고른
    # 자리**. 이 목록은 '카탈로그를 안 거친 것' 이 아니라 '거쳤는데도 같은 것' 까지
    # 함께 봐준다는 뜻이다 — 이 케이스가 재는 것은 en↔ko 차이지 T() 호출 여부가
    # 아니므로 둘을 여기서 갈라낼 수는 없다.
    #
    # 'Null' 이 그 경우다. 16R 주석은 "`col.null` 키가 들어오면 여기서 'Null' 을
    # 지우면 된다" 고 적었는데 **그것은 틀렸다.** 17R 에 키가 들어왔지만
    # (build_html.py:274 가 `T('col.null')` 을 부른다) ko 카탈로그의 값이 en 과
    # 같은 'Null' 이라(lang/ko.py:51) 지우면 이 케이스가 빨개진다. 키가 생겼다고
    # 지우는 것이 아니라 **ko 값이 다른 낱말이 됐을 때** 지운다.
    # 다른 낱말로 바뀌면(뮤턴트 X-htmlnull 처럼) 목록에 없으므로 여기가 빨강이다.
    SAME_OK = {'#', 'DB', 'FK', 'Null'}
    for a, b in zip(en_th, ko_th):
        if a == b and a not in SAME_OK:
            raise Fail(f'the header cell {a!r} is the same word in en and ko — it '
                       f'never went through the catalog (allowed as-is: '
                       f'{sorted(SAME_OK)})')
    # 배지는 '낱말 + 수' 라 수를 뺀 낱말만 본다
    for a, b in zip(en_badge, ko_badge):
        wa, wb = re.sub(r'[\d,\s≈]+', '', a), re.sub(r'[\d,\s≈]+', '', b)
        if wa and wa == wb:
            raise Fail(f'the badge word {wa!r} is the same in en and ko — badges used '
                       f'to read `cols 3` and `3 tables` in all four languages')
    # 그리고 docx 가 쓰는 그 낱말이어야 한다 — 두 문서가 같은 칸을 같은 말로 부른다
    import i18n
    for want in (i18n.t('word.tables'), i18n.t('word.columns'), i18n.t('col.name'),
                 i18n.t('col.type'), i18n.t('col.desc')):
        if want not in en_th + [re.sub(r'[\d,\s≈]+$', '', x) for x in en_badge]:
            raise Fail(f'the document never uses the catalog word {want!r} that docx '
                       f'uses for the same cell')


# ── 16. 영역 id 는 (영역, DB) 쌍에서만 나온다 — 두 방향 ────────────────────

@case('html: an area id is decided by the area and the db together')
def _(work):
    # 16R — 기존 케이스(`an anchor is decided by the table key alone`)는 **한 영역이
    # 여러 DB** 인 방향만 봤다. 반대 방향을 아무도 안 밟아서, 검증자 뮤턴트 M14
    # (`_ident('area_', db or code)`) 가 초록이었다:
    #     pristine → ['area_A_shop-4e0faf3d1a90', 'area_B_shop-f72e04bfa85b']  중복 0
    #     M14      → ['area_shop', 'area_shop']  중복 id 1개, 모호한 href 1개
    # 목차의 두 영역 링크가 둘 다 첫 절로 간다 — 문서가 조용히 남의 영역을 가리킨다.
    write_schema(work, {
        'shop.a1': table('a1', [col('id')], pk=['id'], db='shop'),
        'shop.b1': table('b1', [col('id')], pk=['id'], db='shop')})
    (work / 'erd.spec.json').write_text(json.dumps({
        'areas': [['A', 'First', 'public', ['shop.a1']],
                  ['B', 'Second', 'public', ['shop.b1']]]}), encoding='utf-8')
    run('merge_desc.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')

    ids = re.findall(r'<h3 id="([^"]+)"', html)
    eq(len(ids), 2, f'two areas, two sections — got {ids}')
    eq(len(set(ids)), 2,
       f'two areas inside one database still get two different ids — got {ids}')
    all_ids = re.findall(r'\sid="([^"]+)"', html)
    eq(sorted({i for i in all_ids if all_ids.count(i) > 1}), [],
       'no id appears twice anywhere in the document')
    for href in sorted(set(re.findall(r'href="#([^"]+)"', html))):
        eq(all_ids.count(href), 1, f'href #{href} lands on exactly one id')


# ── 17. 표지 마지막 행을 언제 합치는가 ──────────────────────────────────────

def _cover(work):
    from docx import Document
    return Document(str(work / 'T.docx')).tables[0]


@case('docx: the cover merges its last row only when the right-hand pair is empty')
def _(work):
    # 16R — 이 조건은 `if True:` 에서 `if meta and not str(meta[-1][3]).strip():` 로
    # 바뀌었는데 **보고에 행이 없고 지키는 케이스도 없었다.** 동작은 실제로 갈린다:
    # spec 이 doc.meta 를 안 주면 기본 표지의 마지막 행이 `Columns | 5 | FKs | 0` 인데
    # 늘 합치면 셋이 한 칸이 되어 `5 FKs 0` 이 된다 — 표지가 컬럼 수와 FK 수를
    # 말한다면서 둘을 한 덩어리로 뭉갠 셈이다.
    write_schema(work, {'t': table('t', [col('id'), col('a')], pk=['id'])})
    run('merge_desc.py', work)
    run('build_docx.py', work)
    last = _cover(work).rows[-1]
    cells = [c.text for c in last.cells]
    if 'FKs' not in cells:
        raise Fail(f'the default cover no longer states FKs — this case would measure '
                   f'nothing: {cells}')
    eq(len(set(id(c._tc) for c in last.cells)), 4,
       f'the default cover keeps four separate cells — merging them reads as one '
       f'value: {cells}')
    eq(cells, ['Columns', '2', 'FKs', '0'], 'and it says both numbers')

    # 반대쪽: 오른쪽 쌍이 비어 있으면 왼쪽 내용이 그 자리까지 편다
    (work / 'erd.spec.json').write_text(json.dumps({
        'doc': {'meta': [['Date', '2026-01-01', 'Author', ''],
                         ['Note', 'runs the width of the row', '', '']]}}),
        encoding='utf-8')
    run('build_docx.py', work)
    last = _cover(work).rows[-1]
    if len(set(id(c._tc) for c in last.cells)) != 2:
        raise Fail(f'a last row whose right-hand pair is empty must spread the left '
                   f'content across it: {[c.text for c in last.cells]}')


# ── 18. 손으로 적은 머리말이 비었을 때 ─────────────────────────────────────

@case('html+docx: an empty hand-written intro falls back to the same words')
def _(work):
    # 16R — 마지막 조각이었다.
    #     build_html.py:423  DOC.get('mapping_intro') or T('docx.ch6_intro')
    #     build_docx.py:399  DOC.get('mapping_intro', T('docx.ch6_intro'))
    # `"mapping_intro": ""` 를 적으면 docx 는 빈 문단, HTML 은 카탈로그 문구였다 —
    # 같은 spec 을 준 사람이 두 문서에서 다른 6장을 받는다. 규칙을 두 builder 가 함께
    # import 하는 자리(build_erd.doc_text)에 한 벌만 뒀다.
    from docx import Document
    write_schema(work, {'t': table('t', [col('id')], pk=['id'])})

    def build(intro):
        doc = {'mapping': [['1', 'proposed', 't', 'yes', 'why']]}
        if intro is not None:
            doc['mapping_intro'] = intro
        (work / 'erd.spec.json').write_text(json.dumps({
            'areas': [['A', 'All', 'public', ['t']]], 'doc': doc}), encoding='utf-8')
        run('build_docx.py', work)
        run('build_html.py', work)
        words = '\n'.join(p.text for p in Document(str(work / 'T.docx')).paragraphs)
        return (work / 'T.html').read_text(encoding='utf-8'), words

    run('merge_desc.py', work)
    import i18n
    default = i18n.t('docx.ch6_intro')

    html, words = build('a hand-written lead-in')
    has(html, 'a hand-written lead-in', 'HTML uses what the spec wrote')
    has(words, 'a hand-written lead-in', 'and docx uses the same thing')

    for empty in ('', '   '):
        html, words = build(empty)
        has(html, escape(default),
            f'HTML falls back to the catalog for mapping_intro={empty!r}')
        has(words, default,
            f'docx falls back to the same words for mapping_intro={empty!r} — it used '
            f'to leave an empty paragraph while HTML printed the catalog text')

    html, words = build(None)
    has(html, escape(default), 'and with no key at all, both still agree (HTML)')
    has(words, default, 'and with no key at all, both still agree (docx)')


# ── 19. 신선도 게이트와 '내용이 같으면 안 쓴다' 의 경계 ────────────────────

@case('errors: a re-run that changes nothing does not turn the diagrams stale')
def _(work):
    # 16R — 보고에 없던 변경 둘 중 하나: `merge_desc.py` 는 내용이 같으면 파일을 다시
    # 쓰지 않는다. `build_erd.require_fresh` 는 **mtime 으로** 판정하므로 그 한 줄이
    # 신선도 게이트의 동작을 바꾼다. 둘의 경계에 케이스가 없었다.
    #
    # 재는 것은 두 방향이다. ① 아무것도 안 바뀐 재실행이 멀쩡한 그림을 낡은 것으로
    # 만들지 않는다 (그러면 문서가 안 나오고 멈춘다). ② 정말 바뀌면 멈춘다 — 게이트가
    # 살아 있어야 ①이 '아무것도 안 재는 통과' 가 아니다.
    write_schema(work, {'t': table('t', [col('id'), col('name', 'text')], pk=['id'])})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    before = (work / 'schema.json').stat().st_mtime

    run('merge_desc.py', work)
    after = (work / 'schema.json').stat().st_mtime
    if after != before:
        raise Fail('a re-run that changes nothing rewrote schema.json — every diagram '
                   'is now older than the schema, and the documents stop. (the two '
                   'sides of this boundary are merge_desc.py and build_erd.'
                   'require_fresh)')
    r = run('build_html.py', work, expect_ok=False)
    if r.returncode != 0:
        raise Fail(f'the document stopped although nothing changed:\n'
                   f'{(r.stdout + r.stderr)[-400:]}')

    # 게이트는 살아 있다 — 스키마가 정말 새것이 되면 멈춘다
    later = max(p.stat().st_mtime for p in (work / 'out').glob('*.png')) + 5
    os.utime(work / 'schema.json', (later, later))
    r = run('build_html.py', work, expect_ok=False)
    if r.returncode == 0:
        raise Fail('a schema newer than every diagram no longer stops the document — '
                   'the first half of this case would then measure nothing')


# ── 17R. 라벨을 붙인 두 DB 가 문서에서 다시 한 덩어리가 되는가 ───────────────

@case('merge_schemas: two labelled databases do not collapse into one area')
def _(work):
    # 17R 뮤테이션. `label_part` 가 키와 `db` 는 갈라 놓으면서 `schema` 칸은 그대로
    # 두던 동안, 두 DB 의 테이블이 모두 `public` 하나가 되어 **영역 자동 분류가 두
    # DB 를 한 덩어리로 다시 묶었다** — 키에서 막은 합쳐짐이 그림에서 다시 일어났다.
    # 재는 자리를 파일 칸이 아니라 **그려진 도판**에 둔다: 영역 상세도가 DB 마다
    # 하나씩 나와야 한다.
    work.mkdir(parents=True, exist_ok=True)
    (work / 'schema.shop.json').write_text(json.dumps({
        'orders': table('orders', [col('id')], pk=['id']),
        'items': table('items', [col('id')], pk=['id'])}), encoding='utf-8')
    (work / 'schema.mart.json').write_text(json.dumps({
        'orders': table('orders', [col('id'), col('n')]),
        'facts': table('facts', [col('id')])}), encoding='utf-8')
    merge_run(work, ['shop', 'mart'])
    r = run('build_erd.py', work)
    areas = [n for n in drawn_names(r.stdout) if n.startswith('erd_area_')]
    if len(areas) != 2:
        raise Fail(f'two databases were drawn as {len(areas)} area diagram(s) — '
                   f'{areas}\n{r.stdout}')
    for label in ('shop', 'mart'):
        has(r.stdout, label, f'the area diagrams must still name {label}')


@case('html: the number in a figure caption is the number drawn inside that figure')
def _(work):
    # 17R 뮤테이션. 캡션을 다는 쪽이 번호를 **실린 차례대로** 다시 세면, 부록으로
    # 미룬 전체 상세도에서 두 번호가 갈린다: 그림 안에 "2" 라고 그려진 도판이 "4"
    # 라는 캡션 아래에 실렸다. 어느 쪽이 옳은지는 재지 않는다 — 한 그림이 제 번호를
    # 두 개 갖지 않는 것만 잰다. 번호를 그리는 쪽과 캡션을 다는 쪽이 서로 남이라,
    # 이 대조는 어느 한쪽을 베껴 적지 않는다.
    write_schema(work, {n: table(n, [col('id')], schema=s, pk=['id'])
                        for s, n in (('shop', 'orders'), ('shop', 'items'),
                                     ('mart', 'facts'), ('mart', 'visits'))})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    figs = re.findall(r'<div class="fig">(.*?)</div>\s*<p class="figcap"><b>([^<]*)</b>',
                      html, re.S)
    if len(figs) < 3:
        raise Fail(f'the document carries {len(figs)} inline figures — this case needs '
                   f'the appendix figure to be there to measure anything')
    seen = 0
    for body, cap in figs:
        want = re.findall(r'\d+', cap)
        # 그림 안의 제목 줄은 캡션 라벨과 **같은 문자열**이라야 한다. SVG 를 인라인으로
        # 박으므로 그 줄이 본문 안에 그대로 들어 있다.
        inside = re.search(re.escape(cap.strip().lstrip('[')), body)
        if not want:
            raise Fail(f'a figure caption carries no number: {cap!r}')
        if inside is None:
            raise Fail(f'the figure captioned {cap!r} is drawn with a different number '
                       f'inside it: {re.findall(r"[^<>]*Fig[^<>]*", body)[:2]}')
        seen += 1
    eq(seen, len(figs), 'every figure was checked')

@case('artifacts: the GraphML is written in the encoding its first line declares')
def _(work):
    # 17R 뮤테이션. GraphML 의 첫 줄은 스스로를 `encoding="UTF-8"` 이라고 **선언한다**.
    # 쓰는 자리에 인코딩을 안 주면 실제로 쓰이는 것은 로케일이 정하므로 선언과
    # 알맹이가 어긋난다 — cp949 로케일이면 yEd 가 한글을 깨서 열고, ascii 로케일이면
    # 그 전에 UnicodeEncodeError 로 죽어 그림은 다 그려 놓고 GraphML 만 0바이트였다.
    write_schema(work, {'memo': table(
        'memo', [col('id'), col('body', 'text', comment='본문 설명')],
        pk=['id'], note='주문 메모')})
    r = run('build_erd.py', work,
            env={'LC_ALL': 'C', 'LANG': 'C', 'PYTHONUTF8': '0',
                 'PYTHONCOERCECLOCALE': '0'}, expect_ok=False)
    if r.returncode != 0:
        raise Fail('the locale of the shell decided whether the GraphML could be '
                   f'written:\n{(r.stdout + r.stderr)[-500:]}')
    got = (work / 'T.graphml').read_text(encoding='utf-8')
    has(got, 'encoding="UTF-8"', 'the file still says what it is written in')
    has(got, '본문 설명', 'and the text it declares it can hold is really in it')
