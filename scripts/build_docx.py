#!/usr/bin/env python3
"""ERD 문서 생성 — 어떤 DB에도 쓸 수 있게 내용은 erd.spec.json 의 doc 절에서 읽는다.

  doc.title        표지 제목            doc.subtitle   표지 부제
  doc.meta         [[구분, 내용], …]     표지 정보표
  doc.purpose      1.1 목적             doc.scope      1.2 범위
  doc.sources      [[근거, 내용], …]     1.3 산출 근거
  doc.mapping      [[…], …]             6장 대조표 (없으면 장 생략)
  doc.open_items   [[…], …]             7장 미반영 항목 (없으면 장 생략)
  doc.mapping_note / doc.open_note      각 장 끝 보충 문단
"""
import json
import os
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import erd
from build_erd import doc_counts, doc_tables, doc_text, meta_cells, require_fresh
from erd import AREAS, AREA_NAME, AREA_SCHEMA, LAYERS, ROLE, SCHEMA, layer

# 경로 값(OUT·PROJ)은 `config` 를 통해 **쓰는 자리에서** 묻는다 — 그 이름들은 PEP 562
# 모듈 __getattr__ 로 늦춰져 있어, `from config import OUT` 한 줄이 import 시점에
# mkdir 을 돌려 부르는 사람의 cwd 에 `erd-build/out` 을 남겼다.
import config
from config import DOCNAME
from erd import SPEC
from i18n import LANG, t as T

DOC = SPEC.get('doc', {})
TITLE = DOC.get('title', DOCNAME)

# 문서 폰트는 이름으로 지정한다 — 그림과 달리 여는 PC에 그 폰트가 있어야 한다.
# 그 폰트가 없는 PC에서는 Word 가 알아서 대체하므로 레이아웃만 조금 달라진다.
# 그래서 기본값은 그 말을 쓰는 곳에서 가장 흔한 폰트로 잡는다.
_DOC_FONTS = {
    'ko': ('Pretendard', 'D2Coding'),
    'ja': ('Yu Gothic', 'Consolas'),
    'en': ('Calibri', 'Consolas'),
    'es': ('Calibri', 'Consolas'),
}
_def_font, _def_mono = _DOC_FONTS.get(LANG, _DOC_FONTS['en'])
FONT = os.environ.get('ERD_DOC_FONT', _def_font)
MONO = os.environ.get('ERD_DOC_MONO', _def_mono)    # 없으면 Word 가 대체 — 코드 식별자용


# ── 서식 헬퍼 ────────────────────────────────────────────────────────────────
def set_font(run, name=FONT, size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.append(rf)
    rf.set(qn('w:eastAsia'), name)
    rf.set(qn('w:ascii'), name)
    rf.set(qn('w:hAnsi'), name)


def para(doc, text='', size=10, bold=False, name=FONT, align=None, space_after=6,
         color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    if align:
        p.alignment = align
    if text:
        set_font(p.add_run(text), name, size, bold, color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 15, 2: 12.5, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), FONT, sizes[level], True,
             '1F3864' if level == 1 else ('2E5C8A' if level == 2 else '404040'))
    return p


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), color)
    tcPr.append(sh)


def table(doc, headers, widths=None, style='Table Grid', font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(h), FONT, font_size, True)
        shade(cell, 'D9E2F3')
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    _mark_header_repeat(hdr)
    return t


def _mark_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def row(t, values, widths=None, font_size=8.5, mono_cols=(), bold_cols=(),
        colors=None, aligns=None):
    cells = t.add_row().cells
    # spec 은 사람이 손으로 쓴다. 칸이 남거나 모자란다고 문서 생성이 통째로 죽으면
    # 어디가 잘못됐는지도 알 수 없다 — 표 폭에 맞춰 자르고 채운다.
    values = (list(values) + [''] * len(cells))[:len(cells)]
    for i, v in enumerate(values):
        cell = cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        if aligns and aligns[i]:
            p.alignment = aligns[i]
        set_font(p.add_run(str(v)), MONO if i in mono_cols else FONT, font_size,
                 i in bold_cols, (colors or {}).get(i))
        if widths:
            cell.width = Cm(widths[i])
    return cells


MISSING_FIGS = []


def picture(doc, path, width_cm, caption, max_h_cm=16.2):
    """폭·높이 둘 다에 맞춘다. 폭만 지정하면 세로가 긴 그림이 페이지를 넘어 잘린다.

    가로 A4(29.7×21cm) · 여백 1.5cm → 가용 26.7×18cm. 제목·캡션 몫을 빼고 16.2cm 를 쓴다.

    그림 파일이 없으면 건너뛰고 이름을 적어 둔다. 신선도 관문(require_fresh)은 **없는
    파일을 세지 않으므로** 그대로 지나가고, 그 다음 줄의 PIL 이 없는 파일을 열어
    FileNotFoundError 를 그대로 사용자에게 뱉었다 — 무엇을 어떻게 하라는 말이 한 줄도
    없는 역추적이다. 같은 상황에서 build_html.py 는 도판을 빼고 경고 한 줄로 알린다.
    다만 그쪽은 영역 그림이 빠진 것만 세어서 개요도·전체도가 통째로 없어도 조용하다.
    여기서는 빠진 그림을 하나도 빼놓지 않고 말한다.
    """
    path = Path(path)
    if not path.exists():
        MISSING_FIGS.append(path.name)
        return
    iw, ih = Image.open(path).size
    w_cm = min(width_cm, max_h_cm * iw / ih)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Cm(w_cm))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(14)
    set_font(c.add_run(caption), FONT, 9, False, '595959')


def landscape(doc):
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    s.left_margin = s.right_margin = Cm(1.5)
    s.top_margin = s.bottom_margin = Cm(1.5)
    return s


def portrait(doc):
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = s.bottom_margin = Cm(2.0)
    return s


# ── 6장 대조표 — spec 의 doc.mapping 이 있으면 그것을 쓴다 ────────────────────
# 6장(설계안 대조) · 7장(미반영 항목) 데이터는 spec 의 doc.mapping / doc.open_items
# 에서 온다. 없으면 그 장을 통째로 생략한다.


MAPPING_TABLE = [list(r) for r in DOC.get('mapping', [])]
OPEN_ITEMS = [list(r) for r in DOC.get('open_items', [])]

# 문서가 절을 내주는 테이블. 본문(4장)은 AREAS 를 돌면서 표지·1.2·4장 머리말은
# len(SCHEMA) 를 세고 있었다 — 영역에 안 든 테이블이 하나라도 있으면 '11 테이블' 이라
# 적힌 문서에 절이 10개만 실렸고, 5장 관계표에는 본문에 절이 없는 테이블의 관계가
# 실렸다. 문서가 제 이야기를 할 때 세는 것과 싣는 것을 같은 자리에서 뽑는다.
# (그림 캡션의 수는 그림이 그린 것을 세므로 SCHEMA 그대로다.)
DOC_TABLES = doc_tables()
N_TABLES, N_COLUMNS, N_FKS = doc_counts(DOC_TABLES)


def build():
    # 문서에 박히는 그림은 PNG 뿐이다. 스키마보다 오래된 것이 하나라도 있으면
    # 4장 컬럼표와 2~3장 그림이 서로 다른 스키마를 말하게 된다 — 만들기 전에 멈춘다.
    require_fresh(['erd_overview', 'erd_full'] + [f'erd_area_{a[0]}' for a in AREAS],
                  ('.png',))
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = s.bottom_margin = Cm(2.0)
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # ── 표지 ──
    p = para(doc, DOC.get('title', DOCNAME), size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p.paragraph_format.space_before = Pt(40)
    para(doc, DOC.get('subtitle', ''),
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26, color='595959')

    t = table(doc, [T('word.kind'), T('word.content'), T('word.kind'), T('word.content')],
              [2.6, 5.6, 2.4, 5.4], font_size=9.5)
    meta = meta_cells(DOC.get('meta', [
        [T('docx.doc_name'), DOC.get('title', DOCNAME), T('word.tables'), str(N_TABLES)],
        [T('word.columns'), str(N_COLUMNS), T('word.fkeys'), str(N_FKS)],
    ]))
    for r in meta:
        cells = row(t, r, [2.6, 5.6, 2.4, 5.4], font_size=9.5, bold_cols=(0, 2))
        shade(cells[0], 'F2F2F2')
        shade(cells[2], 'F2F2F2')
    # 마지막 행의 오른쪽 쌍이 비어 있을 때만 왼쪽 내용이 그 자리까지 쓰게 편다.
    # 예전엔 `if True:` 라 **늘** 폈고, 그래서 spec 이 doc.meta 를 안 준 기본 표지에서
    # 마지막 행이 `Columns | 5 | FKs | 0` 인데 셋을 하나로 합쳐 한 칸이 `5 FKs 0` 이
    # 됐다 — 표지가 컬럼 수와 FK 수를 말한다면서 둘을 한 덩어리로 뭉갠 셈이다.
    if meta and not str(meta[-1][3]).strip():
        last = t.rows[-1]
        last.cells[1].merge(last.cells[3])

    doc.add_page_break()

    # ── 1. 개요 ──
    heading(doc, T('docx.ch1'))
    heading(doc, T('docx.ch1_1'), 2)
    para(doc, DOC.get('purpose', T('docx.purpose')))
    heading(doc, T('docx.ch1_2'), 2)
    for line in DOC.get('scope', [T('docx.scope_in', n=N_TABLES), T('docx.scope_out')]):
        para(doc, line)
    heading(doc, T('docx.ch1_3'), 2)
    para(doc, DOC.get('sources_note', T('docx.sources_note')))
    t = table(doc, [T('word.basis'), T('word.content')], [5.6, 11.4], font_size=9)
    for r in DOC.get('sources', [
            ('information_schema', T('docx.src_infoschema')),
            (T('docx.src_comment'), T('docx.src_comment_d')),
            (T('docx.src_orm'), T('docx.src_orm_d'))]):
        row(t, r, [5.6, 11.4], font_size=9, mono_cols=(0,))
    para(doc, '')

    heading(doc, T('docx.ch1_4'), 2)
    t = table(doc, [T('word.notation'), T('word.meaning')], [4.2, 12.8], font_size=9)
    for code, (_f, _h, _b, label) in LAYERS.items():
        row(t, (T('docx.by_color', code=code), label), [4.2, 12.8], font_size=9)
    for r in [('NEW', T('docx.nt_new')),
              (T('word.extended'), T('docx.nt_ext')),
              (T('word.source'), T('docx.nt_src')),
              (T('word.added'), T('docx.nt_added')),
              (T('word.solid'), T('docx.nt_solid')),
              (T('word.dashed'), T('docx.nt_dashed')),
              (T('word.semicircle'), T('docx.nt_hop')),
              (T('word.crowfoot'), T('docx.nt_card')),
              ('NN', 'NOT NULL'),
              ('UQ', T('word.unique'))]:
        row(t, r, [4.2, 12.8], font_size=9)
    para(doc, '')

    # ── 2~3. 그림 (가로 섹션) ──
    landscape(doc)
    heading(doc, T('docx.ch2'))
    para(doc, T('docx.ch2_intro'))
    picture(doc, config.OUT / 'erd_overview.png', 26.0,
            T('word.fig_no', n=1) + ' ' + T('docx.fig_overview', title=TITLE))

    heading(doc, T('docx.ch2_1'), 2)
    para(doc, T('docx.ch2_1_intro', n=len(SCHEMA)))
    picture(doc, config.OUT / 'erd_full.png', 26.0,
            T('word.fig_no', n=2) + ' ' + T('docx.fig_full', title=TITLE))

    heading(doc, T('docx.ch3'))
    para(doc, T('docx.ch3_intro'))
    for i, (code, name, schema, tables) in enumerate(AREAS, start=3):
        heading(doc, T('docx.ch3_area', no=i - 2, code=code, name=name,
                        schema=schema, n=len(tables)), 2)
        picture(doc, config.OUT / f'erd_area_{code}.png', 26.0,
                T('word.fig_no', n=i) + ' ' + T('docx.fig_area', code=code, name=name))

    # ── 4. 테이블별 역할 및 컬럼 설명 ──
    portrait(doc)
    heading(doc, T('docx.ch4'))
    para(doc, T('docx.ch4_intro', tables=N_TABLES, columns=N_COLUMNS))
    n = 0
    for ai, (code, name, schema, tables) in enumerate(AREAS, start=1):
        heading(doc, T('docx.ch4_area', no=ai, code=code, name=name), 2)
        for ti, tname in enumerate(tables, start=1):
            n += 1
            t_ = SCHEMA[tname]
            bd, _c = erd.badge(tname)
            # 역할명이 비면 ` · ` 만 남아 '4.1.1 users · ' 로 끝났다. HTML 은 같은
            # 값이 비면 역할 블록을 아예 안 그린다 — 같은 데이터를 두 문서가 다르게
            # 처리하던 자리다. 구분자는 이을 것이 있을 때만 쓴다.
            role = str(ROLE.get(tname, '') or '').strip()
            heading(doc, f'4.{ai}.{ti} {tname}' + (f' · {role}' if role else ''), 3)
            meta = table(doc, [T('word.schema'), T('word.kind'), T('word.layer'),
                               T('word.columns'), T('word.fkeys')],
                         [2.4, 2.2, 4.6, 1.8, 1.8], font_size=8.5)
            row(meta, (t_.get('schema', 'public'), bd, LAYERS[layer(tname)][3],
                       len(t_['columns']), len(t_['fks'])),
                [2.4, 2.2, 4.6, 1.8, 1.8], font_size=8.5,
                aligns=[None, WD_ALIGN_PARAGRAPH.CENTER, None,
                        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
            if t_.get('note'):
                para(doc, t_['note'], size=9, color='595959', space_after=3)
            ct = table(doc, [T('word.kind'), T('col.name'), T('col.type'), T('col.desc')],
                       [1.3, 4.2, 3.0, 8.5], font_size=8)
            for c in t_['columns']:
                # UNIQUE 는 docx 어디에도 없었다 — 컬럼표에도, 따로 실리는 제약 절도
                # 없다(HTML 은 컬럼표의 키 칸과 접힌 제약 목록 둘 다에 싣는다). 즉
                # 제출용 문서 쪽에서만 스키마 사실 하나가 통째로 사라지고 있었다.
                # 칸을 늘리지 않고 이미 있는 '구분' 칸에 적는다.
                mark = erd.col_role(t_, c) or (
                    'UQ' if any(c['name'] in u for u in t_.get('uniques', [])) else '')
                desc = (T('word.added') + ' ' if c['added'] else '') + c['comment']
                ct.rows  # noqa
                row(ct, (mark, c['name'], c['type'] + (' NN' if c['not_null'] else ''), desc),
                    [1.3, 4.2, 3.0, 8.5], font_size=8, mono_cols=(1, 2), bold_cols=(1,),
                    aligns=[WD_ALIGN_PARAGRAPH.CENTER, None, None, None],
                    colors={0: {'PK': 'B03A2E', 'UQ': '7A4FBF'}.get(mark, '1F618D')}
                    if mark else None)
            para(doc, '', space_after=8)

    # ── 5. 관계 정의 ──
    heading(doc, T('docx.ch5'))
    heading(doc, T('docx.ch5_1'), 2)
    # 본문에 절이 없는 테이블의 관계는 싣지 않는다 — 5장만 보면 있는 절을 찾게 되고,
    # 4장이 세는 것과 5장이 세는 것이 서로 다른 집합이 된다.
    fks = [(n_, fk) for n_ in DOC_TABLES for fk in SCHEMA[n_]['fks']]
    para(doc, T('docx.ch5_1_intro', n=len(fks)))
    t = table(doc, [T('word.child_table'), T('col.name'), T('word.parent_table'),
                    T('col.name'), T('word.delete_rule')],
              [4.4, 3.5, 4.2, 2.3, 2.6], font_size=8)
    for n_, fk in sorted(fks, key=lambda x: (x[0], x[1]['column'])):
        row(t, (n_, fk['column'], fk['ref_table'], fk['ref_column'],
                f"ON DELETE {fk['on_delete']}"),
            [4.4, 3.5, 4.2, 2.3, 2.6], font_size=8, mono_cols=(0, 1, 2, 3),
            colors={4: 'B03A2E' if fk['on_delete'] == 'CASCADE' else '595959'})
    para(doc, '')

    heading(doc, T('docx.ch5_2'), 2)
    para(doc, T('docx.ch5_2_intro'))
    t = table(doc, [T('word.src_side'), T('word.dst_side'), T('word.content')],
              [5.0, 5.4, 6.6], font_size=8.5)
    # 그림에 없는 흐름을 표에만 싣지 않는다 — 문서만 보면 있는 관계로 읽힌다
    for src, dst, label in erd.DERIVES:
        if src not in erd.SCHEMA or dst not in erd.SCHEMA:
            continue
        row(t, (src, dst, label), [5.0, 5.4, 6.6], font_size=8.5, mono_cols=(0, 1))
    para(doc, '')

    # ── 6. 대조표 (spec 에 doc.mapping 이 있을 때만) ──
    if MAPPING_TABLE:
        _chapter_mapping(doc)
    if OPEN_ITEMS:
        _chapter_open(doc)

    path = config.PROJ / f'{DOCNAME}.docx'
    doc.save(path)
    if MISSING_FIGS:
        # 조용히 빠지면 그림 없는 문서가 그림 있는 문서인 척 나간다
        print(T('log.figs_missing', n=len(MISSING_FIGS),
                list=', '.join(MISSING_FIGS[:6]) + (' …' if len(MISSING_FIGS) > 6 else '')))
    return path


def _chapter_mapping(doc):
    heading(doc, T('docx.ch6'))
    para(doc, doc_text(DOC, 'mapping_intro', T('docx.ch6_intro')))
    t = table(doc, ['No', T('word.proposed'), T('word.actual_table'),
                    T('word.applied'), T('word.reason')],
              [1.0, 3.6, 4.6, 1.8, 6.0], font_size=8)
    for r in MAPPING_TABLE:
        row(t, r, [1.0, 3.6, 4.6, 1.8, 6.0], font_size=8, mono_cols=(1, 2),
            aligns=[WD_ALIGN_PARAGRAPH.CENTER, None, None, WD_ALIGN_PARAGRAPH.CENTER, None])
    para(doc, '')
    if DOC.get('mapping_note'):
        para(doc, DOC['mapping_note'], size=9)


def _chapter_open(doc):
    heading(doc, T('docx.ch7'))
    para(doc, T('docx.ch7_intro'))
    t = table(doc, [T('word.priority'), T('word.item'), T('word.target'),
                    T('word.current'), T('word.action')],
              [1.0, 3.0, 3.4, 5.4, 4.2], font_size=8)
    for r in OPEN_ITEMS:
        row(t, r, [1.0, 3.0, 3.4, 5.4, 4.2], font_size=8, mono_cols=(2,),
            aligns=[WD_ALIGN_PARAGRAPH.CENTER, None, None, None, None])
    para(doc, '')
    if DOC.get('open_note'):
        para(doc, DOC['open_note'], size=9)


if __name__ == '__main__':
    p = build()
    print(T('log.docx_saved', name=p.name, kb=f'{p.stat().st_size / 1024:.0f}'))
