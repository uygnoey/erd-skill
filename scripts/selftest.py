#!/usr/bin/env python3
"""회귀 시험 — DB 없이 스킬 전체를 한 번 돌려 보고 결과를 검사한다.

    python3 selftest.py            전부 (옆의 selftest_*.py 까지 한 벌로)
    python3 selftest.py parse      이름에 'parse' 가 든 항목만

이 파일이 회귀 시험의 입구다. 항목은 여기와 `selftest_*.py` 들에 나뉘어 있고, 아래
`load_extras()` 가 옆에 있는 것을 전부 끌어와 한 목록으로 돌린다 — 도우미와 그 목록은
`selftest_kit.py` 에 있다. `install.sh --check` 는 이 파일만 부른다.

그림 품질은 렌더링할 때마다 자체 검증이 찍히지만, 그 밖의 기능은 재는 것이 없었다.
그래서 고칠 때마다 다른 데가 조용히 죽었다 — 인라인 주석은 두 판을 통째로 빈
문자열인 채 지나갔고(pg_dump 는 COMMENT ON 을 쓰므로 그쪽 시험만 통과했다), 자기참조
루프는 두 개가 겹쳐 하나로 보이는 동안 검증이 0 을 찍었다.

여기 담는 것은 **한 번이라도 조용히 깨졌던 것들**이다. 새로 고칠 때마다 그 자리를
지키는 항목을 하나 남긴다.

DB 도 docker 도 필요 없다. 임시 디렉토리에서 돌고 뒤를 치운다.
출력은 영어다 — 사용자가 아니라 이 코드를 고치는 사람이 보는 것이라서.
"""
import json
import os
import re
import subprocess
import sys

from selftest_kit import (Fail, HERE, MEASURES, case, col, ddl, drawn_names, eq, has,
                          hub_schema, load_extras, main, run, table, verify_clean,
                          verify_faults, verify_recs, write_schema)

# ── 말 ──────────────────────────────────────────────────────────────────────
@case('i18n: four catalogs agree')
def _(work):
    sys.path.insert(0, str(HERE))
    import lang.en as en
    for code in ('ko', 'ja', 'es'):
        m = __import__(f'lang.{code}', fromlist=['M']).M
        eq(set(m), set(en.M), f'lang/{code}.py key set')
        for k in en.M:
            eq(sorted(re.findall(r'{(\w+)}', m[k])),
               sorted(re.findall(r'{(\w+)}', en.M[k])), f'{code} placeholders in {k}')
            if not str(m[k]).strip():
                raise Fail(f'{code} {k} is empty')


@case('i18n: every key used by the code exists')
def _(work):
    sys.path.insert(0, str(HERE))
    import lang.en as en
    used = set()
    for f in HERE.glob('*.py'):
        if f.name.startswith('selftest'):
            continue                 # 시험 파일은 배포물이 아니다 — 여기 있는 T( 는 예시다
        for m in re.finditer(r"""T\(\s*['"]([a-z_]+\.[a-z_0-9]+)['"]""",
                             f.read_text(encoding='utf-8')):
            used.add(m.group(1))
    missing = sorted(used - set(en.M))
    eq(missing, [], 'keys used in code but absent from lang/en.py')


@case('i18n: {key} placeholder does not collide with t()')
def _(work):
    from i18n import t
    # t(key, /, **kw) 라야 {key} 를 자리표시자로 쓸 수 있다 — 위치 전용이 아니면
    # 여기서 TypeError 가 난다
    has(t('err.spec_layer', key='ZZZ', value=[]), 'ZZZ',
        'a message whose placeholder is named {key} must still format')


# ── 값 다듬기 ────────────────────────────────────────────────────────────────
@case('clean: newline and control chars collapse')
def _(work):
    from config import clean
    eq(clean('a\nb\tc\x0bd\x1fe   f'), 'a b c d e f', 'clean() flattens to one line')
    eq(clean(None), '', 'clean(None)')
    eq(clean('키 | 파이프'), '키 | 파이프', 'pipe survives')


@case('clean: a newline in a comment does not kill the diagram')
def _(work):
    write_schema(work, {'t': table('t', [col('id'), col('m', 'text',
                                                        comment='one\ntwo\x0bthree')])})
    run('build_erd.py', work)
    if not (work / 'out' / 'erd_full.png').exists():
        raise Fail('no PNG produced')


@case('config: docname is made safe for a filename')
def _(work):
    # 예전엔 `any('/' in n for n in names)` 를 봤다 — names 는 Path.name 이라 구분자가
    # **들어 있을 수 없다**. 즉 조건이 늘 거짓인, 아무것도 못 잡는 줄이었다. 슬래시가
    # 살아남으면 실제로 나는 일은 '이름에 /가 남는 것' 이 아니라 하위 디렉토리가 생겨
    # 산출물이 엉뚱한 자리로 가는 것이다. 그러니 나온 이름과 자리를 직접 못박는다.
    write_schema(work, {'t': table('t', [col('id')])})
    run('build_erd.py', work, env={'ERD_DOCNAME': 'a/b:c'})
    eq(sorted(p.name for p in work.glob('*.graphml')), ['a_b_c.graphml'],
       'every separator becomes _ and the file lands in the project directory itself')
    if (work / 'a').exists():
        raise Fail('a slash turned into a subdirectory — the artifact left its place')
    if list(work.parent.glob('*.graphml')):
        raise Fail(f'an artifact escaped the project directory: '
                   f'{[p.name for p in work.parent.glob("*.graphml")]}')
    # 위로 거슬러 올라가는 이름도 (.. 은 strip('. ') 로 사라진다)
    run('build_erd.py', work, env={'ERD_DOCNAME': '../../etc/passwd'})
    eq(sorted(p.name for p in work.glob('*.graphml')),
       ['_.._etc_passwd.graphml', 'a_b_c.graphml'], 'a traversal cannot climb out')
    # 빈 이름은 '.html' 같은 숨김 파일이 된다 — 기본 이름으로 떨어져야 한다
    run('build_erd.py', work, env={'ERD_DOCNAME': '  . '})
    if not (work / 'ERD.graphml').exists():
        raise Fail('an empty docname must fall back to a visible default name')


# ── DDL 파서 ─────────────────────────────────────────────────────────────────
@case('parse: inline -- comment attaches to its own column')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  id bigint PRIMARY KEY,   -- row id
  amount numeric(12,2),    -- amount
  plain text
);
""")
    got = {c['name']: c['comment'] for c in s['t']['columns']}
    eq(got, {'id': 'row id', 'amount': 'amount', 'plain': ''},
       'trailing comments belong to the column they follow, not the next one')


@case('parse: a comment on a shared line does not swallow the next column')
def _(work):
    s = ddl(work, """
CREATE TABLE t (a int, b int,  -- shared line
  c int);
CREATE TABLE leading (
    id int PRIMARY KEY   -- row id
  , name text            -- leading-comma style
  , v int
);
""")
    eq([c['name'] for c in s['t']['columns']], ['a', 'b', 'c'],
       'a trailing comment must not merge two columns')
    eq([c['name'] for c in s['leading']['columns']], ['id', 'name', 'v'],
       'leading-comma style keeps every column')
    eq({c['name']: c['comment'] for c in s['leading']['columns']},
       {'id': 'row id', 'name': 'leading-comma style', 'v': ''},
       'and each comment stays with its own column')


@case('parse: a constraint written after REFERENCES still counts')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  id int PRIMARY KEY,
  parent_id int REFERENCES users(id) NOT NULL
);
""")
    c = [x for x in s['t']['columns'] if x['name'] == 'parent_id'][0]
    eq(c['not_null'], True, 'NOT NULL after REFERENCES is still NOT NULL')


@case('parse: PRIMARY KEY implies NOT NULL')
def _(work):
    # Postgres puts NOT NULL on every PK column; introspect reports it that way.
    # The diagrams hide a miss (the PK icon wins) but the definition tables do not.
    s = ddl(work, """
CREATE TABLE t (id bigint PRIMARY KEY, v text);
CREATE TABLE pair (a int, b int, note text, PRIMARY KEY (a, b));
CREATE TABLE later (id bigint, v text);
ALTER TABLE later ADD CONSTRAINT later_pk PRIMARY KEY (id);
""")
    eq({c['name']: c['not_null'] for c in s['t']['columns']},
       {'id': True, 'v': False}, 'an inline PK column is NOT NULL')
    eq({c['name']: c['not_null'] for c in s['pair']['columns']},
       {'a': True, 'b': True, 'note': False}, 'a composite table-level PK too')
    eq({c['name']: c['not_null'] for c in s['later']['columns']},
       {'id': True, 'v': False}, 'and a PK added by ALTER TABLE')


@case('parse: columns named after table-level keywords survive')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  id bigint PRIMARY KEY, checksum text, like_count int, unique_code varchar(20),
  constraint_type text, partition_id int, exclude_flag boolean
);
""")
    eq([c['name'] for c in s['t']['columns']],
       ['id', 'checksum', 'like_count', 'unique_code', 'constraint_type',
        'partition_id', 'exclude_flag'], 'keyword-prefixed column names')


@case('parse: string literals cannot break the split')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  id bigint PRIMARY KEY,
  a text DEFAULT '(',
  b text DEFAULT '--none--',
  c text DEFAULT 'it''s fine',
  d text DEFAULT 'see REFERENCES users',
  e text
);
""")
    eq([c['name'] for c in s['t']['columns']], ['id', 'a', 'b', 'c', 'd', 'e'],
       'no column lost to a string literal')
    eq(s['t']['fks'], [], 'REFERENCES inside a string is not a foreign key')
    eq(s['t']['pk'], ['id'], 'PRIMARY KEY inside a string is not a key')
    eq([c['default'] for c in s['t']['columns'] if c['name'] == 'd'],
       ["'see REFERENCES users'"], 'default value kept whole')


@case('parse: dollar-quoted function bodies are skipped')
def _(work):
    s = ddl(work, """
CREATE FUNCTION f() RETURNS void AS $_$
BEGIN
  CREATE TABLE ghost (x int);
  RAISE NOTICE 'it''s $$ tricky';
END
$_$ LANGUAGE plpgsql;
CREATE TABLE real_one (id bigint PRIMARY KEY);
""")
    eq(sorted(s), ['real_one'], 'no ghost table from a function body')


@case('parse: nested block comments and E-strings')
def _(work):
    s = ddl(work, r"""
CREATE TABLE t (
  id int PRIMARY KEY,
  /* outer /* inner */ still a comment */
  x int,
  y text DEFAULT E'it\'s',
  z int
);
""")
    eq([c['name'] for c in s['t']['columns']], ['id', 'x', 'y', 'z'],
       'nested comment / E-string do not swallow columns')


@case('parse: one-line definition keeps every column')
def _(work):
    s = ddl(work, "CREATE TABLE t (a int PRIMARY KEY, b text, c varchar(10));\n")
    eq([c['name'] for c in s['t']['columns']], ['a', 'b', 'c'], 'one-line CREATE TABLE')


@case('parse: pg_dump shapes — schema-qualified, ONLY, ADD CONSTRAINT')
def _(work):
    s = ddl(work, """
CREATE TABLE IF NOT EXISTS public.users (id bigint NOT NULL, email character varying(255));
CREATE TABLE public.orders (id bigint NOT NULL, user_id bigint);
ALTER TABLE ONLY public.users ADD CONSTRAINT users_pkey PRIMARY KEY (id);
ALTER TABLE ONLY public.orders ADD CONSTRAINT orders_user_fk FOREIGN KEY (user_id)
    REFERENCES public.users(id) ON DELETE CASCADE ON UPDATE CASCADE;
COMMENT ON TABLE public.users IS 'members';
COMMENT ON COLUMN public.users.email IS 'login id';
""")
    eq(s['users']['pk'], ['id'], 'PK from ALTER TABLE ONLY')
    eq([(f['column'], f['ref_table'], f['on_delete']) for f in s['orders']['fks']],
       [('user_id', 'users', 'CASCADE')], 'FK and delete rule (ON UPDATE must not leak in)')
    eq(s['users']['note'], 'members', 'COMMENT ON TABLE')
    eq([c['comment'] for c in s['users']['columns'] if c['name'] == 'email'],
       ['login id'], 'COMMENT ON COLUMN')
    eq([c['type'] for c in s['users']['columns'] if c['name'] == 'email'],
       ['varchar(255)'], 'character varying(N) normalizes like introspect')


@case('parse: two-part COMMENT ON COLUMN (no schema)')
def _(work):
    s = ddl(work, """
CREATE TABLE t (id bigint PRIMARY KEY, email text);
COMMENT ON COLUMN t.email IS 'hand written';
""")
    eq([c['comment'] for c in s['t']['columns'] if c['name'] == 'email'],
       ['hand written'], 'two-part column comment')


@case('parse: same name in two schemas stays two tables')
def _(work):
    s = ddl(work, """
CREATE TABLE shop.orders (id bigint PRIMARY KEY, amount numeric(10,2));
CREATE TABLE mart.orders (id bigint PRIMARY KEY, loaded_at timestamptz);
CREATE UNIQUE INDEX uq ON shop.orders (amount);
""")
    eq(sorted(s), ['mart.orders', 'shop.orders'], 'schema-qualified keys')
    eq(s['shop.orders']['uniques'], [['amount']], 'unique index attaches by schema')
    eq(s['mart.orders']['uniques'], [], 'and not to the other schema')


@case('parse: composite FK pairs columns positionally')
def _(work):
    s = ddl(work, """
CREATE TABLE order_items (order_id bigint, line_no int, PRIMARY KEY (order_id, line_no));
CREATE TABLE notes (
  id bigint PRIMARY KEY, o bigint, l int,
  FOREIGN KEY (o, l) REFERENCES order_items(order_id, line_no) ON DELETE CASCADE
);
""")
    eq([(f['column'], f['ref_column']) for f in s['notes']['fks']],
       [('o', 'order_id'), ('l', 'line_no')], 'child column ↔ parent column by position')


@case('parse: REFERENCES serial_numbers does not make a column identity')
def _(work):
    s = ddl(work, """
CREATE TABLE t (id bigint PRIMARY KEY, ref bigint REFERENCES serial_numbers(id));
""")
    c = [x for x in s['t']['columns'] if x['name'] == 'ref'][0]
    eq((c['identity'], c['not_null']), (False, False), 'substring SERIAL must not match')


@case('parse: DDL only, no database')
def _(work):
    s = ddl(work, """
CREATE TABLE orders (id bigint PRIMARY KEY, m bigint REFERENCES merchants(id));
""")
    has(sorted(s), 'merchants', 'referenced-but-undefined table is kept')
    eq(s['merchants']['columns'], [], 'with no columns')
    run('build_erd.py', work)          # 예전엔 여기서 max() 가 죽었다
    if not (work / 'out' / 'erd_full.png').exists():
        raise Fail('a column-less table must render as a title-only box')


# ── 인트로스펙션 ─────────────────────────────────────────────────────────────
# 진짜 DB 없이 introspect 의 **판단**을 재현하는 가짜 psql. 진짜 서버가 하는 것만 한다:
#   · 조회가 JSON 으로 싸여 오면 JSON 한 줄씩, 아니면 받은 -F·-R 로 이어 붙여 내보낸다
#     (전송이 구분자 방식으로 되돌아가면 값 속 \x1e 가 행을 쪼개는 것까지 재현된다)
#   · FAKE_PG_VER 이 11 미만이면 conparentid 를 묻는 조회를 거절한다 (PG 10 이하)
#   · FAKE_FAIL 에 적힌 조각이 든 조회는 **몇 행 흘린 뒤** 죽는다 (문 타임아웃·재기동)
_FAKE_PSQL = '''\
import json
import os
import re
import sys
a = sys.argv
sep = a[a.index('-F') + 1]
rs = a[a.index('-R') + 1] if '-R' in a else chr(10)
q = a[a.index('-c') + 1]
NL, RS, FS = chr(10), chr(30), chr(31)
ver = int(os.environ.get('FAKE_PG_VER', '160000'))
rows = []
if 'server_version' in q:
    rows = [['%d.%d' % (ver // 10000, ver % 100), str(ver)]]
elif 'information_schema.columns' in q:
    rows = [['s1', 'claims', 'id', 'bigint', 'NO', '', 'NO', ''],
            ['s1', 'claims', 'owner_id', 'bigint', 'YES', '', 'NO', ''],
            ['s2', 'owners', 'id', 'bigint', 'NO', '', 'NO', ''],
            ['s1', 'tricky', 'id', 'bigint', 'NO', '', 'NO', ''],
            ['s1', 'tricky', 'note', 'text', 'YES',
             "'line1" + NL + "line2'::text", 'NO', ''],
            ['s1', 'tricky', 'ctrl', 'text', 'YES',
             "'has" + RS + "record" + RS + "sep'::text", 'NO',
             'comment with ' + RS + ' RS and ' + FS + ' FS inside']]
elif 'PRIMARY KEY' in q:
    rows = [['s1', 'claims', 'id'], ['s2', 'owners', 'id'], ['s1', 'tricky', 'id']]
elif "contype='f'" in q:
    if 'conparentid' in q and ver < 110000:
        sys.stderr.write('ERROR:  column con.conparentid does not exist' + NL)
        sys.exit(1)
    rows = [['s1', 'claims', 'owner_id', 'hidden', 'owners', 'id', 'NO ACTION'],
            ['s1', 'tricky', 'id', 's2', 'owners', 'id', 'CASCADE']]
if rows:
    m = re.search(r'_r\\(([^()]*)\\)\\s*$', q)
    if m:                          # 구조화 출력 — 값이 무엇을 담아도 한 행이 한 줄이다
        names = [c.strip() for c in m.group(1).split(',')]
        out = NL.join(json.dumps(dict(zip(names, r))) for r in rows) + NL
    else:
        out = rs.join(sep.join(r) for r in rows) + NL
    sys.stdout.write(out)
fail = os.environ.get('FAKE_FAIL', '')
if fail and fail in q:
    sys.stdout.flush()             # 흘릴 만큼 흘리고 죽는다 — 부분 출력 + 0 아닌 종료
    sys.stderr.write('ERROR:  canceling statement due to statement timeout' + NL)
    sys.exit(1)
'''


def db_fake(work, **env):
    import shlex
    work.mkdir(parents=True, exist_ok=True)
    fake = work / 'fake_psql.py'
    fake.write_text(_FAKE_PSQL, encoding='utf-8')
    e = {'ERD_PSQL': f'{shlex.quote(sys.executable)} {shlex.quote(str(fake))}',
         'ERD_SCHEMAS': 's1,s2'}
    e.update(env)
    return e


@case('introspect: an FK parent you cannot see is dropped, not rewired')
def _(work):
    # hidden.owners 는 목록 밖 — 같은 이름의 s2.owners 로 갈아타면 없는 관계가 그려진다
    r = run('introspect.py', work, env=db_fake(work))
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq(s['claims']['fks'], [], 'an FK to hidden.owners must not become s2.owners')
    has(r.stdout, 'outside the target: 1', 'the dropped FK is counted, not silent')


@case('introspect: a newline inside a default does not forge a row')
def _(work):
    run('introspect.py', work, env=db_fake(work))
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq(sorted(s), ['claims', 'owners', 'tricky'], 'no ghost table from a split row')
    eq([c['default'] for c in s['tricky']['columns'] if c['name'] == 'note'],
       ["'line1\nline2'::text"], 'the default keeps both of its lines')


@case('introspect: partition copies of an FK are filtered in the query')
def _(work):
    # 파티션마다 복제된 제약은 SQL 에서 걸러진다 — 여기서는 그 술어가 지워지지
    # 않았는지만 지킨다 (충실한 재현은 진짜 DB 가 필요하다)
    sys.path.insert(0, str(HERE))
    import introspect
    has(introspect.Q_FK, 'conparentid=0',
        'inherited (per-partition) constraint copies must be filtered out')


# ── 손으로 쓴 DDL 의 주석 소유권 ─────────────────────────────────────────────
# 규칙은 하나다: **주석은 아래로 붙는다**. 같은 줄 앞에 코드가 있으면 그 코드 것이고,
# 제 줄을 쓰는 주석은 바로 아래 코드 것이며, 아래에 코드가 없으면 임자가 없다.
# 콤마로 재던 두 판이 번갈아 한쪽을 깨뜨렸으므로 두 방식을 같은 자리에서 함께 지킨다.
@case('parse: a comment above a column describes that column')
def _(work):
    s = ddl(work, """
CREATE TABLE accounts (
  -- account identifier
  id bigint PRIMARY KEY,
  -- login email address
  email varchar(255) NOT NULL,
  -- account balance in cents
  balance bigint NOT NULL
);
""")
    eq({c['name']: c['comment'] for c in s['accounts']['columns']},
       {'id': 'account identifier', 'email': 'login email address',
        'balance': 'account balance in cents'},
       'comment-above-column must not shift every description down by one')


@case('parse: both comment styles in one table')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  id bigint PRIMARY KEY,   -- row id
  -- how much money
  -- in cents
  amount bigint,
  name text  -- display name
  , tenant_id bigint
  -- unique per tenant
  , UNIQUE (tenant_id, name)
  -- nothing below owns this
);
""")
    eq({c['name']: c['comment'] for c in s['t']['columns']},
       {'id': 'row id', 'amount': 'how much money in cents',
        'name': 'display name', 'tenant_id': ''},
       'trailing, stacked-above and leading-comma comments each keep their own column')
    eq(s['t']['uniques'], [['tenant_id', 'name']], 'the constraint itself still parses')


@case('parse: a comment above a constraint is not a column description')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  tenant_id bigint,   -- owning tenant
  code text,
  -- unique per tenant
  UNIQUE (tenant_id, code)
);
""")
    eq({c['name']: c['comment'] for c in s['t']['columns']},
       {'tenant_id': 'owning tenant', 'code': ''},
       'a note written above a table-level constraint belongs to no column')


@case('parse: a comment after the last column owns nothing')
def _(work):
    s = ddl(work, """
CREATE TABLE t (
  id bigint PRIMARY KEY,
  balance bigint NOT NULL   -- in cents
  -- TODO: add a currency column
);
CREATE TABLE u (
  id bigint PRIMARY KEY,
  -- first and last
  name text,
  -- TODO: add an email column
);
""")
    eq({c['name']: c['comment'] for c in s['t']['columns']},
       {'id': '', 'balance': 'in cents'},
       'a note before the closing paren must not overwrite the last description')
    eq({c['name']: c['comment'] for c in s['u']['columns']},
       {'id': '', 'name': 'first and last'},
       'and the same after a trailing comma')


# 구분자로 풀던 자리들 — 값에 개행이 들어오면 그대로 ValueError 였다.
# 컬럼 이름에도 개행이 들어갈 수 있다: create table t ("a<개행>b" int) 는 합법이다.
#
# 17R. parse_ddl 이 config.psql()(구분자) → config.psql_rows()(행마다 JSON 한 줄)로
# 옮겨서 이 가짜도 같이 옮긴다. selftest_schema.py 의 _FAKE_PSQL_JSON 과 같은
# 방식이다: psql_rows() 가 씌우는 `_r(c0, c1, …)` 별칭을 읽어 그 이름으로 JSON 한
# 줄씩 낸다. **재는 것은 그대로다** — 값에 개행이 들어 있어도 파서가 안 죽는가.
# JSON 이 개행을 \\n 으로 적으므로 한 행은 한 줄이고, 파서가 그것을 풀어 다시
# 개행이 든 컬럼 이름으로 만들어야 아래 eq() 가 통과한다.
_FAKE_PSQL_DDL = '''\
import json
import re
import sys
a = sys.argv
q = a[a.index('-c') + 1]
NL = chr(10)
if "'src'" in q:
    # fetch_ref — 4필드 (table_name, column_name, type, is_nullable)
    rows = [['lookup', 'co' + NL + 'de', 'text', 'NO']]
else:
    # fetch_existing — 5필드. 17R 에 맨 앞에 table_schema 가 붙었다: 조회가
    # ERD_SCHEMAS 를 보게 되면서 어느 스키마에서 온 행인지 스스로 말해야 한다.
    rows = [['public', 'merchants', 'id', 'bigint', 'NO'],
            ['public', 'merchants', 'me' + NL + 'mo', 'text', 'YES']]
m = re.search(r'_r\\(([^()]*)\\)\\s*$', q)
names = [c.strip() for c in m.group(1).split(',')] if m else []
sys.stdout.write(''.join(json.dumps(dict(zip(names, r))) + NL for r in rows))
'''


@case('parse: a newline inside a fetched value does not kill the parser')
def _(work):
    import shlex
    work.mkdir(parents=True, exist_ok=True)
    fake = work / 'fake_psql.py'
    fake.write_text(_FAKE_PSQL_DDL, encoding='utf-8')
    env = {'ERD_PSQL': f'{shlex.quote(sys.executable)} {shlex.quote(str(fake))}',
           'ERD_REF_SCHEMA': 'src', 'ERD_REF_TABLES': 'lookup'}
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / 'a.sql').write_text(
        'CREATE TABLE orders (id bigint PRIMARY KEY, m bigint REFERENCES merchants(id));\n',
        encoding='utf-8')
    run('parse_ddl.py', work, env=env, sql_dir=d)      # 예전엔 여기서 ValueError 였다
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq([c['name'] for c in s['merchants']['columns']], ['id', 'me\nmo'],
       'fetch_existing keeps a column name that contains a newline')
    eq([c['name'] for c in s['lookup']['columns']], ['co\nde'],
       'and so does fetch_ref')


# 17R. 위 가짜가 '값이 험한 서버' 라면, 아래 둘은 **죽는 서버** 와 **public 이 아닌
# 스키마** 다. 둘 다 흉내 내는 것은 프로토콜(행마다 row_to_json 한 줄)과 where 절뿐이다
# — 조회문의 낱말을 베껴 적으면 그 문장을 다듬을 때마다 거짓 빨강이 난다.
_FAKE_PSQL_DIES = '''\
import json
import sys
a = sys.argv
q = a[a.index('-c') + 1]
NL = chr(10)
# 한 행을 흘리고 죽는다 — 문 타임아웃·서버 재기동이 실제로 남기는 모양이다.
sys.stdout.write(json.dumps({'c0': 'public', 'c1': 'merchants', 'c2': 'id',
                             'c3': 'bigint', 'c4': 'NO'}) + NL)
sys.stderr.write('server closed the connection unexpectedly' + NL)
sys.exit(1)
'''

_FAKE_PSQL_SCHEMAS = '''\
import json
import re
import sys
a = sys.argv
q = a[a.index('-c') + 1]
NL = chr(10)
# 진짜 서버가 하는 것만 한다: where 절이 물은 스키마에 있는 행만 돌려준다.
m = re.search(r"c[.]table_schema in [(]([^)]*)[)]", q)
asked = re.findall(r"'([^']*)'", m.group(1)) if m else []
rows = [r for r in [['shop', 'merchants', 'id', 'bigint', 'NO']] if r[0] in asked]
cols = re.search(r'_r[(]([^()]*)[)]\\s*$', q)
names = [c.strip() for c in cols.group(1).split(',')] if cols else []
sys.stdout.write(''.join(json.dumps(dict(zip(names, r))) + NL for r in rows))
'''


def _ddl_with_db(work, fake_src, sql, env=None):
    """가짜 psql 하나를 걸고 parse_ddl 을 돌린다 → 그 판의 결과."""
    import shlex
    work.mkdir(parents=True, exist_ok=True)
    fake = work / 'fake_psql.py'
    fake.write_text(fake_src, encoding='utf-8')
    e = {'ERD_PSQL': f'{shlex.quote(sys.executable)} {shlex.quote(str(fake))}'}
    e.update(env or {})
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / 'a.sql').write_text(sql, encoding='utf-8')
    return e, d


@case('parse: a column query that dies halfway is refused, not written into schema.json')
def _(work):
    # `introspect: a half-read database is refused, not documented` 와 같은 규율을
    # DDL 경로에도 건다. 여기가 `psql()`(returncode 를 버린다) 로 읽던 동안, 몇 행
    # 흘리고 죽은 조회는 **컬럼 몇 개가 통째로 빠진 정의서**를 exit 0 으로 냈다.
    e, d = _ddl_with_db(
        work, _FAKE_PSQL_DIES,
        'CREATE TABLE orders (id bigint PRIMARY KEY, m bigint REFERENCES merchants(id));\n')
    r = run('parse_ddl.py', work, env=e, sql_dir=d, expect_ok=False)
    if r.returncode == 0:
        raise Fail(f'a column query that died partway must not end in exit 0\n{r.stdout}')
    if (work / 'schema.json').exists():
        raise Fail('a half-read column query left a schema.json behind — the next '
                   'step cannot tell it from a complete one:\n'
                   + (work / 'schema.json').read_text(encoding='utf-8')[:400])


@case('parse: an existing table outside public still gets its real columns')
def _(work):
    # ERD_SCHEMAS 를 무시하고 public 만 묻던 동안, 다른 스키마의 기존 테이블은
    # **영영** 컬럼을 못 채우면서 화면에는 한 글자도 안 나왔다 — 이름만 있는 빈
    # 상자가 '컬럼이 없는 테이블' 인 양 정의서에 실렸다.
    e, d = _ddl_with_db(
        work, _FAKE_PSQL_SCHEMAS,
        'CREATE TABLE orders (id bigint PRIMARY KEY, m bigint REFERENCES merchants(id));\n',
        env={'ERD_SCHEMAS': 'shop'})
    run('parse_ddl.py', work, env=e, sql_dir=d)
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq([c['name'] for c in s['merchants']['columns']], ['id'],
       'a table that lives in a schema the user named must not stay an empty box')


_FAKE_PSQL_UTF8 = '''\
import json
import re
import sys
a = sys.argv
q = a[a.index('-c') + 1]
NL = chr(10)
# 서버는 제 로케일과 무관하게 PGCLIENTENCODING(=UTF8) 으로 내보낸다 — 진짜 psql 이
# 하는 것이 그것이다. 그러니 바이트로 쓴다.
rows = [['public', 'merchants', 'id', 'bigint', 'NO'],
        ['public', 'merchants', chr(47700) + chr(47784), 'text', 'YES']]
m = re.search(r'_r[(]([^()]*)[)]\\s*$', q)
names = [c.strip() for c in m.group(1).split(',')] if m else []
out = ''.join(json.dumps(dict(zip(names, r)), ensure_ascii=False) + NL for r in rows)
sys.stdout.buffer.write(out.encode('utf-8'))
'''

ASCII_LOCALE = {'LC_ALL': 'C', 'LANG': 'C', 'PYTHONUTF8': '0',
                'PYTHONCOERCECLOCALE': '0'}


@case('parse: a value only utf-8 can carry comes back whole in an ascii locale')
def _(work):
    # 17R 뮤테이션. psql 파이프를 `text=True` 만으로 읽으면 파이썬은 **로케일 인코딩**
    # 으로 푼다 — ascii 로케일(LC_ALL=C)에서는 한글 코멘트가 든 DB 가 그 자리에서
    # `UnicodeDecodeError` 로 죽어, utf-8 로 못 박아 둔 schema.json 쓰기에 닿지도
    # 못했다. 보내는 값(PGCLIENTENCODING)과 읽는 값은 같은 것이어야 한다.
    e, d = _ddl_with_db(
        work, _FAKE_PSQL_UTF8,
        'CREATE TABLE orders (id bigint PRIMARY KEY, m bigint REFERENCES merchants(id));\n')
    e.update(ASCII_LOCALE)
    r = run('parse_ddl.py', work, env=e, sql_dir=d, expect_ok=False)
    if r.returncode != 0:
        raise Fail('the locale of the shell decided whether the database could be '
                   f'read at all:\n{(r.stdout + r.stderr)[-500:]}')
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq([c['name'] for c in s['merchants']['columns']], ['id', '\uba54\ubaa8'],
       'the value comes back the way the server sent it, whatever the shell speaks')


@case('errors: a run in an ascii locale is not killed by its own output')
def _(work):
    # 17R 뮤테이션. 카탈로그의 문구 자체가 `→`·`·` 를 쓴다. ascii 콘솔에서 그 한
    # 글자에 `UnicodeEncodeError` 가 나면 **일이 다 끝난 뒤 찍는 줄**에서 죽는다 —
    # 파일은 나갔는데 종료코드는 1 이라, 사용자는 무엇이 실패한 것인지 알 수 없다.
    # 재는 대상은 i18n 이 import 시점에 무르게 해 두는 콘솔이다.
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / 'a.sql').write_text('CREATE TABLE t (id bigint PRIMARY KEY);\n',
                             encoding='utf-8')
    r = run('parse_ddl.py', work, env=dict(ASCII_LOCALE), sql_dir=d, expect_ok=False)
    if r.returncode != 0:
        raise Fail('a run in an ascii locale died printing its own summary:\n'
                   + (r.stdout + r.stderr)[-500:])
    if not (work / 'schema.json').exists():
        raise Fail('nothing was written — this case would measure nothing')


@case('parse: ADD COLUMN carries its inline UNIQUE the same way CREATE TABLE does')
def _(work):
    s = ddl(work, 'CREATE TABLE orders (id bigint PRIMARY KEY);\n'
                  'ALTER TABLE orders ADD COLUMN code text UNIQUE;\n')
    eq(s['orders']['uniques'], [['code']],
       'a column added later keeps the UNIQUE written beside it')


@case('parse: ADD COLUMN carries its inline REFERENCES the same way CREATE TABLE does')
def _(work):
    s = ddl(work, 'CREATE TABLE users (id bigint PRIMARY KEY);\n'
                  'CREATE TABLE orders (id bigint PRIMARY KEY);\n'
                  'ALTER TABLE orders ADD COLUMN owner_id bigint REFERENCES users(id);\n')
    eq([(f['column'], f['ref_table']) for f in s['orders']['fks']],
       [('owner_id', 'users')],
       'a relationship written on a later-added column is still a relationship')


@case('parse: ALTER TABLE IF EXISTS still adds its column')
def _(work):
    # 손으로 쓰는 마이그레이션에 흔한 형태다. 이 조각을 넘지 못하던 동안 문 전체가
    # **아무 말 없이** 버려졌다 — 그 컬럼은 정의서에서 그냥 없었다.
    s = ddl(work, 'CREATE TABLE orders (id bigint PRIMARY KEY);\n'
                  'ALTER TABLE IF EXISTS orders ADD COLUMN memo text;\n')
    eq([c['name'] for c in s['orders']['columns']], ['id', 'memo'],
       'ALTER TABLE IF EXISTS is the same statement as ALTER TABLE')


@case('parse: ADD COLUMN IF NOT EXISTS names the column, not "IF"')
def _(work):
    # 세 낱말을 안 넘던 때는 그 첫 마디가 컬럼 이름이 됐다 — 이름 `IF`, 타입 `NOT` 인
    # 가짜 컬럼이 실리고 진짜 컬럼은 사라졌다. 경고는 없었다.
    s = ddl(work, 'CREATE TABLE orders (id bigint PRIMARY KEY);\n'
                  'ALTER TABLE orders ADD COLUMN IF NOT EXISTS memo text;\n')
    eq([(c['name'], c['type']) for c in s['orders']['columns']],
       [('id', 'bigint'), ('memo', 'text')],
       'IF NOT EXISTS is not a column name')


@case('parse: an array default keeps its commas and the constraint written after it')
def _(work):
    # 대괄호를 안 세던 때는 배열 리터럴 안의 콤마가 최상위로 보여 항목이 두 토막이
    # 났다 — 기본값은 `ARRAY['a'` 로 잘리고 `NOT NULL` 은 뒷토막에 실려 사라졌다.
    s = ddl(work, "CREATE TABLE t (\n"
                  "  id bigint PRIMARY KEY,\n"
                  "  tags text[] DEFAULT ARRAY['a','b'] NOT NULL\n"
                  ");\n")
    got = [(c['name'], c['default'], c['not_null']) for c in s['t']['columns']]
    eq(got, [('id', '', True), ('tags', "ARRAY['a','b']", True)],
       'an array literal is one value, and what follows it is still a constraint')


@case('parse: a comma inside a quoted column name does not split the definition')
def _(work):
    # 자를 자리를 이름까지 덮은 사본에서 재지 않으면 이름 **안** 의 콤마가 구분자로
    # 세어져 정의가 이름 한가운데서 쪼개진다.
    s = ddl(work, 'CREATE TABLE t ("a,b" int, x text);\n')
    eq([(c['name'], c['type']) for c in s['t']['columns']],
       [('a,b', 'int'), ('x', 'text')],
       'a quoted name is one name however many commas it holds')


@case('parse: a quoted default that spells out NOT NULL keeps the whole value')
def _(work):
    # 경계를 원본에서 재면 `DEFAULT 'x, NOT NULL y'` 가 `'x,` 로 잘리고, 값을 가린
    # 사본에서 꺼내면 통째로 빈 값이 된다. 값은 원본에서, 경계는 가린 사본에서다.
    s = ddl(work, "CREATE TABLE t (id bigint, note text DEFAULT 'x, NOT NULL y');\n")
    got = [(c['name'], c['default'], c['not_null']) for c in s['t']['columns']]
    eq(got, [('id', '', False), ('note', "'x, NOT NULL y'", False)],
       'a constraint spelled inside a string literal is part of the value')


@case('parse: an opening paren inside a quoted name does not swallow the CREATE TABLE')
def _(work):
    # `parse_create` 가 본문 괄호 짝을 **이름까지 덮은 사본** 에서 세지 않으면
    # `"a(b"` 의 그 괄호가 함께 세어져 짝이 영영 안 닫힌다 — `if depth: continue`
    # 로 그 문이 통째로 빠지고, 테이블 0개에 경고 한 줄 없이 끝났다. 위 콤마
    # 케이스와 다른 자리다(그쪽은 `split_top_level` 의 자를 자리를 잰다).
    s = ddl(work, 'CREATE TABLE t (id int, c text COLLATE "a(b" NOT NULL);\n')
    eq(sorted(s), ['t'], 'the statement still yields its table')
    eq([c['name'] for c in s['t']['columns']], ['id', 'c'],
       'a paren inside a quoted name is not a body paren')


@case('parse: a quote inside a name is not a literal, and a literal\'s " is not a name')
def _(work):
    # 앞의 반: `mask` 가 큰따옴표 식별자를 건너뛰지 않던 때는 이름 안의 `'` 하나가
    # 리터럴의 시작으로 읽혀 그 뒤가 전부 어긋났고, `CREATE TABLE` 이 통째로 사라졌다.
    s = ddl(work, "CREATE TABLE t (id int, c text COLLATE \"it's ok\", z int);\n")
    eq(sorted(s), ['t'], 'the statement still yields its table')
    eq([c['name'] for c in s['t']['columns']], ['id', 'c', 'z'],
       "an apostrophe inside a quoted name is part of the name, not a literal")

    # 뒤의 반이 이 케이스가 **혼자** 재는 자리다. 앞의 반만 두면 아무것도 안 잰다:
    # 큰따옴표를 건너뛰는 가지를 어떻게 부수든 아래 `--` 케이스가 **반드시 함께**
    # 붉어지고, 그쪽은 주석 임자까지 보므로 단독 킬러가 따로 있다. 그래서 앞의 반은
    # 글자 하나(`'`)를 못박는 값만 있고 검출력은 0 이다 — 잰 것: 그 가지를 죽이는
    # 뮤턴트에서 두 케이스가 같이 붉어졌다.
    #
    # 그 값을 채우는 것이 여기다. `mask` 는 큰따옴표를 알아본 대가로 **되돌림 판**을
    # 하나 들고 있다: 건너뛰고 나니 끝나지 않는 리터럴이 남으면(`ran_off`) `"` 를
    # 모르던 방식으로 한 번 더 읽어 그쪽이 성하면 그것을 쓴다. 아래 입력의 `"` 는
    # 식별자의 시작이 아니라 **리터럴 안**의 글자라, 짝으로 집힌 `"` 를 건너뛰면
    # 남은 `'` 가 파일 끝까지 리터럴이 된다 — 되돌림 판이 없으면 이 파일의 테이블이
    # **둘 다** 사라진다(경고 한 줄 없이). 손으로 쓰다 만 DDL 에서 뒤따르는 문을 다
    # 잃는 것이 이 저장소가 가장 싫어하는 모양인데, 그 되돌림 판을 통째로 들어내도
    # 나머지 케이스는 전부 초록이었다.
    s = ddl(work, 'CREATE TABLE q (a "b text DEFAULT \'has " quote\', z int);\n'
                  'CREATE TABLE u (id int);\n',
            sql_dir=work / 'sql2')
    eq(sorted(s), ['q', 'u'],
       'a " that turns out to be inside a literal costs no statement at all')


@case('parse: a -- inside a quoted name is not a comment but a real one still is')
def _(work):
    # 같은 부류의 셋째다. 이름 안의 `--` 를 줄 주석으로 읽으면 그 줄의 나머지와
    # 뒤따르는 컬럼이 함께 가려져 사라졌다. 그러면서 **진짜** 주석은 계속 주석이어야
    # 하므로, 이름이 사는 것과 설명이 제 임자에게 가는 것을 한자리에서 못박는다.
    s = ddl(work, 'CREATE TABLE t (id int, "a--b" text,  -- real note\n  z int);\n')
    eq([(c['name'], c['comment']) for c in s['t']['columns']],
       [('id', ''), ('a--b', 'real note'), ('z', '')],
       'the quoted -- stays in the name and the one after it is still a comment')


@case('parse: an unmatched closing bracket does not swallow the columns after it')
def _(work):
    # `paren_depth` 가 짝을 안 맞추고 그냥 세면 `]` 에서 깊이가 음수로 떨어져
    # `depth == 0` 이 뒤로 영영 거짓이 된다 — b·c 가 통째로 a 의 기본값 속으로
    # 사라졌다. 짝 없는 기호는 없는 셈 쳐야 뒤가 예전처럼 읽힌다.
    s = ddl(work, 'CREATE TABLE q ( a int DEFAULT 1], b text, c text );\n')
    eq([c['name'] for c in s['q']['columns']], ['a', 'b', 'c'],
       'an unmatched ] hides nothing behind it')


@case('parse: an unmatched opening bracket does not swallow the columns after it')
def _(work):
    # 위 케이스의 **반대쪽**이다. `max(0, depth - 1)` 로 누르는 처방은 닫는 쪽만
    # 막고 이쪽은 못 막는다 — 깊이가 1 에서 안 내려와 똑같이 뒤를 삼킨다. 두
    # 방향을 따로 재지 않으면 반만 고친 처방이 초록으로 통과한다.
    s = ddl(work, 'CREATE TABLE r ( a int DEFAULT [1, b text, c text );\n')
    eq([c['name'] for c in s['r']['columns']], ['a', 'b', 'c'],
       'an unmatched [ hides nothing behind it')


@case('parse: a quoted name cannot close a bracket the DEFAULT left open')
def _(work):
    # 위 둘(짝 안 맞는 괄호)과 그 위 셋(따옴표 이름)이 **만나는** 자리다. 어느 한
    # 쪽만 재면 못 잡는다.
    #
    # `_default_end` 가 깊이를 세기 전에 큰따옴표 이름을 덮지 않으면, 짝 없이 열린
    # `[` 가 **이름 안의** `]` 와 짝지어진다. 그러면 그 뒤의 `COLLATE` 이 깊이 1 로
    # 보여 경계에서 빠지고, 기본값이 `[1 COLLATE "a]b"` 통째가 된다 — `_default_end`
    # 주석이 지키겠다고 적어 둔 바로 그것이다(*"이름 안의 `(` 하나가 뒤따르는 제약을
    # 통째로 안 보이게 만든다"*). 덮고 세면 짝 없는 `[` 는 아무와도 안 짝지어져
    # 없는 셈이 되고, `COLLATE` 이 깊이 0 에서 잡혀 값이 `[1` 로 끊긴다.
    #
    # 이름 없이 `DEFAULT [1 COLLATE ...` 만 쓰면 두 코드가 **똑같이** 답한다 —
    # 이름 안의 닫는 기호가 있어야만 갈린다. 그래서 이름과 짝 없는 괄호를 한
    # 입력에 같이 둔다.
    s = ddl(work, 'CREATE TABLE t ( a text DEFAULT [1 COLLATE "a]b", z int );\n')
    eq([(c['name'], c['default']) for c in s['t']['columns']],
       [('a', '[1'), ('z', '')],
       'a ] inside a name closes nothing, so COLLATE is still a boundary')


@case('parse: a ] cannot close a ( that the DEFAULT left open')
def _(work):
    # `paren_depth` 가 스택으로 짝을 맞추는 이유의 **마지막 반**이다. 위 두 케이스
    # (짝 안 맞는 `]`·`[`)는 스택을 순진한 `±1` 이나 `max(0, d-1)` 로 되돌리면 함께
    # 붉어지므로 그쪽은 이미 재고 있다. 안 재던 것은 **짝의 종류를 보는 한 줄**
    # (`stack[-1][1] == c`)이다 — 그것만 지우고 아무거나 팝하게 해도 210 이 전부
    # 초록이었다.
    #
    # 종류를 안 보면 `]` 가 앞서 열린 `(` 를 닫아 **없는 쌍**이 생기고, 그 쌍이
    # 최상위 콤마를 제 안에 가둔다. 아래 입력에서 `f(` 는 끝내 안 닫히고 `[`…`]`
    # 만 진짜 쌍이라 `b` 앞 콤마는 깊이 0 이어야 하는데, 아무거나 팝하면 `(`…`]`
    # 가 쌍이 되어 그 콤마가 깊이 1 로 보인다 — b 가 통째로 a 의 기본값 속으로
    # 사라진다.
    #
    # 괄호는 짝이 맞고(`f(`+`)` 와 바깥 한 쌍) 닫는 괄호 뒤가 콤마가 아니라,
    # `parse_create` 의 버리는 두 갈래에는 닿지 않는다 — 그래서 이 입력은 정말로
    # `paren_depth` 만 묻는다.
    s = ddl(work, 'CREATE TABLE t ( a text DEFAULT f(1, b text[2), c text] );\n')
    eq([(c['name'], c['default']) for c in s['t']['columns']],
       [('a', 'f(1'), ('b', '')],
       'a ] closes no ( , so the comma before b is still top level')


# ── 경계가 깨진 CREATE TABLE ─────────────────────────────────────────────────
# 아래 다섯은 `parse_create` 가 **본문의 괄호 짝이 어긋난 문장**을 어떻게 다루는지를
# 한 갈래씩 나눠 잰다. 한 케이스에 몰지 않는 것은 일부러다 — 빨강이 났을 때 여는
# 쪽이 깨졌는지, 닫는 쪽인지, 알림인지, 되살아남인지가 이름만 보고 갈려야 한다.
#
# 다섯이 함께 못박는 것: **읽지 못한 것은 이름을 대고 버린다.** 조용히 0개로 끝내는
# 것도, 반쯤 읽은 것을 완성본처럼 싣는 것도, 버렸다고 말해 놓고 되살리는 것도 아니다.
def _ddl_run(work, text, name='a.sql'):
    """DDL 한 파일을 돌리고 (스키마, stdout) 을 함께 준다.

    `ddl()` 은 스키마만 돌려준다. 아래 케이스들은 **화면에 나간 말**이 재는 것의
    절반이라 — 조용히 버리는 것을 잡는 케이스가 여럿이다 — 출력까지 봐야 한다.
    """
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / name).write_text(text, encoding='utf-8')
    r = run('parse_ddl.py', work, sql_dir=d)
    return json.loads((work / 'schema.json').read_text(encoding='utf-8')), r.stdout


@case('parse: a CREATE TABLE whose ( never closes is named, not dropped in silence')
def _(work):
    # 여는 괄호가 끝내 안 닫히는 갈래(`if depth:`). 예전엔 `continue` 한 줄이 전부라
    # 테이블 0개에 rc 0 으로 끝났고 화면에는 단서가 한 글자도 없었다 — DDL 을 다시
    # 들여다볼 이유조차 안 남았다.
    #
    # 두 가지를 함께 본다. 버렸다는 것(테이블 0개)과 **이름을 댔다는 것**이다.
    # 이름 쪽을 빼면 알림을 통째로 들어내도 초록이다 — 실제로 그랬다.
    s, out = _ddl_run(work, 'CREATE TABLE r ( a int DEFAULT (1, b text, c text );\n')
    eq(sorted(s), [], 'a table whose body never closes is not filed')
    has(out, 'could not be read: r', 'and the name it was written under is said out loud')


@case('parse: a CREATE TABLE cut short by a stray ) is named, not filed half-read')
def _(work):
    # 남는 `)` 는 본문을 **너무 일찍** 끊는다. `a int DEFAULT 1), b text, c text` 는
    # 그 `)` 를 닫는 괄호로 읽어 b·c 가 문 밖으로 밀려났고, 컬럼 하나만 든 테이블이
    # 아무 말 없이 정의서에 실렸다 — 사람 눈에는 그냥 컬럼이 하나인 테이블이다.
    # 위 케이스와 갈래가 다르다(`) ,` 검출). 그쪽은 깊이가 남는 쪽, 이쪽은 깊이가
    # 일찍 0 이 되는 쪽이라 한쪽만 재면 다른 쪽이 통째로 무방비다.
    s, out = _ddl_run(work, 'CREATE TABLE q ( a int DEFAULT 1), b text, c text );\n')
    eq(sorted(s), [], 'a body cut short by a stray ) is not filed with the columns it kept')
    has(out, 'could not be read: q', 'and the name is said out loud')


@case('parse: the stray ) is caught with blanks between it and the comma')
def _(work):
    # 위 케이스와 **한 글자** 차이다 — `)` 와 콤마 사이의 공백. 검출이 꼬리를
    # `lstrip()` 하지 않으면 이 모양만 빠져나가, 손으로 띄어 쓴 DDL 이 다시 조용히
    # 반만 실린다. 위 케이스는 이것을 못 잰다(공백이 없으므로 둘 다 잡힌다).
    s, out = _ddl_run(work, 'CREATE TABLE q ( a int DEFAULT 1)  , b text, c text );\n')
    eq(sorted(s), [], 'blanks before the comma hide nothing')
    has(out, 'could not be read: q', 'and the name is said out loud')


@case('parse: an unclosed body stops at its own ; instead of eating the next table')
def _(work):
    # 본문이 제 문(`;`)을 넘어가던 때의 모양이다. 짝 안 맞는 `(` 가 **뒤 문장의** `)`
    # 로 맞아떨어져, 남의 `CREATE TABLE` 이 통째로 앞 테이블의 기본값 칸에 실렸다:
    #   r.a 의 기본값 = `(1, b text ); CREATE TABLE u ( a int DEFAULT 1)`
    # 그러면서 r 은 **성한 테이블처럼** 정의서에 실리고 경고도 u 것 하나뿐이었다.
    #
    # 그래서 세 가지를 본다. ⑴ 두 문 다 버려진다 ⑵ 둘 **다** 이름이 불린다
    # ⑶ 뒤따르는 성한 `ok` 는 멀쩡히 살아남는다. ⑶ 이 없으면 '어긋난 문 하나가
    # 파일의 나머지를 다 삼키는' 처방도 초록으로 통과한다.
    s, out = _ddl_run(work,
                      'CREATE TABLE r ( a int DEFAULT (1, b text );\n'
                      'CREATE TABLE u ( a int DEFAULT 1), b text );\n'
                      'CREATE TABLE ok (id int);\n')
    eq(sorted(s), ['ok'], 'neither broken statement is filed, and the healthy one is')
    eq([c['name'] for c in s['ok']['columns']], ['id'], 'the healthy table is whole')
    has(out, 'could not be read: r (a.sql), u (a.sql)', 'both names are said, in file order')


@case('parse: a table that could not be read is not raised again by a later ALTER')
def _(work):
    # 버렸다고 **말한** 것이 되살아나던 자리다. `parse_alter` 는 없는 테이블을 만나면
    # 스텁을 세우는데, 이름만 보던 판은 방금 버린 q 를 컬럼 0개짜리로 다시 세웠다.
    # 세 가지가 한꺼번에 틀렸다: 빠진다고 말해 놓고 안 빠졌고, 배지가 `new` 에서
    # `existing` 으로 뒤집혔고, `ddl_no_db` 가 'ERD_DB/ERD_PSQL 을 주면 채워진다' 고
    # 안내해 있지도 않은 접속 문제를 쫓게 했다 — 신규 테이블이라 DB 에 있을 리 없다.
    #
    # 마지막 한 줄이 이 케이스가 **혼자** 재는 것이다. 스텁이 서는 것만 보면 안내
    # 문구는 따로 안 재어지는데, 사람이 실제로 시간을 버리는 자리는 그 문구다.
    s, out = _ddl_run(work, 'CREATE TABLE q ( a int DEFAULT 1), b text );\n'
                            'ALTER TABLE q ADD COLUMN z int;\n')
    eq(sorted(s), [], 'the ALTER does not raise the table its CREATE could not be read')
    has(out, 'could not be read: q', 'the name is said out loud')
    eq('never defined in the DDL' in out, False,
       'and a table nobody could read is not explained away as a missing connection')


@case('parse: WITH and PARTITION BY tables are read with no warning at all')
def _(work):
    # 위 다섯의 **반대쪽 못**이다. 그쪽은 전부 '버려야 할 것을 버리는가' 를 묻는데,
    # 그것만 재면 검출을 넓히는 쪽으로 틀린 처방이 다 초록이다 — 성한 DDL 을 버려도
    # 아무도 안 묻기 때문이다. 실제로 '콤마가 **어딘가** 있다' 로 넓힌 판은 아래 두
    # 테이블을 통째로 버렸다.
    #
    # 경고가 **0줄** 인 것까지 보는 것이 요점이다. 버리지는 않고 경고만 더 내는 판
    # (넓힌 검출로 이름만 적고 계속 읽는 것)은 표에 남는 것이 하나도 안 달라져,
    # 컬럼만 세는 케이스로는 210 이 전부 초록이었다. 없는 고장을 알리는 것도 고장이다.
    s, out = _ddl_run(work,
                      'CREATE TABLE t (a int) WITH (fillfactor = 70, '
                      'autovacuum_enabled = false);\n'
                      'CREATE TABLE p (a int, b int) PARTITION BY RANGE (a, b);\n')
    eq([(k, [c['name'] for c in v['columns']]) for k, v in sorted(s.items())],
       [('p', ['a', 'b']), ('t', ['a'])], 'both healthy tables are read whole')
    eq([ln for ln in out.splitlines() if 'could not be read' in ln], [],
       'and neither of them is cried over')


@case('parse: ADD COLUMN IF NOT EXISTS written twice adds the column once')
def _(work):
    # 두 마이그레이션 파일이 같은 `ADD COLUMN IF NOT EXISTS` 를 들고 있는 것은 흔한
    # 모양인데, 그냥 append 하던 동안 정의서에 같은 컬럼이 두 줄로 실렸다. 실제
    # PostgreSQL 은 컬럼이 이미 있으면 그 문을 **통째로** 건너뛰므로(NOTICE 한 줄)
    # 먼저 적힌 것이 이긴다.
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / 'a.sql').write_text('CREATE TABLE t (id bigint PRIMARY KEY);\n'
                             'ALTER TABLE t ADD COLUMN IF NOT EXISTS code text UNIQUE;\n',
                             encoding='utf-8')
    (d / 'b.sql').write_text('ALTER TABLE t ADD COLUMN IF NOT EXISTS code text UNIQUE;\n',
                             encoding='utf-8')
    run('parse_ddl.py', work, sql_dir=d)
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq([c['name'] for c in s['t']['columns']], ['id', 'code'],
       'a column already present is not added a second time')


@case('parse: a repeated ADD COLUMN drops its REFERENCES along with the column')
def _(work):
    # 위 케이스와 같은 모양인데 **재는 자리가 다르다.** 위 것은 컬럼 줄만 보므로,
    # 건너뛰기를 "컬럼만 안 붙이고 제약은 그대로 받는" 반쪽으로 바꿔도 그대로
    # 초록이다. 그 반쪽이 바로 `parse_ddl` 주석이 금지한 것이다 — 실제 PostgreSQL
    # 은 `IF NOT EXISTS` 에 걸리면 **문 하나를 통째로** 건너뛰므로 뒤 문의
    # REFERENCES 도 안 붙는다.
    #
    # 왜 UNIQUE 로는 못 재는가: 인라인 UNIQUE 는 아래 케이스가 재는 **다른**
    # 가드(`[cname] not in t['uniques']`)에 한 번 더 걸려 증상이 가려진다.
    # FK 에는 그 가드가 없다 — 두 벌이 그대로 실려 관계선이 겹쳐 그려진다.
    # 그래서 같은 중복을 `REFERENCES` 로 쓴 케이스가 따로 있어야 한다.
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / 'a.sql').write_text('CREATE TABLE users (id bigint PRIMARY KEY);\n'
                             'CREATE TABLE t (id bigint PRIMARY KEY);\n'
                             'ALTER TABLE t ADD COLUMN IF NOT EXISTS uid bigint '
                             'REFERENCES users(id);\n',
                             encoding='utf-8')
    (d / 'b.sql').write_text('ALTER TABLE t ADD COLUMN IF NOT EXISTS uid bigint '
                             'REFERENCES users(id);\n',
                             encoding='utf-8')
    run('parse_ddl.py', work, sql_dir=d)
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq([c['name'] for c in s['t']['columns']], ['id', 'uid'],
       'the second statement adds no second column')
    eq([(f['column'], f['ref_table']) for f in s['t']['fks']], [('uid', 'users')],
       'the skipped statement leaves no relationship behind either')


@case('parse: an inline UNIQUE does not repeat a UNIQUE the table already declares')
def _(work):
    # 위 케이스가 막는 길(컬럼 중복)을 **비켜 가는** 자리다: 테이블 수준
    # `UNIQUE (code)` 만 있고 컬럼은 나중에 ALTER 로 붙으면, 컬럼은 처음 보는
    # 것이라 건너뛰기에 안 걸리는데 인라인 UNIQUE 는 이미 있는 것과 겹친다.
    # 가드가 없으면 uniques 에 같은 것이 두 벌 실려 정의서가 제약을 두 번 적는다.
    d = work / 'sql'
    d.mkdir(parents=True, exist_ok=True)
    (d / 'a.sql').write_text('CREATE TABLE t (id bigint, UNIQUE (code));\n',
                             encoding='utf-8')
    (d / 'b.sql').write_text('ALTER TABLE t ADD COLUMN code text UNIQUE;\n',
                             encoding='utf-8')
    run('parse_ddl.py', work, sql_dir=d)
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq(s['t']['uniques'], [['code']],
       'the same one-column UNIQUE is listed once however many statements spell it')


# ── 이름 안의 이스케이프 큰따옴표 (`"a""b"` = 이름 `a"b`) ────────────────────
# Postgres 는 따옴표 친 이름 **안**의 따옴표를 두 개로 적는다. `_QID` 가 그것을
# 못 넘던 동안 이름이 첫 `""` 에서 끊겨, 그 컬럼이 — 헤더면 그 `CREATE TABLE`
# 통째가 — **경고 한 줄 없이** 사라졌다. 아래 셋이 그 세 자리를 나눠 못박는다.
#
# 나머지 넷은 그 고침이 **딸려 온 대가**를 지킨다. `_QID` 는 갈래가 둘인데
# (`parse_ddl.py` 의 `_QID` 옆 주석) 그 모양은 전부 지우거나 느슨하게 해도 217개가
# 초록이던 자리라, 하나씩 못을 박는다: ②갈래가 **있는가**, ①갈래가 줄을 **안**
# 넘는가, ②갈래가 줄을 **반드시** 넘는가. 마지막 하나는 갈래가 아니라 판단을
# 지킨다 — `blank_quoted`·`_scan` 을 `""` **모르는 채로 두는** 것이다.
@case('parse: an escaped double quote in a column name survives')
def _(work):
    # `"[^"]*"` 로 이름을 재던 때는 `"a""b"` 에서 `"a"` 만 집히고 남은 `"b"` 가 뒤
    # 규칙과 어긋나, 그 컬럼이 아무 말 없이 빠졌다 — HEAD 는 `['id', 'z']` 였다.
    s = ddl(work, 'CREATE TABLE t (id int, "a""b" text, z int);\n')
    eq([(c['name'], c['type']) for c in s['t']['columns']],
       [('id', 'int'), ('a"b', 'text'), ('z', 'int')],
       'the escaped quote is one character of the name, and the type still follows it')


@case('parse: a table named with an escaped double quote is not lost')
def _(work):
    # 같은 고장이 **헤더**에서 나면 잃는 것이 컬럼 하나가 아니라 문 하나다.
    # HEAD 는 이 파일에서 테이블을 `{}` 로 냈다 — 경고도 없이.
    s = ddl(work, 'CREATE TABLE "a""b" (id int primary key, z int);\n')
    eq(sorted(s), ['a"b'], 'the statement yields its table, under the name Postgres gives it')
    eq([c['name'] for c in s['a"b']['columns']], ['id', 'z'], 'with both columns')


@case('parse: an escaped name is reachable from every clause')
def _(work):
    # 위 둘은 `CREATE TABLE` 의 두 자리만 본다. 이름 조각은 그 밖에도 스무 자리
    # 넘게 쓰이는데, 자리마다 **다른 규칙**을 쓰면 반만 고친 판이 초록으로 지나간다.
    #
    # HEAD 가 딱 그 모양이었다: PK·UNIQUE·FK 는 `_names`/`_INPAREN` 으로 읽어
    # `a"b` 를 제대로 적었는데 **컬럼 목록에는 그 이름이 없었다** — 정의서가 있지도
    # 않은 컬럼을 PK 로 적는다. 그래서 마지막 확인이 이 케이스의 알맹이다:
    # 제약에 적힌 이름은 전부 컬럼 목록 안에 있어야 한다.
    #
    # (`ALTER TABLE … ALTER COLUMN … SET NOT NULL` 은 일부러 안 넣었다 — 이 파서는
    #  그 문을 이름이 어떻든 아예 안 읽으므로, 넣어 봐야 이름에 대해 아무것도 안
    #  재고 '재는 척' 만 한다. 대신 `ALTER TABLE` 의 **테이블 이름** 자리에 걸었다.)
    s = ddl(work,
            'CREATE TABLE p ("a""b" int PRIMARY KEY);\n'
            'CREATE TABLE "t""x" (\n'
            '  id int,\n'
            '  "a""b" int NOT NULL,\n'
            '  "c""d" int,\n'
            '  PRIMARY KEY ("a""b"),\n'
            '  UNIQUE ("c""d"),\n'
            '  FOREIGN KEY ("a""b") REFERENCES p ("a""b")\n'
            ');\n'
            'COMMENT ON COLUMN "t""x"."a""b" IS \'note\';\n'
            'ALTER TABLE "t""x" ADD COLUMN "e""f" int;\n'
            'ALTER TABLE "t""x" ADD CONSTRAINT fk2 FOREIGN KEY ("e""f") '
            'REFERENCES p ("a""b");\n'
            'CREATE UNIQUE INDEX ix ON "t""x" ("e""f");\n')
    eq(sorted(s), ['p', 't"x'], 'both tables are read')     # HEAD: ['p'] 하나뿐이었다
    eq([c['name'] for c in s['p']['columns']], ['a"b'],
       'a table whose only column is an escaped name is not an empty box')
    t = s['t"x']
    eq([(c['name'], c['comment']) for c in t['columns']],
       [('id', ''), ('a"b', 'note'), ('c"d', ''), ('e"f', '')],
       'CREATE TABLE, COMMENT ON COLUMN and ADD COLUMN all find it')
    eq(t['pk'], ['a"b'], 'PRIMARY KEY names it')
    eq(t['uniques'], [['c"d'], ['e"f']], 'UNIQUE and CREATE UNIQUE INDEX name it')
    eq([(f['column'], f['ref_table'], f['ref_column']) for f in t['fks']],
       [('a"b', 'p', 'a"b'), ('e"f', 'p', 'a"b')],
       'and REFERENCES names it on both ends, from CREATE and from ALTER alike')
    have = {c['name'] for c in t['columns']}
    named = set(t['pk']) | {c for u in t['uniques'] for c in u} \
        | {f['column'] for f in t['fks']}
    eq(sorted(named - have), [],
       'no constraint names a column the table does not have')


@case('parse: a quoted name containing a newline still parses')
def _(work):
    # `_QID` ②갈래(`"[^"]*\n[^"]*"`)의 지킴이다. ①갈래는 속도 때문에 줄을 안 넘어서,
    # ②를 빼면 **HEAD 가 읽던 것**을 못 읽는다 — 이 테이블이 통째로 사라진다.
    # 문법으로 되는 이름이고 HEAD 가 냈으므로, 고치면서 잃으면 그것이 회귀다.
    s = ddl(work, 'CREATE TABLE "a\nb" (id int primary key);\n')
    eq(sorted(s), ['a\nb'], 'a newline inside a quoted name costs no statement')


@case('parse: an unmatched quote does not swallow the statement after it')
def _(work):
    # `_QID` ①갈래에 **`\n` 을 못박은** 것의 지킴이다. ①이 줄을 넘을 수 있으면 짝
    # 없는 `"""` 하나가 닫는 따옴표를 다음 줄 너머에서 찾아, **두 문을 통째로**
    # 한 이름으로 먹는다 — 뒤의 성한 `c` 가 사라진다(HEAD 는 냈다).
    #
    # 모양을 가린다. 삼킴은 ①이 넘어간 자리에서 **뒤 문의 따옴표가 짝을 채워 줄
    # 때만** 벌어진다. `"""` 테이블 하나 뒤에 평범한 `CREATE TABLE c` 를 두는
    # (가장 자연스럽게 적을) 모양은 채울 따옴표가 없어 ①의 `\n` 을 지워도 안 붉는다.
    # 컬럼 이름·PK 목록·인덱스·COMMENT 로 쓴 여섯 모양도 마찬가지였다. 무는 것은
    # 뒤 문이 `REFERENCES """` 로 따옴표를 하나 더 내놓는 이 모양과, 짝 없는 `"""`
    # 테이블이 둘 연달아 오는 모양뿐이다 — 여덟 모양을 물려 보고 고른 것이다.
    #
    # 이름 있는 테이블만 센다. `REFERENCES """` 가 이름이 빈 상자를 하나 세우는데,
    # 그것은 지금도 HEAD 도 같은 모양이라 이 케이스가 재는 것이 아니다 —
    # 구멍 ①(짝 없는 따옴표를 조용히 잃는 것)의 자리이지 이 못의 자리가 아니다.
    s, out = _ddl_run(work,
                      'CREATE TABLE """ (id int);\n'
                      'CREATE TABLE c (id int, p int REFERENCES """ (id));\n')
    eq([n for n in sorted(s) if n], ['c'],
       'the healthy statement after an unmatched quote is still read')
    eq([col['name'] for col in s['c']['columns']], ['id', 'p'], 'with both its columns')
    eq([n for n in s if '\n' in n], [],
       'and no table is filed under a name made of two swallowed statements')


@case('parse: an unclosed name list packed with "" finishes in time')
def _(work):
    # `_QID` ②갈래에 **`\n` 을 못박은** 것의 지킴이다. 그 못이 없으면 두 갈래가
    # 줄바꿈 없는 이름을 **똑같이** 물어, `_INPAREN` 의 `(?:이름|[^)])+\)` 가 짝
    # 없는 괄호에서 되짚어 올 때 갈래가 겹친 만큼 글자 수에 지수로 터진다.
    # 무는 것은 같으므로 결과로는 못 잰다 — **시간**으로만 잰다.
    #
    # 잰 값(이 저장소, 2026-08-05): 지금 0.52초 · 겹치는 판 48초. 문턱 2초는 그
    # 사이 어디든 좋게 넉넉하다. 실패가 '느리다' 로 보이도록 rc 가 아니라 시계를
    # 본다 — 겹치는 판에서는 이 케이스 하나가 시험 전체보다 오래 걸린다.
    import time
    sql = ('CREATE TABLE t (id int);\n'
           'ALTER TABLE t ADD CONSTRAINT k PRIMARY KEY (' + '""a' * 14 + ';\n')
    began = time.perf_counter()
    s = ddl(work, sql)
    took = time.perf_counter() - began
    eq(sorted(s), ['t'], 'the healthy table is still read')
    if took > 2.0:
        raise Fail(f'a 14-deep "" name list took {took:.1f}s (ceiling 2.0s) — the two '
                   f'_QID branches are matching the same text again; see the comment '
                   f'beside _QID in parse_ddl.py')


@case('parse: an unmatched " before a "" is masked from the left')
def _(work):
    # 이름 조각(`_QID`)은 `""` 를 넘게 고쳤지만 **덮는 쪽**(`blank_quoted`·`_scan`)은
    # 일부러 안 고쳤다 — 성한 입력에서는 어느 쪽이든 결과가 같아서다. 그 '안 고친다'
    # 는 판단을 아무도 안 지키고 있었다: 둘을 `""` 아는 판으로 바꿔도 217이 전부
    # 초록이었다. 그러면 다음 사람이 주석만 지우고 고쳐 놓아도 아무 데서도 안 붉는다.
    #
    # 갈리는 것은 **짝이 안 맞는** 따옴표가 든 입력이다. 아래 `"a(""b` 는 따옴표가
    # 셋이라, 지금 판은 앞의 둘을 짝지어 덮어 그 사이의 `(` 를 가린다. `""` 를 아는
    # 판은 첫 따옴표가 `""` 를 넘어가다 닫는 짝을 못 찾아 **아무것도 안 덮고**, 그
    # `(` 가 본문 괄호로 세어져 짝이 영영 안 닫힌다 — 테이블이 통째로 버려진다.
    # (`parse: an opening paren inside a quoted name…` 과 같은 고장, 다른 입력이다.)
    s, out = _ddl_run(work, 'CREATE TABLE t (id int, c text COLLATE "a(""b, z int);\n')
    eq(sorted(s), ['t'], 'the statement still yields its table')
    eq([c['name'] for c in s['t']['columns']], ['id', 'c', 'z'],
       'a ( between an unmatched " and a "" is still masked, so the body closes')
    eq([ln for ln in out.splitlines() if 'could not be read' in ln], [],
       'and nothing is cried over')


# ── 이름 안의 `;` 가 문을 끊던 자리 (`parse_alter` 의 `_stmts`) ───────────────
# `parse_alter` 는 문을 `(ADD\s+COLUMN\b.*?);` 로 **첫 `;` 까지** 잘랐다. 그런데
# Postgres 에서 `"a;b"` 는 `a;b` 라는 이름 하나다 — 이름 안의 `;` 에서 문이 끊겼다.
# `parse_create` 는 같은 일을 `blank_quoted` 로 이름까지 덮은 사본에서 재고 있었고,
# 이 자리만 안 그랬다. 고침은 그 사본(`msc_q`)에서 끝을 재는 것이다.
#
# **자르는 자리는 둘이다** — ADD COLUMN 루프와 ADD CONSTRAINT 루프. 한쪽만 되돌리는
# 뮤턴트를 만들어 보면 두 부류가 서로를 **전혀 못 잡는다**(직접 잰 격자, 2026-08-05):
#
#            ADD COLUMN 만 되돌림   ADD CONSTRAINT 만 되돌림   둘 다(=고침 전체 되돌림)
#   아래 ①              빨강                 초록                        빨강
#   아래 ②              빨강                 초록                        빨강
#   아래 ③              초록                 빨강                        빨강
#   아래 ④              빨강                 빨강                        빨강
#
# 그래서 ①·②(ADD COLUMN 쪽)만 아무리 늘려도 ADD CONSTRAINT 쪽은 무방비다. ③이
# 그 자리의 유일한 못이다. 여기서 하나를 뺄 때는 그 열이 통째로 초록이 되는지
# **먼저 확인해라** — 이 고침은 되돌려도 `all 224 passed` 이던 자리다.
#
# 이름 안의 줄바꿈·연속공백이 접히는 것(`"a\nb"` → `a b`)은 **HEAD 부터의 성질이고
# 이 고침과 무관**하므로 일부러 안 박았다. 지금 동작을 못박으면 그것을 고치려는
# 사람이 여기서 막힌다.
@case('parse: a ; inside a quoted name does not cut the ADD COLUMN statement short')
def _(work):
    # ① 잃는 쪽. 끊긴 자리 뒤가 통째로 안 읽히므로 `"a;b"` 만이 아니라 **같은 문에
    # 이어 적은 성한 `memo` 까지** 함께 사라졌다 — 경고 한 줄 없이. HEAD 는 이
    # 파일에서 `['id']` 를 냈다.
    s = ddl(work, 'CREATE TABLE t (id bigint PRIMARY KEY);\n'
                  'ALTER TABLE t ADD COLUMN "a;b" int, ADD COLUMN memo text;\n')
    eq([(c['name'], c['type']) for c in s['t']['columns']],
       [('id', 'bigint'), ('a;b', 'int'), ('memo', 'text')],
       'a quoted name is one name however many ; it holds, and the column written '
       'after it in the same statement is read too')


@case('parse: a statement spelled inside a quoted name does not reach another table')
def _(work):
    # ② **짓는** 쪽 — 같은 고장의, 잃는 것보다 나쁜 얼굴이다. 이름 안에 문을 통째로
    # 적으면 예전 자르기는 이름 속 `;` 에서 문을 끝내고 그 **다음 글자부터** 다시
    # 머리를 찾아, 이름의 나머지를 진짜 문으로 읽었다: 남의 테이블 `u` 에 아무 데도
    # 안 적힌 컬럼 `z` 가 붙었다(HEAD 실측: u = ['id', 'z']).
    # PG16 대조 — `u` 에는 아무것도 안 붙고, `t` 에 그 긴 이름의 컬럼 하나가 붙는다.
    s = ddl(work, 'CREATE TABLE t (id int);\n'
                  'CREATE TABLE u (id int);\n'
                  'ALTER TABLE t ADD COLUMN '
                  '"a int; ALTER TABLE u ADD COLUMN z text" bigint;\n')
    eq([c['name'] for c in s['u']['columns']], ['id'],
       'a table named only inside another table quoted name gains nothing')
    eq([c['name'] for c in s['t']['columns']],
       ['id', 'a int; ALTER TABLE u ADD COLUMN z text'],
       'the whole quoted text is the one column name it is')


@case('parse: a ; inside a quoted name does not misname what ADD CONSTRAINT reads')
def _(work):
    # ③ **ADD CONSTRAINT 루프의 유일한 못이다.** 위 둘은 이 자리에서 초록이다 —
    # 루프가 둘이고 각자 제 몫만 자르기 때문이다(위 격자).
    #
    # 여기서도 잃는 쪽과 짓는 쪽이 함께 나온다. PK·UNIQUE 는 이름 속 `;` 뒤가 잘려
    # **빈 채로** 남고(HEAD: pk []), FK 는 더 나쁘다 — `REFERENCES parts ("part;no")`
    # 가 `("part` 에서 끊겨 참조 컬럼을 못 읽자 `_refs` 의 기본값이 켜져서, **`parts`
    # 에 있지도 않은 `id`** 를 가리키는 관계가 지어졌다(HEAD 실측: ('pn','parts','id')).
    # 정의서와 ERD 가 없는 컬럼으로 선을 긋는다.
    s = ddl(work,
            'CREATE TABLE parts ("part;no" int PRIMARY KEY, "lot;id" int);\n'
            'CREATE TABLE orders (pn int, "ord;no" int);\n'
            'ALTER TABLE orders ADD CONSTRAINT pk PRIMARY KEY ("ord;no");\n'
            'ALTER TABLE parts ADD CONSTRAINT uq UNIQUE ("lot;id");\n'
            'ALTER TABLE orders ADD CONSTRAINT fk FOREIGN KEY (pn) '
            'REFERENCES parts ("part;no");\n')
    # 짓는 쪽을 먼저 본다 — 하나가 붉으면 뒤는 안 돌므로, 더 나쁜 얼굴이 화면에 뜬다.
    eq([(f['column'], f['ref_table'], f['ref_column']) for f in s['orders']['fks']],
       [('pn', 'parts', 'part;no')],
       'REFERENCES points at the column that is actually there')
    have = {c['name'] for c in s['parts']['columns']}
    eq(sorted(f['ref_column'] for f in s['orders']['fks'] if f['ref_column'] not in have),
       [], 'no relationship names a column the referenced table does not have')
    eq(s['orders']['pk'], ['ord;no'], 'PRIMARY KEY reads the whole quoted name')
    eq(s['parts']['uniques'], [['lot;id']], 'and so does UNIQUE')


@case('parse: ALTER heads with no ; after them do not cost the square of the file')
def _(work):
    # ④ 고침이 **딸려 온 값**의 지킴이다. 예전 자르기는 머리마다 `.*?;` 로 파일
    # 끝까지 훑고 실패했다 — 뒤에 `;` 가 하나도 없으면 머리 수 × 남은 길이, 곧
    # O(n²) 다. 지금은 `find(';')` 가 한 번에 -1 을 내고 그 자리에서 멎는다.
    # 결과는 두 판이 **똑같다**(어느 쪽도 이 문들을 안 읽는다) — 시간으로만 갈린다.
    #
    # `;` 를 잃은 덤프·이어붙이다 깨진 파일에서 나오는 모양이다. 지금 것은 '아무것도
    # 못 읽었다' 를 1초 안에 말하고, 예전 것은 같은 말을 하는 데 분 단위로 갔다.
    #
    # 잰 값(이 저장소, 2026-08-05, 12,000문): 지금 0.45초 · ADD COLUMN 쪽만 되돌린
    # 판 7.7초 · ADD CONSTRAINT 쪽만 되돌린 판 7.9초 · 둘 다 되돌린 판 14.9초.
    # 문턱 3.0초는 그 사이다 — 이 기계보다 **6배 느린** 기계에서도 안 붉고, 한쪽만
    # 되돌린 판은 **2.5배 빠른** 기계에서도 붉는다. 두 루프를 다 걸도록 두 종류를
    # 섞어 적는다: ADD COLUMN 만 적으면 ADD CONSTRAINT 쪽 열이 초록이 된다.
    import time
    sql = ['CREATE TABLE t (id int);\n']
    for i in range(6000):
        sql.append(f'ALTER TABLE t ADD COLUMN c{i} int\n')
        sql.append(f'ALTER TABLE t ADD CONSTRAINT k{i} UNIQUE (id)\n')
    began = time.perf_counter()
    s = ddl(work, ''.join(sql))
    took = time.perf_counter() - began
    eq(sorted(s), ['t'], 'the healthy CREATE TABLE is still read')
    eq([c['name'] for c in s['t']['columns']], ['id'],
       'and a statement that never reaches a ; is read by neither the old nor the '
       'new cut — this case measures the time, not the reading')
    if took > 3.0:
        raise Fail(f'12,000 ALTER heads with no ; took {took:.1f}s (ceiling 3.0s) — '
                   f'a statement end is being looked for with a scan per head again; '
                   f'see _stmts in parse_ddl.py')


# ── 설명 ─────────────────────────────────────────────────────────────────────
@case('merge_desc: common dictionary follows ERD_LANG')
def _(work):
    write_schema(work, {'t': table('t', [col('id'), col('created_at', 'timestamptz')])})
    run('merge_desc.py', work, env={'ERD_LANG': 'ko'})
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    got = [c['comment'] for c in s['t']['columns'] if c['name'] == 'created_at']
    eq(got, ['생성 일시'], 'COMMON dictionary is localized')


@case('merge_desc: a second run does not relabel where descriptions came from')
def _(work):
    write_schema(work, {'t': table('t', [col('id'), col('x', 'text', comment='given')])})
    a = run('merge_desc.py', work).stdout
    b = run('merge_desc.py', work).stdout
    def src(out):
        return re.search(r'\{.*\}', out).group(0)
    eq(src(b), src(a), 'provenance stats are stable across runs')


@case('merge_desc: descriptions come back from a previous edition')
def _(work):
    write_schema(work, {'t': table('t', [col('id'), col('x', 'text')])})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    doc = work / 'T.html'
    doc.write_text(doc.read_text(encoding='utf-8').replace(
        '<td></td></tr>', '<td>polished by hand</td></tr>', 1), encoding='utf-8')
    write_schema(work, {'t': table('t', [col('id'), col('x', 'text')])})
    run('merge_desc.py', work, env={'ERD_DOC_HTML': str(doc)})
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    joined = ' '.join(c['comment'] for c in s['t']['columns'])
    has(joined, 'polished by hand', 'a hand-edited description survives regeneration')


@case('merge_desc: a heading with no table does not swallow the next table descriptions')
def _(work):
    # 한 몫의 끝을 `</table>` 로 잡던 동안, 표가 없는 h4 는 **다음 h4 의 표까지**
    # 한 덩어리로 먹었고 finditer 는 먹은 자리를 건너뛰었다 — 그 테이블의 설명이
    # 통째로, 아무 말 없이 유실됐다. 설명을 지키자는 기능이 설명을 잃는 셈이었다.
    # 흔한 스키마 정의서에는 표 없는 제목(개요·변경 이력)이 실제로 섞여 있다.
    work.mkdir(parents=True, exist_ok=True)
    doc = work / 'prev.html'
    doc.write_text(
        '<html><body>\n'
        '<h4>overview</h4><p>this heading carries no table at all</p>\n'
        '<h4>beta</h4><table><tr><td>1</td><td>x</td><td>kept by hand</td></tr></table>\n'
        '</body></html>\n', encoding='utf-8')
    write_schema(work, {'beta': table('beta', [col('id'), col('x', 'text')])})
    run('merge_desc.py', work, env={'ERD_DOC_HTML': str(doc)})
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    got = [c['comment'] for c in s['beta']['columns'] if c['name'] == 'x']
    eq(got, ['kept by hand'],
       'a table listed after a table-less heading keeps its descriptions')


@case('merge_desc: a description filled in by hand leaves the "no description" list')
def _(work):
    # 앞선 실행이 'none' 이라 적어 둔 자리에 설명이 생기면 그 사이에 사람이
    # schema.json 을 손으로 채운 것이다. 그것을 그대로 물려받으면 설명이 **있는**
    # 컬럼이 몇 번을 돌려도 '아직 설명 없는 컬럼' 목록에 오른다.
    write_schema(work, {'t': table('t', [col('id'), col('x', 'text')])})
    first = run('merge_desc.py', work).stdout
    has(first, 't.x', 'a column with no description is listed as one to fill in')
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    for c in s['t']['columns']:
        if c['name'] == 'x':
            c['comment'] = 'typed in by a person'
    (work / 'schema.json').write_text(json.dumps(s, ensure_ascii=False),
                                      encoding='utf-8')
    second = run('merge_desc.py', work).stdout
    if re.search(r'^\s+t\.x\b', second, re.M):
        raise Fail('a column a person has already described is still listed as '
                   f'having none:\n{second}')


# ── 서버가 다르고, 값이 험하고, 조회가 죽는 날 ────────────────────────────────
@case('introspect: a server without conparentid keeps its foreign keys')
def _(work):
    # conparentid 는 PG 11 부터다. 10 이하에 그대로 물으면 FK 조회가 통째로 실패하는데
    # 요약은 그것을 'FK 0' 이라고 참인 양 찍었다 — 관계가 하나도 없는 문서가 exit 0.
    r = run('introspect.py', work, env=db_fake(work, FAKE_PG_VER='100021'))
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq([(f['ref_table'], f['on_delete']) for f in s['tricky']['fks']],
       [('owners', 'CASCADE')], 'an old server must still report its foreign keys')
    if 'query failed' in r.stdout:
        raise Fail(f'the FK query must not be sent as-is to an old server:\n{r.stdout}')
    sys.path.insert(0, str(HERE))
    import introspect
    has(introspect.q_fk(110000), 'conparentid', 'PG 11+ still filters partition copies')
    if 'conparentid' in introspect.q_fk(100021):
        raise Fail('PG 10 must not be asked for a column it does not have')


@case('introspect: a control character in a value cannot forge a row')
def _(work):
    # 값 하나에 든 \x1e 가 행을 쪼개 테이블 1개짜리 DB 를 4개로 읽혔다. 구분자를 무엇으로
    # 골라도 값이 그 바이트를 담을 수 있다 — 전송이 구분자를 쓰지 않아야 끝나는 이야기다.
    run('introspect.py', work, env=db_fake(work))
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq(sorted(s), ['claims', 'owners', 'tricky'], 'no ghost table from \x1e in a value')
    got = {c['name']: (c['default'], c['comment']) for c in s['tricky']['columns']}
    eq(got['ctrl'], ("'has\x1erecord\x1esep'::text",
                     'comment with RS and FS inside'),
       'the value arrives whole and does not shift into the next column')
    eq(got['note'][0], "'line1\nline2'::text", 'and a newline still survives too')


@case('introspect: a half-read database is refused, not documented')
def _(work):
    # 일곱 조회 중 하나만 죽어도 예전엔 완성본이 나왔다: Q_PK 가 죽으면 모든 pk 가 [],
    # 요약은 멀쩡, exit 0. 몇 행 흘린 뒤 죽은 조회는 경고조차 없었다.
    r = run('introspect.py', work, expect_ok=False,
            env=db_fake(work, FAKE_FAIL='PRIMARY KEY'))
    if 'Traceback' in r.stderr:
        raise Fail(f'should be a message:\n{r.stderr[-300:]}')
    has(r.stdout, 'database query failed',
        'a query that died after streaming rows must still warn')
    has(r.stdout + r.stderr, 'primary keys', 'the message names what could not be read')
    if (work / 'schema.json').exists():
        raise Fail('a half-read schema must not be written as if it were complete')

    # 없어도 그림이 나오는 부가정보는 멈추지 않는다 — 대신 무엇이 빠졌는지 이름을 댄다
    r = run('introspect.py', work, env=db_fake(work, FAKE_FAIL='pg_indexes'))
    has(r.stdout, 'indexes', 'an optional query names itself in the summary when it fails')
    if not (work / 'schema.json').exists():
        raise Fail('an optional query must not stop the run')


# ── spec ────────────────────────────────────────────────────────────────────
@case('spec: a typo names what is wrong instead of a traceback')
def _(work):
    write_schema(work, {'t': table('t', [col('id')])})
    (work / 'erd.spec.json').write_text(json.dumps(
        {'areas': [['A', 'ok', 'public', ['t', 'ghost']], ['B', 'empty', 'public', []]]}), encoding='utf-8')
    out = run('build_erd.py', work).stdout
    has(out, 'ghost', 'the missing table is named')
    (work / 'erd.spec.json').write_text('{ "areas": [', encoding='utf-8')
    r = run('build_erd.py', work, expect_ok=False)
    if 'Traceback' in r.stderr:
        raise Fail(f'broken spec should not traceback:\n{r.stderr[-400:]}')


@case('spec: a bad layer colour is refused, not pasted into the document')
def _(work):
    write_schema(work, {'t': table('t', [col('id')])})
    (work / 'erd.spec.json').write_text(json.dumps(
        {'areas': [['A', 'ok', 'public', ['t']]],
         'layers': {'A': ['#12345', '#35507D', '#4A80C0', 'x']}}), encoding='utf-8')
    r = run('build_erd.py', work, expect_ok=False)
    if 'Traceback' in r.stderr:
        raise Fail('bad colour should be a message, not a traceback')
    has(r.stdout + r.stderr, '#12345', 'the offending value is shown')


# ── 검증을 검증한다 ──────────────────────────────────────────────────────────
# 그림 품질에 대한 이 스킬의 주장은 전부 자체검증 한 줄에 얹혀 있다. 그래서 **그 줄을
# 읽는 쪽**이 눈을 감으면 나머지 시험이 다 통과해도 아무것도 지켜지지 않는다. 실제로
# 그 줄에 (허용)·[경고] 꼬리가 붙자 숫자를 긁던 정규식이 마지막 항목을 놓쳤고, 가로선
# 중첩이 44 여도 render 항목은 전부 통과라고 했다. 여기 담는 것은 그 자리다.
@case('verify: a clean verdict is computed here, not taken from the record')
def _(work):
    # 이 케이스가 막는 것은 그림 버그가 아니라 **시험이 눈을 감는 것**이다. 검사하는
    # 쪽이 재는 쪽의 판정(warn·tolerated)을 다시 믿기 시작하면, 아래 세 가지가 전부
    # 통과로 돌아선다 — 실제로 셋 다 한 번씩 전부 통과였다.
    honest = {'file': 'erd_area_A.png',
              'counts': {'label_table': 0, 'label_x': 0, 'thru': 0,
                         'v_overlap': 0, 'h_overlap': 2},
              'tolerated': [], 'warn': ['h_overlap']}
    if not verify_faults(honest):
        raise Fail('a real regression must be a fault')
    # ① 재는 쪽이 경고를 지워도 판정은 그대로여야 한다
    if not verify_faults({**honest, 'warn': []}):
        raise Fail('an empty warn list must not make a nonzero count clean')
    # ② tolerate 를 전 항목으로 넓혀도 마찬가지다
    if not verify_faults({**honest, 'warn': [], 'tolerated': list(MEASURES)}):
        raise Fail('the code tolerating a check must not silence the test')
    # ③ 재기를 그만두고 n/a 를 찍는 것은 '깨끗하다' 가 아니다
    blind = {'file': 'erd_area_A.png', 'counts': {k: None for k in MEASURES},
             'tolerated': [], 'warn': []}
    eq(len(verify_faults(blind)), len(MEASURES), 'every unmeasured check is a fault')
    # 라벨을 안 그리는 개요도에서만 라벨 항목이 n/a 여도 된다 — 그 밖은 아니다
    ov = {'file': 'erd_overview.png',
          'counts': {'label_table': None, 'label_x': None, 'thru': 0,
                     'v_overlap': 0, 'h_overlap': 0}, 'tolerated': [], 'warn': []}
    eq(verify_faults(ov), [], 'the overview draws no labels — n/a is honest there')
    eq(verify_faults({**ov, 'file': 'erd_area_A.png'})[:1],
       ['label_table: n/a — a check that stopped measuring is not a clean check'],
       'the same n/a on a diagram that does draw labels is an escape hatch')
    # ④ 항목이 이름째 사라지면 남은 것이 전부 0 이라 조용히 통과한다
    gone = {'file': 'erd_area_A.png',
            'counts': {k: 0 for k in MEASURES if k != 'thru'}, 'tolerated': [], 'warn': []}
    if not verify_faults(gone):
        raise Fail('a check that vanished from the record must be a fault')
    extra = {'file': 'erd_area_A.png',
             'counts': {**{k: 0 for k in MEASURES}, 'brand_new': 0},
             'tolerated': [], 'warn': []}
    if not verify_faults(extra):
        raise Fail('a check the code added but this test ignores must be a fault')
    # ⑤ 봐주기는 케이스가 숫자로 적을 때만, 그 숫자까지만
    eq(verify_faults(honest, {'h_overlap': 2}), [], 'an explicit allowance covers it')
    if not verify_faults(honest, {'h_overlap': 1}):
        raise Fail('an allowance must not cover more than the number it names')


@case('verify: the sweep sees every diagram the case drew, not only the last run')
def _(work):
    # 14라운드. '케이스가 통과할 때마다 그 케이스가 남긴 기록을 전부 훑는다' 는
    # 13라운드의 주장이었는데, 훑기가 읽는 파일은 **판마다 처음부터 다시 쓰였다**
    # (erd.py 의 _LOG_STARTED 는 프로세스 전역이다). 그래서 build_erd.py 를 두 번
    # 이상 부르는 열 케이스에서 앞 판의 기록이 통째로 사라졌다 — 174장을 그리고
    # 114장(66%)만 봤다. 지워진 60장은 같은 그림을 다시 그린 것이 아니라 **서로 다른
    # 스키마**였다(ERD_MAX_AREAS 는 스키마 하나·둘·셋을, 오른쪽 끝 검사는 네 벌을
    # 각각 그린다). 그 60장에 회귀를 넣어 보면 시험은 초록이었다.
    #
    # 지금은 run() 이 부를 때마다 다른 기록 파일을 준다. 여기서 못박는 것은 그
    # 성질이다 — 앞 판에만 있던 그림이 훑기에 남아 있는가.
    t = {}
    for grp in ('order', 'user', 'item'):
        for i in range(3):
            t[f'{grp}_{grp}_{i}'] = table(f'{grp}_{grp}_{i}', [col('id')])
    write_schema(work, t)
    first = drawn_names(run('build_erd.py', work).stdout)
    # 두 번째 판은 영역이 하나뿐이라, 앞 판의 영역 그림 이름이 이번 판에는 없다
    write_schema(work, {'solo': table('solo', [col('id')])})
    second = drawn_names(run('build_erd.py', work).stdout)

    only_first = set(first) - set(second)
    if not only_first:
        raise Fail('the fixture no longer separates the two runs — both drew the same '
                   'file names, so nothing here can tell a kept record from a lost one')
    swept = [r['file'] for r in verify_recs(work, scope='all')]
    eq(len(swept), len(first) + len(second),
       'every diagram the case drew must leave a record the sweep reads')
    missing = sorted(only_first - set(swept))
    if missing:
        raise Fail(f'{", ".join(missing)} was drawn by the first run and is in no record '
                   f'the sweep reads — a regression in it would never turn this suite red')
    # 그리고 그 그림들은 **마지막 판만** 보면 정말로 안 보인다. 이 줄이 없으면 위
    # 단정은 '어차피 한 판밖에 없었다' 로도 통과한다.
    if only_first & {r['file'] for r in verify_recs(work)}:
        raise Fail('the last run redrew them after all — this case proves nothing')


@case('verify: the printed line and the machine record say the same thing')
def _(work):
    # 사람이 읽는 줄과 기계가 읽는 기록이 갈라지면 둘 중 하나는 거짓말이 된다.
    #
    # 예전엔 항목마다 `has(line, ...)` 로 조각을 찾고 `'[warn]' in line == bool(warn)`
    # 을 봤다. 그런데 붙박이 fixture 는 어느 그림도 경고를 내지 않아서, 그 마지막 줄은
    # 모든 그림에서 False == False 를 비교했다 — 경고 꼬리는 한 번도 재지 않았다.
    # (허용) 서식도 마찬가지로 한 번도 지나가지 않았다.
    #
    # 그래서 ① 세 갈래(숫자·n/a·허용·경고)가 다 나오는 입력을 쓰고 ② 조각이 아니라
    # **줄 전체를 기록에서 되만들어** 글자까지 맞춘다. 꼬리·서식·항목 순서가 다 걸린다.
    sys.path.insert(0, str(HERE))
    import lang.en as en
    write_schema(work, hub_schema(24))
    seen = {'a number': 0, 'n/a': 0, 'tolerated': 0, 'the warn tail': 0}

    def check(stdout):
        """방금 돈 판이 찍은 줄과 그 판이 남긴 기록을 짝지어 글자까지 맞춘다.

        기록은 판마다 처음부터 다시 쓰이므로(erd.verify_log), 한 판이 끝난 그 자리에서
        읽어야 그 판의 것만 본다. 줄 순서 = 기록 순서다.
        """
        lines = [x.strip() for x in stdout.split('\n') if 'verify ' in x]
        recs = verify_recs(work)
        eq(len(lines), len(recs), 'one machine record per printed verify line')
        for line, r in zip(lines, recs):
            parts = []
            for k in MEASURES:
                v = r['counts'][k]
                if v is None:
                    shown, kind = en.M['verify.na'], 'n/a'
                elif v and k in r['tolerated']:
                    shown, kind = en.M['verify.tolerated'].format(n=v), 'tolerated'
                else:
                    shown, kind = str(v), 'a number'
                seen[kind] += 1
                parts.append(f"{en.M['verify.' + k]} {shown}")
            want = en.M['log.verify'].format(name=r['file'], report=' · '.join(parts))
            if r['warn']:
                want += en.M['verify.warn'].format(
                    list=', '.join(en.M['verify.' + k] for k in r['warn']))
                seen['the warn tail'] += 1
            eq(line, want.strip(),
               f'{r["file"]}: the printed line is exactly what the record says')

    check(run('build_erd.py', work).stdout)

    # 같은 전체도를 '봐주기 없이' 한 번 더 그린다 — 경고 꼬리가 붙은 줄을 실제로 만든다.
    # tolerate 는 draw_erd 의 인자라, 코드를 고치지 않고 시험 쪽에서 끌 수 있다.
    probe = work / 'probe.py'
    probe.write_text(
        "import erd\n"
        "pos, boxes, groups = erd.layout_global()\n"
        "erd.draw_erd(erd.OUT / 'erd_full.png', list(erd.SCHEMA), pos, boxes, 'again',\n"
        "             with_desc=True, scale=2, legend=True, groups=groups,\n"
        "             derives=True, tolerate=())\n", encoding='utf-8')
    check(run(str(probe), work, env={'PYTHONPATH': str(HERE)}).stdout)

    # The router may legitimately eliminate the previously tolerated overlap.
    # Keep requiring numeric and n/a branches, but do not force a warning/tolerance
    # branch that a cleaner layout no longer produces.
    for kind in ('a number', 'n/a'):
        if not seen[kind]:
            raise Fail(f'no diagram exercised {kind} — the fixture proves nothing about it')


@case('verify: a diagram without labels does not report a label check it never ran')
def _(work):
    # 개요도는 관계 라벨을 아예 그리지 않는다(edge_labels=False). 그러면 라벨 겹침은
    # 잴 것이 없는데도 예전엔 0 이 찍혔다 — 검사하지 않은 것을 '깨끗하다' 로 읽히게
    # 하는 것은 이 저장소가 몇 판째 반복해 온 바로 그 잘못이다.
    write_schema(work, {
        'a': table('a', [col('id'), col('b_id')], pk=['id'],
                   fks=[{'column': 'b_id', 'ref_table': 'b', 'ref_column': 'id',
                         'on_delete': 'CASCADE'}]),
        'b': table('b', [col('id')], pk=['id'])})
    out = run('build_erd.py', work).stdout
    ov = verify_recs(work, 'overview')[0]
    for k in ('label_table', 'label_x'):
        if ov['counts'][k] is not None:
            raise Fail(f'the overview draws no labels — {k} cannot be a number')
    has([x for x in out.split('\n') if 'overview' in x and 'verify' in x][0],
        'n/a', 'the printed line says so too, instead of a reassuring 0')
    # 라벨을 그리는 그림에서는 반대로 반드시 숫자여야 한다 (n/a 로 도망가지 못하게)
    for r in verify_recs(work, 'area'):
        for k in ('label_table', 'label_x'):
            if r['counts'][k] is None:
                raise Fail(f'{r["file"]} draws labels — {k} must actually be measured')


@case('verify: a plain hub-and-spoke schema raises no warning on its first run')
def _(work):
    # 21테이블 두 허브 — 아무 데도 이상한 구석이 없는 입력이다. 그런데도 첫 판부터
    # `erd_area_B: 가로선 중첩 2 [경고]` 가 나왔다. [경고] 는 회귀 신호인데 평범한
    # 입력에서 울리면 사람은 그것을 무시하는 법부터 배운다 — 신호가 죽는다.
    t = {}
    for hub, n in (('shop', 10), ('user', 9)):
        t[f'{hub}_hub'] = table(f'{hub}_hub', [col('id'), col('name', 'text')], pk=['id'])
        for i in range(n):
            nm = f'{hub}_t{i:02d}'
            cols, fks = [col('id'), col('hub_id')], [
                {'column': 'hub_id', 'ref_table': f'{hub}_hub', 'ref_column': 'id',
                 'on_delete': 'CASCADE'}]
            if i >= 2:                      # 형제를 하나 건너 참조 — 통로가 붐빈다
                cols.append(col('prev_id'))
                fks.append({'column': 'prev_id', 'ref_table': f'{hub}_t{i - 2:02d}',
                            'ref_column': 'id', 'on_delete': 'SET NULL'})
            t[nm] = table(nm, cols, pk=['id'], fks=fks)
    write_schema(work, t)
    r = run('build_erd.py', work)
    verify_clean(work, what='an ordinary schema must not warn on its first run')
    if '[warn]' in r.stdout:
        raise Fail(f'a [warn] a user would learn to ignore:\n{r.stdout}')


@case('verify: the label overlap check measures ink, not the placement padding')
def _(work):
    # 14라운드. 위 케이스와 같은 부류인데 이번엔 `라벨↔라벨` 이었다. 자기참조를 절반쯤
    # 섞은 무작위 스키마 80개에서 셋이 [경고] 를 달았고, 잘라서 눈으로 본 셋 다 두 줄로
    # 나란히 놓인 멀쩡히 읽히는 라벨이었다. 재던 사각형이 글자가 아니라 **자리 잡기용
    # 여백**이었던 탓이다 — 좌우로 LABEL_PAD_X 씩, 세로는 글꼴과 무관하게 고정 높이라
    # 11pt 글자 잉크보다 넉넉하다. 어지러운 배치가 있어야 재현되는 자리라 그림으로는
    # 못 박아 두지 못한다. 그래서 규칙 자체를 여기에 박는다.
    #
    # 세 자리를 한꺼번에 지킨다: 여백이 걷혔는가, 여백만큼 떨어진 것을 안 세는가,
    # 그리고 **진짜 겹친 것은 여전히 세는가**. 마지막이 빠지면 이 수정은 그냥
    # 경고를 끈 것이 된다.
    write_schema(work, {'a': table('a', [col('id')])})
    e = {k: v for k, v in os.environ.items() if not k.startswith('ERD_')}
    e.update({'ERD_WORK': str(work), 'ERD_PROJ': str(work),
              'ERD_LANG': 'en', 'ERD_DOCNAME': 'T'})
    probe = '''
import json, erd
S = 2
f = erd.load_fonts(S)['edge']
P = erd.LABEL_PAD_X


def placed(cy, w=60.0):
    """자기참조 라벨이 쓰는 자리 잡기용 상자 — 중심이 글자의 anchor y 다."""
    return (100.0 - w / 2 - P, cy - 8, 100.0 + w / 2 + P, cy + 8)


def laps(c0, c1):
    a = erd.label_ink_box(placed(c0), f, S)
    b = erd.label_ink_box(placed(c1), f, S)
    return not (a[2] <= b[0] or a[0] >= b[2] or a[3] <= b[1] or a[1] >= b[3])


box = placed(100.0)
ink = erd.label_ink_box(box, f, S)
print(json.dumps({
    'pad': P,
    'inside': [ink[0] > box[0], ink[1] > box[1], ink[2] < box[2], ink[3] < box[3]],
    # 실제로 [경고] 를 냈던 간격들이다 (자리 잡기용 상자로는 3px·1px 겹쳤다)
    'apart_13': laps(100.0, 113.0),
    'apart_15': laps(100.0, 115.0),
    # 글자끼리 진짜로 포개지는 간격 — 여기서도 0 이면 이 검사는 죽은 것이다
    'apart_5': laps(100.0, 105.0),
    'apart_0': laps(100.0, 100.0),
}))
'''
    r = subprocess.run([sys.executable, '-c', probe], capture_output=True, text=True, encoding='utf-8',
                       env=e, cwd=str(HERE))
    if r.returncode != 0:
        raise Fail(f'the label geometry probe died:\n{r.stdout}\n{r.stderr}')
    got = json.loads(r.stdout.strip().splitlines()[-1])
    eq(got['inside'], [True] * 4,
       'the measured box must sit strictly inside the placement box on all four sides')
    eq(got['apart_13'], False, 'two labels 13px apart read as two lines — not an overlap')
    eq(got['apart_15'], False, 'nor 15px apart')
    eq(got['apart_5'], True, 'but 5px apart the glyphs really do collide — still counted')
    eq(got['apart_0'], True, 'and two labels on the same spot are the bug this check exists for')


# ── 세는 함수가 살아 있는가 ─────────────────────────────────────────────────
# 12라운드의 가장 큰 발견은 '회귀 시험이 반증 불가능했다' 였다. 13라운드는 '깨끗하다'
# 의 **규칙**을 시험이 갖게 해서 고쳤다고 적었는데, 14라운드에 뮤테이션 91개를 돌려
# 보니 시험이 갖게 된 것은 규칙뿐이고 **숫자**는 여전히 검사 대상인 erd.py 가 주고
# 있었다. 기록 188장의 census 가 그 이유를 그대로 보여 준다:
#
#   label_table {0:131, None:57}   label_x {0:131, None:57}
#   thru {0:188}   v_overlap {0:188}   h_overlap {0:185, 5:3}
#
# 다섯 중 넷은 **한 번도 0 이 아닌 적이 없다.** 규칙이 '실패' 방향으로 발동한 적이
# 없으니, 카운터를 상수 0 으로 바꾸거나(`('thru', thru_nodes())` → `('thru', 0)`)
# 관통 판정을 `> 4` 에서 `> 4000` 으로 벌려도 101개가 전부 통과했다. 12라운드가 잡은
# `warn` 이 이름만 `counts` 로 바뀐 셈이다. **그물은 있는데 걸리는 물고기가 없었다.**
#
# 그래서 아래 넷은 erd.py 의 산출물을 통하지 않는다. 세는 함수를 떼어내 **손으로 만든
# 겹치는 기하**를 직접 먹이고, 겹치면 반드시 0 이 아닌 값이 나오는 것을 못박는다.
# 재는 것은 레이아웃 품질이 아니라 **자가 살아 있는지**다 — 라우터가 좋아져 실제
# 겹침이 0 이 되어도 이 넷은 흔들리지 않는다. 마지막 하나는 반대쪽을 막는다: 자는
# 멀쩡한데 기록에 적히는 숫자가 그 자에서 나오지 않는 자리.
def _counting_ns(src=None):
    """`erd.py` 의 세는 함수들을 떼어내 단독으로 부를 수 있는 네임스페이스로.

    `overlaps`·`thru_nodes`·`lab_hits` 는 `draw_erd()` 안의 지역 함수라 밖에서 부를
    수 없다. 그래서 소스를 파싱해 그 def 만 꺼내 모듈 수준으로 컴파일한다 — 바깥에서
    빌려 쓰던 이름(`segs_v`·`node_rects`·`lab_boxes`)은 전역 조회가 되므로, 부르기
    직전에 이 네임스페이스에 넣어 주면 된다.

    erd 를 import 하지 않는 이유가 있다. import 하면 schema.json·글꼴·PIL 이 전부
    딸려 오고, 그중 아무것이나 어긋나면 **세는 함수와 상관없는 이유로** 이 케이스가
    빨강이 된다. 여기서 묻는 것은 오로지 '이 자가 아직 눈금을 가지고 있는가' 다.

    `src` 는 시험이 손댄 소스를 먹여 보는 자리다 (기본은 `erd.py` 를 읽는다).
    15라운드에 이 함수 자체를 '정확히 동작하는 손수 구현' 으로 갈아치우는 뮤턴트가
    통과했다 — 그러면 아래 단위 케이스 셋이 `erd.py` 가 아니라 **시험이 만든 스텁**을
    채점한다. 재는 손이 만든 것을 재는 손이 채점하는, 이 기록이 세 번 이름 붙인
    그 자리다. 그래서 인자를 하나 열어 두고, 소스를 한 줄 바꿔 먹였을 때 동작이
    정말 따라 바뀌는지를 케이스 하나가 못박는다.
    """
    import ast
    want = ('boxes_touch', 'overlaps', 'thru_nodes', 'lab_hits')
    tree = ast.parse(src if src is not None
                     else (HERE / 'erd.py').read_text(encoding='utf-8'))
    found = {}
    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef) and node.name in want:
            found.setdefault(node.name, node)
    missing = [n for n in want if n not in found]
    if missing:
        raise Fail(f'erd.py no longer defines {", ".join(missing)} — the counting '
                   f'functions were renamed, and a renamed counter is an unchecked one')
    mod = ast.Module(body=[found[n] for n in want], type_ignores=[])
    ast.fix_missing_locations(mod)
    ns = {'DEBUG_OVERLAP': False, 'T': lambda *a, **kw: ''}
    exec(compile(mod, str(HERE / 'erd.py'), 'exec'), ns)      # noqa: S102
    return ns


# ── 자에 눈금이 판 크기까지 있는가 ───────────────────────────────────────────
# 15라운드가 만든 자 넷은 전부 **손으로 만든 작은 기하**만 먹였다: node_rects 1개,
# lab_boxes 2~3개, segs 2개. 16라운드 검증자가 그 좁음을 뚫었다 — 세는 함수 첫 줄에
# 크기 조건 하나를 넣으면 **실제 사용자 스키마가 전부 무측정**인데 161개가 초록이었다.
#
#   lab_hits()   첫 줄에  if len(lab_boxes) > 3:   return 0     → all 161 passed
#   thru_nodes() 첫 줄에  if len(node_rects) > 12: return 0     → all 161 passed
#
# 작은 판에서만 이빨이 있는 자는 자가 없는 것과 같다. 이 저장소의 시험 안에만 해도
# 200테이블 판이 있고, 사용자 스키마는 그보다 크다.
#
# 그래서 같은 불변식을 **크기 사다리**에서 되풀이한다. 판 하나를 크게 만드는 것으로는
# 모자란다 — 문턱을 그 위로 올리면 그만이니까. 사다리는 2 에서 257 까지 배로 올라가고
# (실제 도판 규모를 넘는다), 칸마다 답을 **구성으로** 안다: 답을 계산해 주는 두 번째
# 구현을 여기 두면 재는 손이 제 답안을 제가 채점하는 그 자리로 되돌아간다.
SCALES = (2, 3, 5, 9, 17, 33, 65, 129, 257)
_ENDS_A, _ENDS_B = ((0.0, 0.0), (9.0, 9.0)), ((5.0, 5.0), (7.0, 7.0))
_E = ((0.0, 0.0), (1.0, 1.0))


def _labs_apart(n):
    """서로 100px 떨어진 라벨 상자 n 개 — 겹침 쌍은 구성상 0."""
    return [(i * 100.0, 0.0, i * 100.0 + 10.0, 10.0) for i in range(n)]


def _labs_paired(n):
    """둘씩 5px 포개진 라벨 상자 n 개 — 겹침 쌍은 구성상 n//2."""
    return [((i // 2) * 100.0 + (5.0 if i % 2 else 0.0), 0.0,
             (i // 2) * 100.0 + (5.0 if i % 2 else 0.0) + 10.0, 10.0) for i in range(n)]


def _labs_piled(n):
    """n 개가 서로 전부 포개진 더미 — 겹침 쌍은 구성상 n(n-1)/2."""
    return [(float(i), 0.0, float(i + n + 1), 10.0) for i in range(n)]


def _thru_board(n):
    """테이블 n 개짜리 판. 관통 4n 개 중 정확히 2n 건이 관통이다.

    표 i 는 x 로 200 씩 떨어져 있으므로 어느 선도 이웃 표까지 닿지 않는다 —
    답(2n)이 세는 함수의 규칙이 아니라 **판의 생김새**에서 나온다.
    """
    rects = [(i * 200.0, 0.0, i * 200.0 + 100.0, 60.0) for i in range(n)]
    segs_v = ([(i * 200.0 + 50.0, -20.0, 80.0, _E, False) for i in range(n)]      # 관통
              + [(i * 200.0 + 150.0, -20.0, 80.0, _E, False) for i in range(n)])  # 옆
    segs_h = ([(30.0, i * 200.0 - 20.0, i * 200.0 + 80.0, _E, False) for i in range(n)]
              + [(80.0, i * 200.0 - 20.0, i * 200.0 + 80.0, _E, False) for i in range(n)])
    return rects, segs_v, segs_h


def _ov_board(n):
    """구간 4n 개. 같은 자리를 40px 나란히 달리는 쌍이 n, 6px 만 스치는 쌍이 n."""
    segs = []
    for i in range(n):
        c = i * 1000.0
        segs.append((c, 0.0, 50.0, _ENDS_A, False))
        segs.append((c, 10.0, 60.0, _ENDS_B, False))       # 40px 나란히 → 1건
        segs.append((c + 500.0, 0.0, 50.0, _ENDS_A, False))
        segs.append((c + 500.0, 44.0, 90.0, _ENDS_B, False))   # 6px 스침 → 0건
    return segs


@case('verify: the counting namespace grades erd.py itself, not a stand-in')
def _(work):
    """세 단위 케이스가 채점하는 것이 정말 `erd.py` 인가.

    15라운드. `_counting_ns()` 를 **정확히 동작하는 손수 구현**으로 갈아치우니 아래
    셋이 전부 초록이었다 — erd.py 의 세는 함수를 통째로 지워도 마찬가지다. 12라운드가
    이름 붙인 '재는 손이 만든 것' 이 정확히 여기 다시 있었다: 도구를 재는 케이스가
    도구가 만든 것을 잰다. **`_counting_ns()` 자체를 재는 것이 하나도 없었다.**

    두 방향으로 막는다.

    1. 나온 함수가 `erd.py` 의 **그 줄**에서 나왔는가 — 파일 이름과 첫 줄 번호를,
       이 케이스가 따로 파싱한 AST 와 맞춘다. 손으로 쓴 스텁은 제 줄 번호를 단다.
    2. 정말 **읽은 것을 컴파일**하는가 — 소스를 한 줄 손봐 먹였을 때 동작이 따라
       바뀌어야 한다. 안 바뀌면 그 함수는 소스를 안 읽고 있다는 뜻이고, 그러면
       1번은 흉내 낼 수 있어도 이쪽은 못 흉내 낸다.
    """
    import ast
    want = ('boxes_touch', 'overlaps', 'thru_nodes', 'lab_hits')
    src = (HERE / 'erd.py').read_text(encoding='utf-8')
    defs = {}
    for n in ast.walk(ast.parse(src)):
        if isinstance(n, ast.FunctionDef) and n.name in want:
            defs.setdefault(n.name, n)
    eq(sorted(defs), sorted(want), 'erd.py defines all four counting functions')

    ns = _counting_ns()
    for name in want:
        fn = ns[name]
        if not callable(fn):
            raise Fail(f'{name} is not callable')
        eq(fn.__code__.co_filename, str(HERE / 'erd.py'),
           f'{name} must be compiled out of erd.py, not written here')
        eq(fn.__code__.co_firstlineno, defs[name].lineno,
           f'{name} must come from the line erd.py defines it on — a hand-written '
           f'stand-in passes every unit case below while erd.py counts nothing')

    # 소스를 손봐 먹인다: 함수의 첫 문장 앞에 표식 하나를 끼워 넣는다.
    # 들여쓰기는 그 첫 문장에서 가져오므로 중첩 깊이가 바뀌어도 따라간다.
    #
    # 15라운드 판은 **`boxes_touch` 하나만** 이렇게 확인했다. 16라운드 검증자가 그
    # 4분의 3 을 뚫었다: `boxes_touch` 만 소스에서 컴파일하고 `overlaps`·`thru_nodes`·
    # `lab_hits` 셋은 손으로 박은 뒤 `co_filename`·`co_firstlineno` 를 `code.replace()`
    # 로 위조하면 위의 두 대조가 전부 통과하고 `all 161 passed` 가 찍혔다. 줄 번호는
    # 위조할 수 있어도 **먹인 소스를 따라 바뀌는 것**은 위조할 수 없다 — 손으로 박은
    # 함수는 소스를 안 읽으므로 표식을 내놓지 못한다. 그래서 넷 다 그렇게 묻는다.
    for name in want:
        node = defs[name]
        first = node.body[0]
        lines = src.splitlines(keepends=True)
        mark = f'16R tampered {name}'
        lines.insert(first.lineno - 1,
                     ' ' * first.col_offset + f'return {mark!r}\n')
        tampered = _counting_ns(''.join(lines))
        blanks = [None] * len(node.args.args)
        eq(tampered[name](*blanks), mark,
           f'the namespace must compile the source it was handed — {name} ignored a '
           f'line planted at the top of its own body, so nothing here is reading '
           f'erd.py for it at all (a hand-written stand-in with a forged '
           f'co_filename passes every other check on this page)')
        try:                    # 손 안 댄 쪽은 표식을 모른다 (먹이가 없어 죽어도 좋다)
            got = ns[name](*blanks)
        except Exception:                                         # noqa: BLE001
            got = None
        if got == mark:
            raise Fail(f'the untouched namespace returned the doctored answer for '
                       f'{name} — the two namespaces are the same object')
    eq(ns['boxes_touch']((0, 0, 10, 10), (99, 99, 100, 100)), False,
       'and the untouched namespace is still the real ruler')


@case('verify: the overlap counter still counts two lines drawn on the same spot')
def _(work):
    ov = _counting_ns()['overlaps']
    A, B, C, D = (0.0, 0.0), (9.0, 9.0), (5.0, 5.0), (7.0, 7.0)

    def seg(coord, s0, s1, ends, pin=False):
        # (좌표, 구간 시작, 구간 끝, 그 선의 양 끝점, 컬럼 행에 못박힌 꼬리인가)
        return (coord, s0, s1, ends, pin)

    eq(ov([seg(100, 0, 50, (A, B)), seg(100, 10, 60, (C, D))]), 1,
       'two unrelated lines lying on the same coordinate for 40px is one overlap')
    # 11라운드의 버그: 끝점을 **둘 다** 공유하면 합류가 아니라 같은 선을 두 번 그린 것
    eq(ov([seg(100, 0, 50, (A, B)), seg(100, 10, 60, (A, B))]), 1,
       'the same line drawn twice is not a join — r11 waved it through')
    eq(ov([seg(100, 0, 50, (A, B)), seg(100, 10, 60, (C, B))]), 0,
       'lines that really do meet at one point may run together')
    eq(ov([seg(100, 0, 50, (A, B)), seg(103, 10, 60, (C, D))]), 0,
       '3px apart is not the same spot')
    eq(ov([seg(100, 0, 50, (A, B)), seg(102, 10, 60, (C, D))]), 1, 'but 2px apart is')
    eq(ov([seg(100, 0, 50, (A, B)), seg(100, 44, 90, (C, D))]), 0,
       'a 6px graze is not an overlap')
    eq(ov([seg(100, 0, 50, (A, B)), seg(100, 42, 90, (C, D))]), 1,
       'an 8px run together is')
    eq(ov([seg(100, 0, 50, (A, B), True), seg(100, 10, 60, (C, D), True)]), 0,
       'two pinned exit tails cannot be moved, so they are not counted')
    eq(ov([seg(100, 0, 50, (A, B), True), seg(100, 10, 60, (C, D), False)]), 1,
       'but a pinned tail under a lane the router chose is the router losing')

    # 그리고 도판 규모에서도. `if len(segs) > K: return 0` 한 줄이면 큰 판이 전부
    # 무측정인데 위 아홉 줄은 하나도 안 붉는다 — 자에 눈금이 어디까지 있는지는
    # 그 크기에서 재 봐야만 알 수 있다.
    for n in SCALES:
        eq(ov(_ov_board(n)), n,
           f'{4 * n} segments on a diagram-sized board must still yield {n} overlaps — '
           f'a counter that only bites on a two-segment fixture leaves every real '
           f'schema unmeasured')


@case('verify: the through-table counter still counts a line crossing a table')
def _(work):
    ns = _counting_ns()
    E = ((0.0, 0.0), (1.0, 1.0))
    box = [(0.0, 0.0, 100.0, 60.0)]

    def thru(v, h, rects=box):
        ns['segs_v'], ns['segs_h'], ns['node_rects'] = v, h, rects
        return ns['thru_nodes']()

    eq(thru([(50.0, -20.0, 80.0, E, False)], []), 1,
       'a vertical line straight through a table is a through-table hit')
    eq(thru([], [(30.0, -20.0, 200.0, E, False)]), 1,
       'and so is a horizontal one — half a check is not a check')
    eq(thru([(50.0, -20.0, 80.0, E, False)], [(30.0, -20.0, 200.0, E, False)]), 2,
       'both halves count into the same number')
    eq(thru([(50.0, 55.0, 100.0, E, False)], []), 1,
       '5px of penetration is still a line inside a table')
    eq(thru([(50.0, 57.0, 100.0, E, False)], []), 0,
       'but 3px is the border, not a crossing')
    eq(thru([(150.0, -20.0, 80.0, E, False)], []), 0,
       'a line beside the table is not through it')

    # 그리고 도판 규모에서도. 검증자가 `thru_nodes()` 첫 줄에 `if len(node_rects) > 12:
    # return 0` 을 넣었더니 161개가 전부 초록이었다 — 위 여섯 줄이 표 **한 개**짜리
    # 판만 먹이기 때문이다. 사다리는 실제 도판(이 시험 안에 200테이블 판이 있다)을
    # 넘는 곳까지 올라간다.
    for n in SCALES:
        rects, sv, sh = _thru_board(n)
        eq(thru(sv, sh, rects), 2 * n,
           f'a board of {n} tables and {len(sv) + len(sh)} segments must still report '
           f'{2 * n} crossings — a size-conditional counter reads exactly like a clean '
           f'diagram on every schema a user actually has')


@case('verify: the two label counters still count boxes that really overlap')
def _(work):
    ns = _counting_ns()
    bt = ns['boxes_touch']
    # 라벨↔테이블(lab_hit)과 라벨↔라벨(lab_hits)이 **같은 자**를 쓴다. 예전엔 한쪽이
    # `<` 라 스치기만 해도 세고 다른 쪽은 `<=` 라 안 셌다 — 같은 것을 재면서 반대
    # 규칙이었다. 그래서 여기서 자 하나를 세 방향으로 지킨다.
    eq(bt((0, 0, 10, 10), (5, 5, 15, 15)), True, 'boxes that share ink do overlap')
    eq(bt((0, 0, 10, 10), (10, 0, 20, 10)), False, 'boxes that only touch do not')
    eq(bt((0, 0, 10, 10), (11, 0, 20, 10)), False, 'and boxes apart do not')
    # 원점 근처에서만 정직한 자도 자가 아니다 — 200테이블 판의 좌표는 만 단위다.
    for d in (5000, 50000):
        eq(bt((d, d, d + 10, d + 10), (d + 5, d + 5, d + 15, d + 15)), True,
           f'and the same ruler still reads True {d}px from the origin')
        eq(bt((d, d, d + 10, d + 10), (d + 10, d, d + 20, d + 10)), False,
           f'and still reads False {d}px from the origin')

    def hits(boxes):
        ns['lab_boxes'] = boxes
        return ns['lab_hits']()

    eq(hits([(0, 0, 10, 10), (5, 5, 15, 15), (100, 100, 110, 110)]), 1,
       'one overlapping pair out of three labels')
    eq(hits([(0, 0, 10, 10), (5, 5, 15, 15), (6, 6, 16, 16)]), 3,
       'three labels piled on each other are three pairs')
    eq(hits([(0, 0, 10, 10), (10, 0, 20, 10)]), 0, 'labels that only touch are readable')

    # 그리고 도판 규모에서도. 검증자가 `lab_hits()` 첫 줄에 `if len(lab_boxes) > 3:
    # return 0` 을 넣었더니 161개가 전부 초록이었다 — 위 세 줄이 라벨 두세 개짜리
    # 판만 먹이기 때문이다. 같은 한 줄로 라벨이 실제로 서로 포개진 판까지 전부
    # '깨끗함' 으로 보고된다.
    for n in SCALES:
        eq(hits(_labs_apart(n)), 0,
           f'{n} labels laid out 100px apart are {n} readable labels')
        eq(hits(_labs_paired(n)), n // 2,
           f'{n} labels in overlapping couples are {n // 2} clashes — a counter that '
           f'only bites below four labels leaves every real diagram unmeasured')
        eq(hits(_labs_piled(n)), n * (n - 1) // 2,
           f'{n} labels piled on one another are {n * (n - 1) // 2} pairs')


@case('verify: the counters the drawing path really uses are fed, and still bite')
def _(work):
    """단위 케이스 셋의 **반대편** — 세는 함수가 실제 도판에서 무엇을 먹는가.

    15라운드. 위 셋은 세는 함수에 먹이를 **자기가** 넣는다. 그래서 원리상 못 잡는
    뮤턴트가 두 부류 있었고, 둘 다 `all 141 passed` 로 빠져나갔다.

      · 굶기기 — `lab_boxes = [… for b in placed[:0]]` 한 줄이면 라벨 두 항목이
        영구히 0 이다. 자는 멀쩡하고, 먹일 것이 없을 뿐이다
      · 실제 경로에서만 거짓 — `lab_hits()` 안에서 단독 네임스페이스에는 없는
        이름(`pos`)으로 갈래를 파면 단위 시험은 통과하고 도판만 0 이 된다

    그래서 이 케이스는 **진짜 판을 돌리면서** draw_erd 프레임을 붙잡는다. 돌아올 때
    그 프레임의 지역값에서 먹이의 크기를 읽고(굶기기), 그 판이 실제로 쓴 함수 객체를
    그대로 꺼내 손으로 만든 겹치는 기하를 자유변수 셀에 꽂아 불러 본다(이빨).
    단위 케이스가 `erd.py` 소스를 떼어 낸 사본을 재는 것과 달리, 여기서 재는 것은
    그 판이 실제로 부른 바로 그 함수다.
    """
    write_schema(work, {
        'a': table('a', [col('id'), col('b_id')], pk=['id'],
                   fks=[{'column': 'b_id', 'ref_table': 'b', 'ref_column': 'id',
                         'on_delete': 'CASCADE'}]),
        'b': table('b', [col('id')], pk=['id'])})
    probe = work / 'probe.py'
    probe.write_text('''\
import json, os, sys
import build_erd

GRAB = []


def tracer(frame, event, arg):
    if event == 'call' and frame.f_code.co_name == 'draw_erd':
        frame.f_trace_lines = False        # 줄 이벤트까지 받으면 판이 기어간다

        def at_return(f, ev, a):
            if ev == 'return' and 'lab_hits' in f.f_locals:
                GRAB.append(f.f_locals)
            return at_return
        return at_return
    return tracer


sys.settrace(tracer)
try:
    build_erd.main()
finally:
    sys.settrace(None)


def cell(fn, name, value):
    fn.__closure__[fn.__code__.co_freevars.index(name)].cell_contents = value


feeds = [{'file': os.path.basename(str(g['path'])),
          'edge_labels': bool(g['edge_labels']),
          'tnames': len(g['tnames']), 'placed': len(g['placed']),
          'lab_boxes': len(g['lab_boxes']), 'node_rects': len(g['node_rects']),
          'segs': len(g['segs_v']) + len(g['segs_h'])} for g in GRAB]

teeth = {}
labelled = [x for x in GRAB if x['edge_labels']]
if labelled:            # 없으면 먹이 쪽 판정이 먼저 그 사실을 말한다
    g = labelled[-1]
    A, B, C, D = (0.0, 0.0), (9.0, 9.0), (5.0, 5.0), (7.0, 7.0)
    E = ((0.0, 0.0), (1.0, 1.0))
    ov, th, lh = g['overlaps'], g['thru_nodes'], g['lab_hits']
    teeth['same_spot'] = ov([(100, 0, 50, (A, B), False),
                             (100, 10, 60, (C, D), False)])
    teeth['twice'] = ov([(100, 0, 50, (A, B), False), (100, 10, 60, (A, B), False)])
    teeth['join'] = ov([(100, 0, 50, (A, B), False), (100, 10, 60, (C, B), False)])
    cell(th, 'node_rects', [(0.0, 0.0, 100.0, 60.0)])
    cell(th, 'segs_h', [])
    cell(th, 'segs_v', [(50.0, -20.0, 80.0, E, False)])
    teeth['thru'] = th()
    cell(th, 'segs_v', [(150.0, -20.0, 80.0, E, False)])
    teeth['beside'] = th()
    cell(lh, 'lab_boxes', [(0, 0, 10, 10), (5, 5, 15, 15)])
    teeth['lab_pair'] = lh()
    cell(lh, 'lab_boxes', [(0, 0, 10, 10), (10, 0, 20, 10)])
    teeth['lab_touch'] = lh()
    # 그리고 도판 규모에서도 (16라운드). 위 일곱은 상자 두 개·표 한 개짜리 판이라,
    # 세는 함수 첫 줄의 크기 조건 하나로 전부 지나간다.
    N = 64
    cell(lh, 'lab_boxes', [(i, 0, i + N + 1, 10) for i in range(N)])
    teeth['lab_many'] = lh()                       # 서로 전부 포갠 더미 → N(N-1)/2
    cell(th, 'node_rects', [(i * 200.0, 0.0, i * 200.0 + 100.0, 60.0)
                            for i in range(N)])
    cell(th, 'segs_v', [(i * 200.0 + 50.0, -20.0, 80.0, E, False) for i in range(N)]
         + [(i * 200.0 + 150.0, -20.0, 80.0, E, False) for i in range(N)])
    cell(th, 'segs_h', [(30.0, i * 200.0 - 20.0, i * 200.0 + 80.0, E, False)
                        for i in range(N)])
    teeth['thru_many'] = th()                      # 관통 N(세로) + N(가로) = 2N
    big = []
    for i in range(N):
        c = i * 1000.0
        big += [(c, 0.0, 50.0, (A, B), False), (c, 10.0, 60.0, (C, D), False),
                (c + 500.0, 0.0, 50.0, (A, B), False),
                (c + 500.0, 44.0, 90.0, (C, D), False)]
    teeth['ov_many'] = ov(big)                     # 나란히 달리는 쌍만 → N
print('15R-PROBE ' + json.dumps({'feeds': feeds, 'teeth': teeth}))
''', encoding='utf-8')
    out = run(str(probe), work, env={'PYTHONPATH': str(HERE)}).stdout
    said = [ln for ln in out.splitlines() if ln.startswith('15R-PROBE ')]
    if not said:
        raise Fail(f'the probe never reached draw_erd:\n{out}')
    got = json.loads(said[-1][len('15R-PROBE '):])

    # ① 먹이 — 세는 함수가 실제로 무엇을 받는가
    feeds = got['feeds']
    if not feeds:
        raise Fail('the build drew no diagram this probe could watch')
    labelled = [f for f in feeds if f['edge_labels']]
    if not labelled:
        raise Fail('no diagram drew relationship labels — then the two label counters '
                   'measured nothing and this case would prove nothing')
    for f in feeds:
        if f['node_rects'] != f['tnames'] or not f['tnames']:
            raise Fail(f'{f["file"]}: the table-rectangle feed is {f["node_rects"]} for '
                       f'{f["tnames"]} tables — a starved counter reads exactly like a '
                       f'clean diagram')
        if not f['segs']:
            raise Fail(f'{f["file"]}: not one line segment reached the overlap counters')
    for f in labelled:
        if not f['placed']:
            raise Fail(f'{f["file"]}: the diagram says it draws labels but placed none')
        if f['lab_boxes'] != f['placed']:
            raise Fail(f'{f["file"]}: {f["placed"]} labels drawn but {f["lab_boxes"]} '
                       f'handed to the label counters — starving the ruler is how '
                       f'label_table and label_x stay 0 for ever')

    # ② 이빨 — 그 판이 실제로 부른 그 함수 객체가 겹침에 반응하는가
    eq(got['teeth'], {'same_spot': 1, 'twice': 1, 'join': 0, 'thru': 1, 'beside': 0,
                      'lab_pair': 1, 'lab_touch': 0,
                      'lab_many': 64 * 63 // 2, 'thru_many': 128, 'ov_many': 64},
       'the counting functions the real drawing path used must still count — a counter '
       'that only tells the truth to the unit cases (or only on a two-box fixture) is '
       'worse than none')


@case('verify: the five numbers in the record come from the counting functions')
def _(work):
    # 위 셋은 '자가 살아 있는가' 를 묻는다. 이 케이스는 반대쪽이다 — 자는 멀쩡한데
    # 기록에 적히는 숫자가 그 자에서 안 나오는 자리. 실제로 `('thru', thru_nodes())`
    # 를 `('thru', 0)` 으로 바꾸면 넷 다 초록이었다 (188장 어디서도 thru 가 0 이 아닌
    # 적이 없어 사람도 눈치채지 못한다). 그래서 여기서는 값이 아니라 **배선**을 본다.
    #
    # 15라운드. 그런데 그 배선을 **이름 집합**으로 봤다: 식 안에 `lab_hits` 라는
    # 이름이 나오기만 하면 통과였다. 그래서 `lab_hit = 0 * sum(…)` 이 그대로
    # 빠져나갔다 — 이름은 전부 제자리에 있고 값만 상수 0 이다. `x and False`,
    # `x if False else 0`, `lab_boxes[:0]` 도 모두 같은 구멍으로 지나간다.
    # 이름이 아니라 **모양**을 못박는다: 그 자리에 적힐 수 있는 식은 그 호출 하나뿐이고,
    # 껍데기는 `… if edge_labels else None` 하나만 허용한다. 모양을 정해 두면 상수 곱은
    # 이름이 아무리 멀쩡해도 `BinOp` 라서 걸린다. 레이아웃을 정말 고쳐 쓸 일이 생기면
    # 이 케이스가 붉어지는데, 그때가 바로 '기록의 숫자가 어디서 나오는지' 를 다시
    # 적어야 하는 때다.
    import ast
    tree = ast.parse((HERE / 'erd.py').read_text(encoding='utf-8'))
    draw = next((n for n in ast.walk(tree)
                 if isinstance(n, ast.FunctionDef) and n.name == 'draw_erd'), None)
    if draw is None:
        raise Fail('erd.py no longer has draw_erd()')

    def show(node):
        try:
            return ast.unparse(node)
        except Exception:                                         # noqa: BLE001
            return f'<{type(node).__name__}>'

    def _here(node):
        """draw_erd 안의 문장 전부 — 다만 **중첩 함수 안은 안 본다** (같은 이름과 안 섞이게).

        15라운드 판은 `draw.body` 의 직계 자식만 훑었다. 그러면 `if True:` 한 겹만
        씌워도 그 대입이 안 보인다.
        """
        for ch in ast.iter_child_nodes(node):
            if isinstance(ch, (ast.FunctionDef, ast.AsyncFunctionDef, ast.Lambda)):
                continue
            yield ch
            yield from _here(ch)

    def top_assign(name):
        """draw_erd 안에서 그 이름에 값을 묶는 **모든** 자리.

        15라운드 판은 **처음 만난** 대입 하나만 돌려줬다. 그러면 같은 이름에 두 번째
        대입을 하는 것만으로 아래 모양 검사가 통째로 지나간다 — 첫 줄은 규칙대로 두고
        다음 줄에 `checks = [(k, 0) for k, _ in checks]` 를 적으면 된다. 실제로 그렇게
        해서 161개가 초록이었다. 그래서 묶는 자리를 전부 모으고, **전부가** 규칙을
        지켜야 통과다.

        대입이 아닌 묶기(증분 대입·`for`·`with … as`·`except … as`·바다코끼리)는 값의
        모양을 읽을 수 없으므로 그 자리에서 죽는다 — 읽을 수 없는 것을 '괜찮다' 로
        넘기는 것이 이 케이스가 세 번 물린 모양이다.
        """
        out = []
        for st in _here(draw):
            if isinstance(st, ast.Assign):
                for t in st.targets:
                    if getattr(t, 'id', '') == name:
                        out.append(st.value)
                    elif any(getattr(x, 'id', '') == name for x in ast.walk(t)
                             if isinstance(x, ast.Name)):
                        raise Fail(f'draw_erd() binds `{name}` inside the unpacking '
                                   f'`{show(t)}` — this check can only read a plain '
                                   f'`{name} = <expression>`')
                continue
            if isinstance(st, (ast.AugAssign, ast.AnnAssign, ast.NamedExpr)):
                if getattr(st.target, 'id', '') == name:
                    raise Fail(f'draw_erd() rebinds `{name}` with `{show(st)}` — a '
                               f'second binding decides the number, and a shape check '
                               f'that only reads the first one is checking nothing')
                continue
            targets = []
            if isinstance(st, (ast.For, ast.AsyncFor)):
                targets = [st.target]
            elif isinstance(st, (ast.With, ast.AsyncWith)):
                targets = [i.optional_vars for i in st.items if i.optional_vars]
            elif isinstance(st, ast.ExceptHandler) and st.name == name:
                raise Fail(f'draw_erd() rebinds `{name}` as an except target')
            for t in targets:
                if any(getattr(x, 'id', '') == name for x in ast.walk(t)
                       if isinstance(x, ast.Name)):
                    raise Fail(f'draw_erd() rebinds `{name}` as a loop or with target '
                               f'— this check can only read a plain assignment')
        # 묶기만 막으면 반쪽이다. `lab_boxes.clear()` 는 대입이 아니지만 먹이를 똑같이
        # 굶긴다 — 이름은 그대로, 값만 비어 버린다. 이 이름들에 붙는 메서드 호출은
        # 하나도 없어야 한다(오늘 erd.py 에도 없다).
        for st in _here(draw):
            if (isinstance(st, ast.Call) and isinstance(st.func, ast.Attribute)
                    and getattr(st.func.value, 'id', '') == name):
                raise Fail(f'draw_erd() calls `{show(st)}` — a method call on `{name}` '
                           f'can empty the feed without touching the assignment this '
                           f'check reads')
        return out

    def unguard(key, node):
        if not isinstance(node, ast.IfExp):
            return node
        if not (isinstance(node.test, ast.Name)
                and isinstance(node.orelse, ast.Constant) and node.orelse.value is None):
            raise Fail(f'{key} is written as {show(node)} — the only shape allowed here '
                       f'is `<count> if <flag> else None`; anything else can turn the '
                       f'number off without turning the check red')
        return node.body

    def is_call(node, fname, args=()):
        return (isinstance(node, ast.Call) and isinstance(node.func, ast.Name)
                and node.func.id == fname and not node.keywords
                and len(node.args) == len(args)
                and all(isinstance(a, ast.Name) and a.id == w
                        for a, w in zip(node.args, args)))

    def one(name, values):
        """묶는 자리는 하나여야 하고 없어도 안 된다 — 그 하나를 돌려준다."""
        if not values:
            raise Fail(f'draw_erd() no longer assigns `{name}` anywhere this check can '
                       f'read')
        if len(values) > 1:
            raise Fail(f'draw_erd() assigns `{name}` {len(values)} times '
                       f'({" / ".join(show(v) for v in values)}) — the last one decides '
                       f'the number, so every one of them must hold the shape; write it '
                       f'once')
        return values[0]

    checks = one('checks', top_assign('checks'))
    if not isinstance(checks, ast.List):
        raise Fail('draw_erd() no longer builds a list called `checks` — the record is '
                   'assembled somewhere this test cannot read, so nobody is checking '
                   'that its numbers come from the counters')
    for el in checks.elts:
        if not (isinstance(el, ast.Tuple) and len(el.elts) == 2
                and isinstance(el.elts[0], ast.Constant)):
            raise Fail(f'`checks` holds `{show(el)}` — every entry must be a '
                       f'(name, count) pair written right there, or this test cannot '
                       f'tell where the number came from')
    wiring = [(el.elts[0].value, el.elts[1]) for el in checks.elts]
    eq([k for k, _ in wiring], list(MEASURES),
       'the record carries exactly the five measures the test knows, in order')
    shape = {'label_table': ('lab_hit', lambda v: isinstance(v, ast.Name)
                             and v.id == 'lab_hit'),
             'label_x': ('lab_hits()', lambda v: is_call(v, 'lab_hits')),
             'thru': ('thru_nodes()', lambda v: is_call(v, 'thru_nodes')),
             'v_overlap': ('overlaps(segs_v)',
                           lambda v: is_call(v, 'overlaps', ('segs_v',))),
             'h_overlap': ('overlaps(segs_h)',
                           lambda v: is_call(v, 'overlaps', ('segs_h',)))}
    for k, node in wiring:
        want, ok = shape[k]
        v = unguard(k, node)
        if not ok(v):
            raise Fail(f'the record writes {k} as `{show(v)}` — it must be exactly '
                       f'`{want}`\n      a counter wired to a constant (or multiplied '
                       f'by one, or `and False`) reads exactly like a clean diagram')

    # `lab_hit` 은 이름이라 한 겹 더 따라간다. 그리고 그 셈에 먹이를 대는 두 목록도
    # 함께 본다 — 자를 그대로 두고 **먹이만 굶기는** 뮤턴트가 따로 있었다.
    lab_hit = one('lab_hit', top_assign('lab_hit'))
    if not (isinstance(lab_hit, ast.Call) and isinstance(lab_hit.func, ast.Name)
            and lab_hit.func.id == 'sum' and len(lab_hit.args) == 1
            and isinstance(lab_hit.args[0], ast.GeneratorExp)):
        raise Fail(f'lab_hit is `{show(lab_hit)}` — it must be exactly '
                   f'`sum(<generator over lab_boxes × node_rects>)`')
    gen = lab_hit.args[0]
    if not (isinstance(gen.elt, ast.Constant) and gen.elt.value == 1):
        raise Fail(f'lab_hit sums `{show(gen.elt)}` instead of 1 per hit')
    eq(sorted(getattr(c.iter, 'id', show(c.iter)) for c in gen.generators),
       ['lab_boxes', 'node_rects'],
       'lab_hit must walk every label against every table — a sliced or filtered feed '
       'is how the count stays 0 with the ruler untouched')
    conds = [t for c in gen.generators for t in c.ifs]
    if not (len(conds) == 1 and is_call(conds[0], 'boxes_touch', ('b', 'o'))):
        raise Fail(f'lab_hit decides with `{"; ".join(show(c) for c in conds) or "nothing"}`'
                   f' — the label/table check must use the same ruler as label/label')
    for name, src in (('lab_boxes', 'placed'), ('node_rects', 'tnames')):
        v = one(name, top_assign(name))
        if not (isinstance(v, ast.ListComp) and len(v.generators) == 1
                and isinstance(v.generators[0].iter, ast.Name)
                and v.generators[0].iter.id == src and not v.generators[0].ifs):
            raise Fail(f'{name} is built as `{show(v)}` — it must be one comprehension '
                       f'over the whole of `{src}`, with nothing sliced off the end; '
                       f'starving a counter is indistinguishable from a clean diagram')


@case('verify: relationships that cross the corridors do not stack on an exit row')
def _(work):
    # 위 허브-앤-스포크는 진출입 꼬리끼리 스치는 쪽이었다. 이쪽은 다른 갈래다:
    # 열을 건너뛰는 선이 통로에서 고른 lane 이, **나중에** 정해진 다른 관계의 진출입
    # y 위에 그대로 얹혔다. 진출입 y 를 다 잡은 뒤에야 lane 을 고르게 바꿔 고쳤다.
    # 관계가 고르게 퍼진 스키마에서는 재현되지 않아, 어긋난 참조를 그대로 박아 둔다.
    link = {'ord_t00': ['inv_t03'], 'ord_t02': ['inv_t02', 'inv_t00'],
            'ord_t03': ['ord_t04'], 'ord_t04': ['ord_t03', 'usr_t04'],
            'usr_t01': ['inv_t00', 'ord_t02'], 'usr_t02': ['inv_t00'],
            'usr_t03': ['inv_t02'], 'usr_t04': ['inv_t03', 'usr_t03'],
            'inv_t00': ['usr_t03'], 'inv_t01': ['usr_t03'], 'inv_t02': ['inv_t04']}
    t = {}
    for p in ('ord', 'usr', 'inv'):
        for i in range(5):
            nm = f'{p}_t{i:02d}'
            cols, fks = [col('id'), col('v', 'text')], []
            for j, ref in enumerate(link.get(nm, [])):
                cols.append(col(f'r{j}'))
                fks.append({'column': f'r{j}', 'ref_table': ref, 'ref_column': 'id',
                            'on_delete': 'CASCADE'})
            t[nm] = table(nm, cols, pk=['id'], fks=fks)
    write_schema(work, t)
    r = run('build_erd.py', work)
    verify_clean(work, what='a corridor lane must not land on a column exit row')
    if '[warn]' in r.stdout:
        raise Fail(f'a [warn] a user would learn to ignore:\n{r.stdout}')


@case('selftest: the suite does not inherit the caller ERD_* environment')
def _(work):
    # `ERD_LABEL=shop python3 selftest.py` 는 2개, `ERD_EXCLUDE=.* ` 는 19개를
    # 깨뜨렸다. 문서가 권하는 다중 DB 흐름을 그대로 따른 사람이 `install.sh --check`
    # 에서 '설치가 고장 났다' 는 말을 들었다. 시험은 코드를 재야지 껍데기를 재면 안 된다.
    write_schema(work, {'t': table('t', [col('id'), col('v', 'text')])})
    clean = run('build_erd.py', work).stdout
    poison = {'ERD_EXCLUDE': '.*', 'ERD_LABEL': 'shop', 'ERD_LANG': 'ko',
              'ERD_DOCNAME': 'leaked', 'ERD_SVG_TITLE': '1'}
    keep = {k: os.environ.get(k) for k in poison}
    os.environ.update(poison)
    try:
        eq(run('build_erd.py', work).stdout, clean,
           'the caller environment must not reach the scripts under test')
    finally:
        for k, v in keep.items():
            os.environ.pop(k, None) if v is None else os.environ.__setitem__(k, v)


@case('selftest: the harness fails where it says it fails')
def _(work):
    # **재는 쪽은 아무도 재 주지 않는다** — 이 기록이 세 라운드째 같은 이름을 붙이고
    # 있는 자리다. 14라운드의 뮤테이션이 그 목록을 늘렸다: `eq()` 가 절대 실패하지
    # 않게 해도, `case()` 가 이름 중복을 그냥 넘겨도, `verify_recs()` 의 '기록이
    # 없으면 실패' 를 `return []` 로 바꿔도, `verify_clean()` 이 판정을 그만둬도,
    # 훑기의 '그린 장수 ↔ 기록 장수' 대조를 지워도 101개가 전부 통과했다.
    # 도구가 제 실패를 말하지 못하면 그 도구가 낸 초록은 아무것도 뜻하지 않는다.
    import selftest_kit as kit

    def must_fail(fn, what):
        try:
            fn()
        except (kit.Fail, AssertionError, RuntimeError):
            return
        raise Fail(f'{what} went through without a word')

    # 15라운드. 프로브가 `eq(1, 2, 'x')` 와 `has('abcdef', 'zzz', 'x')` 딱 둘이라,
    # **그 두 값에서만** 정직한 뮤턴트가 그대로 빠져나갔다: `eq()` 가 list·tuple·dict·
    # set 에는 절대 실패하지 않게 해도 초록이었고(이 저장소 단정의 대부분이 목록·사전
    # 비교다), `has()` 가 네 글자 이상 바늘에 절대 실패하지 않게 해도 초록이었다
    # (긴 문자열 검색이 대부분이다). 도구가 어떤 값에 정직한지를 재려면 프로브가 실제로
    # 쓰이는 값의 모양을 덮어야 한다 — 재는 쪽의 fixture 가 좁으면 재는 쪽도 좁게
    # 정직하다. 반대 방향(같은 값에는 지나가야 한다)도 함께 둔다.
    must_fail(lambda: eq(1, 2, 'x'), 'eq() on two different values')
    must_fail(lambda: eq('abc', 'abd', 'x'), 'eq() on two different strings')
    must_fail(lambda: eq([1, 2], [1, 3], 'x'), 'eq() on two different lists')
    must_fail(lambda: eq((1, 2), (1, 3), 'x'), 'eq() on two different tuples')
    must_fail(lambda: eq({'a': 1}, {'a': 2}, 'x'), 'eq() on two different dicts')
    must_fail(lambda: eq({1, 2}, {1, 3}, 'x'), 'eq() on two different sets')
    must_fail(lambda: eq(None, 0, 'x'), 'eq() on a missing value against a zero')
    # 16라운드. 프로브의 모양을 짐작으로 고르지 않는다 — 한 벌을 계측해서 골랐다.
    # `eq` 971회: list 782 · int 78 · tuple 36 · str 29 · dict 23 · **bool 13** ·
    # set 8 · None 2. 위 일곱 줄에 **bool 이 없었다**. `eq(True, False)` 는 이 저장소가
    # 열세 번 묻는 모양이고(`'[warn]' in out` 류), 불리언에만 정직한 뮤턴트는 그
    # 열세 자리를 통째로 재운다. 길이 조건으로 갈래를 파는 거짓말도 함께 막는다 —
    # 그 부류가 이번 라운드에 세는 함수 쪽에서 실제로 통했다.
    must_fail(lambda: eq(True, False, 'x'), 'eq() on two different booleans')
    must_fail(lambda: eq(False, None, 'x'), 'eq() on a false against a missing value')
    must_fail(lambda: eq([('a', 1)], [('a', 2)], 'x'),
              'eq() on two lists of pairs — the shape most of this suite compares')
    must_fail(lambda: eq(list(range(50)), list(range(49)) + [99], 'x'),
              'eq() on two fifty-item lists that differ in the last one')
    must_fail(lambda: eq('x' * 4000 + 'a', 'x' * 4000 + 'b', 'x'),
              'eq() on two long strings that differ at the end')
    for same in (1, 'abc', [1, 2], (1, 2), {'a': 1}, {1, 2}, None, True, False,
                 list(range(50))):
        eq(same, same, 'eq() must let an equal value through')
    must_fail(lambda: has('abcdef', 'zzz', 'x'), 'has() on a needle that is not there')
    must_fail(lambda: has('abcdef', 'abcdefg', 'x'), 'has() on a seven-letter needle')
    must_fail(lambda: has('<html>' + 'x' * 5000, 'ERD_VERIFY_LOG', 'x'),
              'has() on a long needle missing from a long haystack')
    must_fail(lambda: has([], 'anything', 'x'), 'has() on an empty haystack')
    has('abcdef', 'cde', 'has() must let a needle that is there through')
    # 그리고 공백 든 바늘. 계측하면 `has` 150회 중 **42회(28%)** 가 공백을 담고 있다 —
    # 사용자에게 보이는 **문장**을 찾는 자리가 전부 그렇다. 16라운드 검증자가
    # `has()` 를 '공백 든 바늘에는 절대 실패 안 함' 으로 바꿨더니 161개가 초록이었고,
    # 위 다섯 줄은 한 글자도 공백을 안 담고 있어서 그럴 수밖에 없었다.
    must_fail(lambda: has('abcdef', 'ab cd', 'x'), 'has() on a needle with a space')
    must_fail(lambda: has('a b c d', 'b  c', 'x'),
              'has() on a needle whose spacing is wrong')
    must_fail(lambda: has('6 cases need a real server and were NOT run',
                          '7 cases need a real server', 'x'),
              'has() on a sentence needle that is off by one number')
    must_fail(lambda: has('<svg>\n<g id="a"/>\n</svg>', 'stroke-width: 2\n', 'x'),
              'has() on a multi-line needle missing from a document')
    must_fail(lambda: has('  ' * 400 + 'tail', 'ERD_MAX_AREAS says so\n', 'x'),
              'has() on a long needle with a space and a newline')
    has('a b c d', 'b c', 'has() must let a needle with a space through')
    has('<svg>\n<g id="a"/>\n</svg>', '<g id="a"/>\n</svg>',
        'and one that spans two lines')

    saved_cases = list(kit.CASES)
    try:
        kit.case('__probe__ duplicate')(lambda w: None)
        must_fail(lambda: kit.case('__probe__ duplicate')(lambda w: None),
                  'a second case registered under a name already taken')
    finally:
        kit.CASES[:] = saved_cases

    # 등록이 아무것도 못 찾는 것은 '고칠 게 없다' 가 아니라 실패다. 글로브가 어긋나면
    # 39개가 사라지는데 종료코드 0 에 마지막 줄은 초록 `all 62 passed` 였고,
    # `install.sh --check` 는 그 마지막 줄만 사용자에게 보여 준다.
    empty = work / 'nothing-beside-me'
    empty.mkdir(parents=True, exist_ok=True)
    saved_here = kit.HERE
    try:
        kit.HERE = empty
        must_fail(kit.load_extras, 'load_extras() with no selftest_*.py to be found')
    finally:
        kit.HERE = saved_here

    rec = {'file': 'erd_area_A.png', 'counts': {k: 0 for k in MEASURES},
           'warn': [], 'tolerated': []}
    eq(verify_faults(rec), [], 'an all-zero area record is clean')
    saved = list(kit._LOGS), kit._DREW[0], kit._SWEPT[0]
    try:
        kit._LOGS.clear()
        must_fail(lambda: verify_recs(work),
                  'verify_recs() with not one record — "found nothing" is not "clean"')
        # 먼저 **깨끗한** 기록 하나로 '그린 장수 ↔ 기록 장수' 만 묻는다. 더러운
        # 기록으로 물으면 훑기가 어느 쪽 규칙 때문에 죽었는지 구분되지 않는다 —
        # 실제로 그렇게 썼다가 대조를 지운 뮤턴트를 놓쳤다.
        log = work / 'probe.jsonl'
        log.write_text(json.dumps(rec) + '\n', encoding='utf-8')
        kit._LOGS.append(log)
        kit._DREW[0] = 2
        must_fail(lambda: kit.sweep_verify('__probe__ none', work),
                  'the sweep with 2 diagrams drawn and only 1 recorded')
        kit._DREW[0] = 1
        kit.sweep_verify('__probe__ none', work)        # 맞으면 지나간다
        log.write_text(json.dumps(dict(rec, counts=dict(rec['counts'], thru=3),
                                       tolerated=['thru'])) + '\n', encoding='utf-8')
        must_fail(lambda: verify_clean(work),
                  'verify_clean() on a record the drawing code itself called tolerable')
        must_fail(lambda: kit.sweep_verify('__probe__ none', work),
                  'the sweep over a diagram that is not clean')
        # 봐주기는 상한이 아니라 정확값이다 — 넓혀 놓고 그 아래 회귀를 숨길 수 없게
        log.write_text(json.dumps(dict(rec, counts=dict(rec['counts'], h_overlap=2)))
                       + '\n', encoding='utf-8')
        kit.RENDER_ALLOW['__probe__ allow'] = {'h_overlap': 9}
        must_fail(lambda: kit.sweep_verify('__probe__ allow', work),
                  'an allowance of 9 over diagrams that peak at 2')
        kit.RENDER_ALLOW['__probe__ allow'] = {'h_overlap': 2}
        kit.sweep_verify('__probe__ allow', work)      # 정확히 맞으면 지나간다
    finally:
        kit.RENDER_ALLOW.pop('__probe__ allow', None)
        kit._LOGS[:], kit._DREW[0], kit._SWEPT[0] = saved[0], saved[1], saved[2]


# ── 조용히 줄기를 막는 톱니 ─────────────────────────────────────────────────
# 14라운드의 신고제(`EXPECT_CASES`)는 '조용히 0' 을 막는다. 그런데 15라운드에
# `selftest_schema.py` 에서 `@case` 셋을 떼고 `EXPECT_CASES` 를 40 → 37 로 함께
# 고쳐 봤더니 `all 138 passed` 로 **조용히 줄었다**. 신고는 '내가 올린 수' 를 말할
# 뿐이라, 신고와 실제를 같은 손으로 함께 낮추면 아무 데서도 안 붉는다.
#
# 그래서 바닥을 따로 둔다. 이 표는 **그 파일 바깥**(입구 파일)에 있으므로 한 파일만
# 손보는 뮤턴트로는 못 내린다. 케이스를 더하는 것은 자유롭고(바닥은 하한이다),
# 줄이려면 여기 숫자를 함께 내려야 한다 — 그러면 그 삭제가 diff 에 남는다.
# `RENDER_ALLOW` 를 '코드를 고쳐서는 늘릴 수 없는 자리' 에 둔 것과 같은 이유다.
#
# 목록에 없는 새 파일은 1 이 바닥이다 — 새 파일에 손으로 적을 것을 늘리지 않는다.
# 값은 15라운드가 끝나는 자리(2026-08-04)의 실측이다.
# 값은 16라운드가 끝나는 자리(2026-08-04)의 실측이다.
# 값은 17라운드 회귀 보강 자리(2026-08-05)의 실측이다 — 그 라운드가 더한 만큼 올린다.
# 바닥을 올리지 않으면 이번에 더한 케이스는 다음 라운드가 조용히 뺄 수 있다.
#
# 2026-08-05: 검토 기록 문서를 저장소에서 뺐다. 그 문서를 열던 `docs:` 케이스 셋이
# 잴 것을 잃었으므로 함께 지웠고, `selftest_install` 의 바닥을 13 → 10, 총계를
# 202 → 199 로 내린다. 없는 파일을 재는 케이스는 아무것도 안 재는 것이다.
#
# 2026-08-05: 그때 함께 잃은 것이 하나 있었다 — 지운 케이스 가운데 하나는 네 언어
# 설치 문서의 `all N passed` 도 같이 재고 있었고, 그 자리가 비면서 그 숫자를 읽는
# 것이 아무것도 없어졌다. `selftest_install` 에 설치 문서만 보는 축소판을 도로
# 세운다: 바닥 10 → 11, 총계 199 → 200.
#
# 2026-08-05: 그 축소판이 설치 문서 넷만 보고 있었는데, 같은 부류의 자리가 둘 더
# 있었다 — `SKILL.md`·`SKILL.ko.md` 의 실행법 블록이 적는 개수다. 거기에는 `101` 과
# `39` 가 남아 있었다: `199` 시절보다도 오래된 수라, 앞 라운드가 `grep "199\|200"`
# 으로 훑었을 때 걸리지 않았다. `39` 가 든 문장은 적힐 당시엔 참이었고(그때 시험
# 파일이 셋뿐이라 그 입구는 정말 제 것만 돌았다) `64b643d` 가 `selftest_r14_*` 넷을
# 더하면서 거짓이 됐다. 손으로 적은 수를 손으로 찾는 것은 이렇게 실패하므로
# 같은 규칙으로 재는 케이스를 하나 더 세운다: 바닥 11 → 12, 총계 200 → 201.
CASE_FLOOR = {'selftest_schema': 40, 'selftest_build': 26,
              'selftest_config': 29, 'selftest_install': 12,
              'selftest_render': 11}

# 그런데 이 표는 **파일 이름으로** 걸려 있다. 15라운드 수정자가 스스로 신고하고
# 16라운드 검증자가 확정한 우회가 그래서 있다: `selftest_render.py` 를
# `selftest_r15_render.py` 로 개명하면서 케이스 4개를 지우면 개명된 이름은 표에 없어
# 바닥이 1 이 되고, `all 157 passed` 가 초록으로 찍힌다. 이름표를 옮기면 그 이름표에
# 걸린 바닥이 함께 사라진다 — 이름에 거는 한 원리적으로 그렇다.
#
# 그래서 이름을 하나도 안 보는 바닥을 하나 더 둔다. 파일이 몇 개든 무슨 이름이든,
# **이 벌 전체가 몇 개인가**는 개명으로 바뀌지 않는다. 개명만 하는 것은 이 바닥을
# 통과하고(그래도 좋다 — 개명 자체는 아무것도 안 잃는다), 개명에 삭제를 섞는 순간
# 총계가 내려가 여기서 붉어진다. 도커 케이스는 서버가 있어야만 등록되므로 뺀다.
#
# 2026-08-05: 이름 안의 이스케이프 큰따옴표(`"a""b"`)를 고치면서 그 자리를 지키는
# 케이스 일곱을 더한다: 217 → 224.
#
# 2026-08-05: `parse_alter` 가 이름 안의 `;` 에서 문을 끊던 것을 고쳤는데, 그 고침을
# **통째로 되돌려도 224개가 전부 초록**이었다 — 자르는 자리가 둘(ADD COLUMN·ADD
# CONSTRAINT)인데 어느 쪽에도 못이 없었다. 넷을 더한다: 224 → 228. 어느 케이스가
# 어느 루프를 잡는지는 그 케이스들 위의 격자표에 적어 두었다.
#
# 2026-08-06: '같은 이름이 두 자리를 가리킬 때' 넷을 더한다: 228 → 232. 넷 다 고침을
# 통째로 되돌려도 228개가 전부 초록이던 자리다 — PK 조회의 조인(제약 이름이 스키마
# 안에서 유일하다고 본 것), 인라인 PK·UNIQUE 판정에 새는 부모 이름, 두 번 적은 유니크
# 인덱스, 그리고 `doc` 안쪽의 모양(한 줄짜리 문자열이 글자마다 문단이 됐다).
TOTAL_FLOOR = 257

# 이 파일이 올리는 케이스 수. 옆의 다섯 파일이 세 라운드째 지키고 있는 규율인데
# **입구 파일만 면제**였다 — 그래서 여기 70개가 신고도 바닥도 없이 있었다.
# 케이스를 더하거나 빼면 이 수와 `selftest_kit.ENTRY_FLOOR['selftest']` 를 함께 고친다.
EXPECT_CASES = 119


@case('selftest: every case file beside the kit is registered and says how many it added')
def _(work):
    # 13라운드가 36개를 만들고 아무 데도 등록하지 않았다. 14라운드가 글로브로 등록했는데
    # 그 장치는 **제가 실패한 것을 스스로 말하지 못했다** — 등록을 통째로 지워도
    # `all 62 passed` 가 초록으로 찍혔다. 총계를 상수로 박으면 다음 라운드에 곧
    # 어긋나므로, 파일마다 제가 올린 수를 신고하게 하고 여기서 그 신고서를 읽는다.
    import selftest_kit as kit
    running = getattr(sys.modules.get('__main__'), '__file__', '')
    want = {p.stem for p in HERE.glob('selftest_*.py')} - {'selftest_kit'}
    if running:
        want -= {os.path.splitext(os.path.basename(running))[0]}
    got = dict(kit.LOADED)
    eq(sorted(got), sorted(want),
       'every selftest_*.py beside the kit reports itself — a file that loads silently '
       'can also fail to load silently')
    for stem, n in sorted(got.items()):
        if n <= 0:
            raise Fail(f'{stem}.py registered {n} cases')
        mod = sys.modules[stem]
        declared = getattr(mod, 'EXPECT_CASES', None)
        # 15라운드. 신고는 선택이 아니다. 신고하지 않는 파일은 `load_extras()` 의
        # '정확히 맞아야 한다' 를 통째로 비켜 가고, 그러면 그 파일에서 케이스가
        # 사라지는 것을 아무것도 막지 않는다 — 아래 톱니도 신고서를 읽으므로 함께
        # 눈이 먼다. 여기서 붉히는 쪽을 골랐지, `load_extras()` 에서 죽이는 쪽은
        # 고르지 않았다: 신고 한 줄을 잊은 것 때문에 **시험이 한 개도 안 도는** 것은
        # 이 저장소가 세 라운드째 물린 바로 그 모양이라, 한 줄 빨강으로 파일 이름을
        # 대는 편이 낫다.
        if declared is None:
            raise Fail(f'{stem}.py registered {n} cases but declares no EXPECT_CASES — '
                       f'add `EXPECT_CASES = {n}` to it; a file that does not say how '
                       f'many it adds can also lose them without saying so')
        eq(n, declared, f'{stem}.py says EXPECT_CASES = {declared}')
        floor = CASE_FLOOR.get(stem, 1)
        if n < floor:
            raise Fail(f'{stem}.py is down to {n} cases from {floor} — a case may only '
                       f'leave with its floor: lower CASE_FLOOR[{stem!r}] in '
                       f'selftest.py in the same change, so the removal shows up in '
                       f'the diff instead of in a tally nobody compares')
    # ── 입구 파일 자신 ──────────────────────────────────────────────────────
    # 15라운드 판은 여기서 `own > 0` 만 봤다. 그래서 입구 파일의 70개(시험 벌의 43%)에
    # 대해서는 신고도 바닥도 없었다: `@case` 넷을 no-op 데코레이터로 바꾸면 `✗` 한 건
    # 없이 `all 157 passed` 가, 케이스 셋을 지우면 `all 158 passed` 가 찍혔다.
    # 위 다섯 파일에 건 규율을 입구 파일에도 똑같이 건다 — 신고(EXPECT_CASES)와
    # 바닥(kit.ENTRY_FLOOR). 바닥이 kit 에 있는 이유는 위 표와 같다: 지키려는 파일
    # 안에 두면 그 파일만 손보는 변경으로 함께 내릴 수 있다.
    own = len(kit.CASES) - sum(got.values())
    if own <= 0:
        raise Fail(f'the entry file registered {own} cases of its own — the tally '
                   f'{len(kit.CASES)} is coming from somewhere unexpected')
    entry = os.path.splitext(os.path.basename(running))[0] if running else ''
    if not entry:
        raise Fail('this suite is running without a file name of its own, so the entry '
                   'file cannot be held to the floor every other file is held to')
    entry_mod = sys.modules.get('__main__')
    declared = getattr(entry_mod, 'EXPECT_CASES', None)
    if declared is None:
        raise Fail(f'{entry}.py is the entry file and registered {own} cases but '
                   f'declares no EXPECT_CASES — add `EXPECT_CASES = {own}` to it; the '
                   f'file that runs the suite is the one file nothing was counting')
    # 입구 파일이 도커 케이스를 안에서 등록할 수 있다 (selftest_schema.py 가 그렇다).
    # 그쪽 `EXPECT_CASES` 는 이미 그 수를 함께 세므로 신고 대조는 그대로 맞고, 바닥은
    # 하한이라 6개가 더 붙어도 넘지 않는다.
    eq(own, declared, f'{entry}.py is the entry file and says EXPECT_CASES = {declared}')
    entry_floor = kit.ENTRY_FLOOR.get(entry, 1)
    if own < entry_floor:
        raise Fail(f'{entry}.py is the entry file and is down to {own} cases from '
                   f'{entry_floor} — a case may only leave with its floor: lower '
                   f'ENTRY_FLOOR[{entry!r}] in selftest_kit.py in the same change, so '
                   f'the removal shows up in the diff instead of in a tally nobody '
                   f'compares')

    # ── 이름을 하나도 안 보는 바닥 ──────────────────────────────────────────
    # 위의 두 표는 전부 **파일 이름**으로 걸려 있다. 개명하면 그 이름에 걸린 바닥이
    # 함께 사라진다(TOTAL_FLOOR 옆의 설명 참고). 총계는 개명으로 안 바뀐다.
    total = sum(1 for n, _f in kit.CASES if not n.startswith('db: '))
    if total < TOTAL_FLOOR:
        raise Fail(f'the suite is down to {total} cases from {TOTAL_FLOOR} — whatever '
                   f'file they left, and whatever that file is called now, lower '
                   f'TOTAL_FLOOR in selftest.py in the same change so the loss shows '
                   f'up in the diff')


# ── 그림 ─────────────────────────────────────────────────────────────────────
@case('render: two self-references draw as two loops')
def _(work):
    write_schema(work, {'cat': table(
        'cat', [col('id'), col('parent_id'), col('root_id')], pk=['id'],
        fks=[{'column': 'parent_id', 'ref_table': 'cat', 'ref_column': 'id',
              'on_delete': 'SET NULL'},
             {'column': 'root_id', 'ref_table': 'cat', 'ref_column': 'id',
              'on_delete': 'SET NULL'}])})
    run('build_erd.py', work)
    verify_clean(work, what='self-loops must not overlap')
    # 루프의 팔 높이는 라우터가 고른 값이라 중첩 검사에서 면제되지 않는다. 두 루프가
    # 포개지면 v_overlap·h_overlap 이 잡는다 — 이 항목이 정말로 재고 있는지(0 이
    # 아니라 '재지 않음' 이 아닌지) 못박는다. 이 항목이 장식이던 판이 있었다.
    for r in verify_recs(work):
        for k in ('v_overlap', 'h_overlap'):
            if r['counts'][k] is None:
                raise Fail(f'{r["file"]}: {k} was not measured at all')
    for r in verify_recs(work, 'area'):
        if r['counts']['label_x'] is None:
            raise Fail('the area diagram draws both loop labels — label↔label must be measured')
    svg = (work / 'out' / 'erd_area_A.svg').read_text(encoding='utf-8')
    eq(svg.count('parent_id'), 2, 'each loop keeps its own label')   # 라벨 + 컬럼행


@case('render: a hub with many children keeps lines out of the tables')
def _(work):
    t = {'hub': table('hub', [col('id'), col('name', 'text')], pk=['id'])}
    for i in range(24):
        n = f'c{i:02d}'
        t[n] = table(n, [col('id'), col('hub_id')], pk=['id'],
                     fks=[{'column': 'hub_id', 'ref_table': 'hub', 'ref_column': 'id',
                           'on_delete': 'CASCADE'}])
    write_schema(work, t)
    run('build_erd.py', work)
    for r in verify_recs(work, 'area'):
        if r['counts']['thru'] != 0:
            raise Fail(f'a line runs through a table: {r["file"]} thru={r["counts"]["thru"]}')


@case('render: many unrelated tables do not become a vertical ribbon')
def _(work):
    # 예전엔 60테이블 하나에 1:4 였다. 열 수를 옛 고정값(4)으로 되돌리면 그 자리는
    # 1:2.30 이라 통과한다 — 고칠 때 잡으라고 둔 시험이 정작 그 회귀를 통과시킨다.
    # 리본이 되는지는 **테이블이 늘 때 비율이 어떻게 되는가**로 재야 보인다:
    # 고정 4열이면 60→200 에서 1:2.30 → 1:7.26 으로 자라고(열은 그대로, 세로만 길어짐),
    # 열 수가 따라 늘면 1:1.49 → 1:1.82 로 거의 그대로다.
    from PIL import Image
    ratio = {}
    for n in (60, 200):
        write_schema(work, {f't{i:03d}': table(f't{i:03d}', [col('id'), col('v', 'text')])
                            for i in range(n)})
        run('build_erd.py', work)
        w, h = Image.open(work / 'out' / 'erd_full.png').size
        ratio[n] = h / w
        if ratio[n] > 3:
            raise Fail(f'{n} tables → aspect ratio 1:{ratio[n]:.2f} — '
                       f'unreadable once fitted into a document')
    if ratio[200] > ratio[60] * 1.5:
        raise Fail(f'the sheet only grows downward: 1:{ratio[60]:.2f} at 60 tables → '
                   f'1:{ratio[200]:.2f} at 200 — columns are not keeping up')


@case('render: badge clears the name on a title-only box')
def _(work):
    # 컬럼을 모르는 테이블(참조만 되고 정의가 없는 것)도 배지를 단다. 상자 폭이
    # 제목 폭만 보고 정해지면 배지가 이름 위에 얹힌다 — 픽셀 대신 기하로 잰다:
    # 오른쪽 정렬로 그려질 배지의 왼쪽 끝이 제목의 오른쪽 끝보다 오른쪽이어야 한다.
    write_schema(work, {
        'orders': table('orders', [col('id'), col('m')], origin='new', pk=['id'],
                        fks=[{'column': 'm', 'ref_table': 'merchants',
                              'ref_column': 'id', 'on_delete': 'NO ACTION'}]),
        'merchants': table('merchants', [])})
    probe = work / 'probe.py'
    probe.write_text(
        "import erd\n"
        "f = erd.load_fonts()\n"
        "for name in erd.SCHEMA:\n"
        "    box = erd.measure(name)\n"
        "    bd, _c = erd.badge(name)\n"
        "    name_right = erd.PAD + erd.tw(name, f['title'])\n"
        "    badge_left = box['w'] - erd.PAD - erd.tw(bd, f['badge'])\n"
        "    assert badge_left > name_right, (\n"
        "        f'{name}: badge at x={badge_left} overlaps name ending at x={name_right}')\n",
        encoding='utf-8')
    run(str(probe), work, env={'PYTHONPATH': str(HERE)})


# ── 산출물 ───────────────────────────────────────────────────────────────────
@case('artifacts: html, docx and graphml describe the same schema')
def _(work):
    t = {'a': table('a', [col('id'), col('x', 'text', comment='note a')], pk=['id']),
         'b': table('b', [col('id'), col('a_id')], pk=['id'],
                    fks=[{'column': 'a_id', 'ref_table': 'a', 'ref_column': 'id',
                          'on_delete': 'CASCADE'}])}
    write_schema(work, t)
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    run('build_docx.py', work)
    # 세 벌 다 **자리를 못박아** 본다. 예전엔 GraphML·docx 쪽이 맨 `in` 검사였다:
    # graphml_node() 가 빈 문자열을 돌려주게 해도 통과했다 — <node> 가 하나도 없는
    # GraphML 인데, 남은 <edge> 하나가 'a' 와 'b' 를 둘 다 품고 있어서다. GraphML 은
    # 이 스킬이 내건 네 산출물 중 하나인데 텅 빈 채로 나갈 수 있었다.
    html = (work / 'T.html').read_text(encoding='utf-8')
    for name in t:
        has(html, f'<h4 id="tb_{name}">', f'{name} has its own section in the HTML')

    import xml.etree.ElementTree as ET
    NS = {'g': 'http://graphml.graphdrawing.org/xmlns',
          'y': 'http://www.yworks.com/xml/graphml'}
    root = ET.parse(work / 'T.graphml').getroot()
    nodes = root.findall('.//g:node', NS)
    titles = []
    for n in nodes:
        for lab in n.findall('.//y:NodeLabel', NS):
            if 'label.name' in (lab.get('configuration') or ''):
                titles.append((lab.text or '').split('  ·  ')[0])
    eq(sorted(titles), sorted(t), 'one GraphML <node> per table, each carrying its name')
    edges = root.findall('.//g:edge', NS)
    eq([(e.get('source'), e.get('target')) for e in edges],
       [(nodes[1].get('id'), nodes[0].get('id'))], 'the one FK is the one GraphML edge')
    eq([e.findtext('g:data[@key="d3"]', namespaces=NS) for e in edges], ['a_id : id'],
       'and it says which columns it joins')

    # docx 도 4장의 테이블별 절 제목으로 못박는다 — 어디든 이름이 스치기만 하면
    # 통과하던 검사는 컬럼표가 통째로 빠져도 몰랐다.
    from docx import Document
    doc = Document(str(work / 'T.docx'))
    heads = [p.text for p in doc.paragraphs if re.match(r'^4\.\d+\.\d+ ', p.text)]
    eq(sorted(h.split(' ', 1)[1].split(' · ')[0] for h in heads), sorted(t),
       'each table gets its own section in chapter 4 of the docx')
    cells = {c.text for tb in doc.tables for row in tb.rows for c in row.cells}
    for name in t:
        for c in t[name]['columns']:
            if c['name'] not in cells:
                raise Fail(f'the docx column table lost {name}.{c["name"]}')


@case('artifacts: graphml and svg are well-formed xml')
def _(work):
    write_schema(work, {'t': table('t', [
        col('id'), col('x', 'text', comment='an & ampersand <tag> "quote" \x0b')])})
    run('build_erd.py', work)
    import xml.etree.ElementTree as ET
    ET.parse(work / 'T.graphml')
    for svg in (work / 'out').glob('*.svg'):
        ET.parse(svg)


@case('artifacts: html escapes what a comment may contain')
def _(work):
    write_schema(work, {'t': table('t', [
        col('id'), col('x', 'text', comment='<script>alert(1)</script>')])})
    run('build_erd.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    if '<script>alert(1)</script>' in html:
        raise Fail('a comment must not become live markup')


@case('artifacts: no unresolved message keys anywhere')
def _(work):
    # 표식(⟨키⟩)은 i18n.py 것이다. 예전엔 그 모양을 여기에 그대로 적어 두어서, 저쪽이
    # 표식을 바꾸면 이 케이스는 아무것도 못 보면서 계속 통과했다. 표식을 **저쪽에
    # 물어서** 만든다 — 그리고 물어본 답이 키를 담고 있지 않으면 그것부터 실패다.
    write_schema(work, {'t': table('t', [col('id')])})
    sys.path.insert(0, str(HERE))
    from i18n import t
    probe = 'zz_probe.no_such_key'
    mark = t(probe)
    if probe not in mark or mark == probe:
        raise Fail(f'i18n no longer marks an unresolved key ({mark!r}) — '
                   f'this case would look at a document it cannot read')
    lead, tail = mark.split(probe)
    pat = re.escape(lead) + r'[a-z_][a-z._0-9]*' + re.escape(tail)
    for lang in ('en', 'ko', 'ja', 'es'):
        run('merge_desc.py', work, env={'ERD_LANG': lang})
        run('build_erd.py', work, env={'ERD_LANG': lang})
        r = run('build_html.py', work, env={'ERD_LANG': lang})
        text = (work / 'T.html').read_text(encoding='utf-8')
        eq(re.findall(pat, text), [], f'{lang}: unresolved keys in the document')
        # 콘솔로 새는 것도 마찬가지다 — 사용자가 먼저 보는 자리다
        eq(re.findall(pat, r.stdout + r.stderr), [], f'{lang}: unresolved keys on the console')


@case('artifacts: docx pictures fit the page')
def _(work):
    write_schema(work, {f't{i}': table(f't{i}', [col('id'), col('v', 'text')])
                        for i in range(12)})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    r = run('build_docx.py', work)
    from docx import Document
    shapes = Document(str(work / 'T.docx')).inline_shapes
    # 그림이 한 장도 없어도 이 반복문은 조용히 지나갔다 — add_picture 를 통째로 지워도
    # 통과했다. '그림이 페이지에 맞는다' 와 '그림이 없다' 를 구분하지 못하는 검사다.
    # 그려 놓은 PNG 마다 한 장씩 들어가야 한다: 개요도 · 전체도 · 영역별 상세도.
    pngs = sorted(p.name for p in (work / 'out').glob('*.png'))
    eq(len(shapes), len(pngs), f'one picture per diagram drawn ({pngs})')
    if len(pngs) < 3:
        raise Fail(f'the fixture must draw more than a token diagram: {pngs}')
    if 'warn' in r.stdout:
        raise Fail(f'no diagram should have been left out:\n{r.stdout}')
    for s in shapes:
        if s.width.cm > 26.7 or s.height.cm > 18.0:
            raise Fail(f'picture {s.width.cm:.1f}×{s.height.cm:.1f}cm exceeds the page')
        if s.width.cm < 1 or s.height.cm < 1:
            raise Fail(f'picture {s.width.cm:.1f}×{s.height.cm:.1f}cm is not a diagram')


@case('artifacts: the counters describe the PNG, and each builder says which one it ships')
def _(work):
    # 14라운드. 검증 숫자가 말하는 그림과 문서에 실리는 그림이 서로 다르다는 것이
    # 나왔다. `build_html.py` 는 SVG 를 인라인으로 박고 `build_docx.py` 는 PNG 를
    # 넣는데, 다섯 카운터는 **PNG 한 장만** 잰다 — SVG 는 draw_erd 가 배율을 1 로
    # 되돌려 한 번 더 그리는 세 번째 렌더라, 글자 폭이 배율에 비례하지 않는 만큼
    # 라벨 자리가 달라진다. 자기참조를 섞은 무작위 스키마 20벌을 재 보니 라벨 하나가
    # 134px 옮겨 앉은 판까지 있었다 (라벨의 **글자 목록**은 언제나 같았다).
    #
    # 그 어긋남 자체는 erd.py 안에서만 고칠 수 있다. 여기서 못박는 것은 **문서가
    # 적어 둔 사실** 이다 — 숫자는 PNG 것이고, HTML 은 SVG 를, docx 는 PNG 를
    # 싣는다. 어느 쪽이 조용히 바뀌면 SKILL.md 의 그 문단이 거짓이 된다.
    # SVG 가 PNG 배율을 함께 쓰게 되면 아래 단정은 그대로 참이고, 그때 비로소
    # 숫자가 두 산출물을 함께 말한다.
    import zipfile
    write_schema(work, {
        'a': table('a', [col('id'), col('b_id')], pk=['id'],
                   fks=[{'column': 'b_id', 'ref_table': 'b', 'ref_column': 'id',
                         'on_delete': 'CASCADE'}]),
        'b': table('b', [col('id')], pk=['id'])})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    for r in verify_recs(work):
        if not r['file'].endswith('.png'):
            raise Fail(f'the counters name {r["file"]} — this test and SKILL.md both say '
                       f'they measure the PNG')
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    has(html, '<div class="fig"><svg', 'the html inlines the SVG, not the measured PNG')
    if 'data:image/png' in html:
        raise Fail('the html fell back to PNG — then SKILL.md must stop saying it ships '
                   'the vector')
    run('build_docx.py', work)
    with zipfile.ZipFile(work / 'T.docx') as z:
        media = sorted(n.rsplit('/', 1)[-1] for n in z.namelist()
                       if n.startswith('word/media/'))
    if not media:
        raise Fail('the docx carries no picture at all — nothing here is measured')
    odd = [n for n in media if not n.lower().endswith('.png')]
    if odd:
        raise Fail(f'the docx no longer ships the measured raster: {odd}')


# ── 두 번째 실행 ─────────────────────────────────────────────────────────────
# 한 번 돌렸을 때는 다 맞는데 두 번째부터 어긋나는 것들. 첫 실행만 보면 안 보인다.
def with_manual(work, manual):
    """MANUAL 을 채워 둔 채로 merge_desc 를 돌린다 — 사용자가 사전에 적어 넣은 상태.

    사전은 소스에 적는 것이라 밖에서 넣을 길이 없다. 그래서 임시 디렉토리에 작은
    실행기를 하나 놓고 그것을 돌린다 (run() 은 절대경로면 그대로 쓴다).
    """
    drv = work / 'manual_run.py'
    drv.parent.mkdir(parents=True, exist_ok=True)
    drv.write_text('import sys\n'
                   f'sys.path.insert(0, {str(HERE)!r})\n'
                   'import merge_desc\n'
                   f'merge_desc.MANUAL = {manual!r}\n'
                   'merge_desc.main()\n', encoding='utf-8')
    return run(str(drv), work)


def undescribed(out):
    """'columns still without a description' 목록에 찍힌 키."""
    tail = out.split('without a description:')[-1]
    return re.findall(r'^\s+(\S+)\s+\(', tail, re.M)


@case('merge_desc: the key the tool prints is the key that works')
def _(work):
    write_schema(work, {
        'analytics.events': table('events', [col('kind', 'text')], schema='analytics'),
        'staging.events': table('events', [col('kind', 'text')], schema='staging')})
    keys = undescribed(with_manual(work, {}).stdout)
    eq(sorted(keys), ['analytics.events.kind', 'staging.events.kind'],
       'the list names each column by its own table key')
    out = with_manual(work, {k: f'desc of {k}' for k in keys}).stdout
    eq(undescribed(out), [], 'pasting those keys into MANUAL describes every column')
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq({k: t['columns'][0]['comment'] for k, t in s.items()},
       {k.rsplit('.', 1)[0]: f'desc of {k}' for k in keys},
       'and each key lands on the table it names')


@case('merge_desc: a bare key does not bleed into a same-named table')
def _(work):
    write_schema(work, {
        'shop.users': table('users', [col('grade', 'text')], schema='shop'),
        'mart.users': table('users', [col('grade', 'text')], schema='mart'),
        'shop.orders': table('orders', [col('channel', 'text')], schema='shop')})
    out = with_manual(work, {'users.grade': 'shop grade',
                             'shop.users.grade': 'only shop',
                             'orders.channel': 'inbound channel'}).stdout
    s = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    eq({k: t['columns'][0]['comment'] for k, t in s.items()},
       {'shop.users': 'only shop', 'mart.users': '', 'shop.orders': 'inbound channel'},
       'the qualified key wins, the ambiguous bare key is dropped, '
       'and a bare key still works where the name is unique')
    has(out, 'shop.users.grade', 'the ignored bare key is reported with what to use instead')


@case('artifacts: a diagram older than the schema does not reach the document')
def _(work):
    write_schema(work, {'keep': table('keep', [col('id')]),
                        'legacy_audit': table('legacy_audit', [col('id')])})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    run('build_html.py', work)
    has((work / 'T.html').read_text(encoding='utf-8'), 'legacy_audit',
        'the first edition describes both tables')

    # 시간이 흐른 척한다: 스키마만 새로 쓰고 그림은 한 시간 전으로 돌려 둔다
    write_schema(work, {'keep': table('keep', [col('id')])})
    old = (work / 'schema.json').stat().st_mtime - 3600
    for p in (work / 'out').iterdir():
        os.utime(p, (old, old))

    for script in ('build_html.py', 'build_docx.py'):
        r = run(script, work, expect_ok=False)
        if r.returncode == 0:
            raise Fail(f'{script} embedded diagrams older than the schema')
        if 'Traceback' in r.stderr:
            raise Fail(f'{script} should say what is wrong, not traceback:\n{r.stderr[-300:]}')
        has(r.stdout + r.stderr, 'build_erd.py', f'{script} says how to fix it')
    has((work / 'T.html').read_text(encoding='utf-8'), 'legacy_audit',
        'the refused build leaves the previous edition intact, not half-rewritten')

    run('build_erd.py', work)                      # 다시 그리면 지나간다
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    if 'legacy_audit' in html:
        raise Fail('after redrawing, the dropped table must be gone from text and figures')


# ── 오류 경로 ────────────────────────────────────────────────────────────────
@case('errors: a missing database is explained, not tracebacked')
def _(work):
    r = run('introspect.py', work, expect_ok=False)
    if 'Traceback' in r.stderr:
        raise Fail(f'should be a message:\n{r.stderr[-300:]}')
    has(r.stdout + r.stderr, 'ERD_PSQL', 'the message names what to set')


@case('errors: an unusable font is explained')
def _(work):
    write_schema(work, {'t': table('t', [col('id')])})
    r = run('build_erd.py', work, env={'ERD_FONT': '/etc/hosts'}, expect_ok=False)
    if 'Traceback' in r.stderr:
        raise Fail('a non-font file should be a message, not a PIL traceback')
    has(r.stdout + r.stderr, 'ERD_FONT', 'the message names the variable')


@case('errors: a diagram file that is simply missing is a message, not a traceback')
def _(work):
    # 신선도 관문(require_fresh)은 **없는 파일을 세지 않는다** — 일부러 그렇다. 그래서
    # 그림 하나를 지운 뒤 문서를 만들면 관문은 그대로 지나가고, build_docx.py 는 PIL 이
    # 없는 파일을 여는 자리에서 FileNotFoundError 를 사용자에게 그대로 뱉었다.
    # build_html.py 는 같은 상황을 도판만 빼고 경고 한 줄로 넘긴다. 둘이 같아야 한다.
    write_schema(work, {'a': table('a', [col('id')]), 'b': table('b', [col('id')])})
    run('merge_desc.py', work)
    run('build_erd.py', work)
    for ext in ('.png', '.svg'):     # HTML 은 PNG 가 없으면 SVG 로 떨어진다
        (work / 'out' / f'erd_full{ext}').unlink()
    # 빠진 도판이 조용히 빠지면 그림 없는 문서가 그림 있는 문서인 척 나간다 — 둘 다
    # 어느 그림을 못 넣었는지 이름을 대야 한다 (HTML 쪽은 영역 그림만 세고 있었다)
    for script in ('build_html.py', 'build_docx.py'):
        r = run(script, work)
        if 'Traceback' in r.stderr:
            raise Fail(f'{script} tracebacked on a missing diagram:\n{r.stderr[-400:]}')
        has(r.stdout, 'erd_full', f'{script} says which diagram it had to leave out')
    from docx import Document
    pngs = list((work / 'out').glob('*.png'))
    eq(len(Document(str(work / 'T.docx')).inline_shapes), len(pngs),
       'the document still carries every diagram that does exist')
    # 다시 그리면 아무 말 없이 원래대로
    run('build_erd.py', work)
    r = run('build_docx.py', work)
    if 'warn' in r.stdout:
        raise Fail(f'a complete set of diagrams must not warn:\n{r.stdout}')


@case('errors: an old schema.json without new keys still renders')
def _(work):
    # 예전엔 종료 코드 0 말고는 아무것도 보지 않았다. 그림을 한 장도 안 그리고 끝나도
    # 통과였다 — '죽지 않았다' 와 '렌더된다' 는 다른 말이다.
    work.mkdir(parents=True, exist_ok=True)
    (work / 'schema.json').write_text(json.dumps(
        {'t': {'name': 't', 'columns': [{'name': 'id', 'type': 'bigint', 'comment': ''}],
               'fks': [{'column': 'p', 'ref_table': 'u', 'ref_column': 'id'}]},
         'u': {'name': 'u', 'columns': [{'name': 'id', 'type': 'bigint', 'comment': ''}]}}), encoding='utf-8')
    run('build_erd.py', work)
    from PIL import Image
    for stem in ('erd_overview', 'erd_full'):
        p = work / 'out' / f'{stem}.png'
        if not p.exists():
            raise Fail(f'{stem}.png was never drawn — exit 0 is not a rendered diagram')
        if min(Image.open(p).size) < 50:
            raise Fail(f'{stem}.png is {Image.open(p).size} — an empty canvas, not a diagram')
    if not list((work / 'out').glob('erd_area_*.png')):
        raise Fail('no area diagram came out of an old schema.json')
    # 새 키가 없어도 옛 스키마의 테이블·관계가 문서까지 간다
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    for name in ('t', 'u'):
        has(html, f'<h4 id="tb_{name}">', f'{name} reached the document from an old schema')


# ── 같은 이름이 두 자리를 가리킬 때 ──────────────────────────────────────────
# 아래 넷은 전부 '이름 하나가 뜻을 하나만 가리킨다' 를 믿었다가 틀린 자리다. 제약
# 이름은 스키마가 아니라 테이블 안에서 유일하고, 컬럼 선언 안의 낱말은 그 컬럼 것일
# 수도 부모 이름일 수도 있고, 같은 인덱스를 두 번 적은 DDL 이 유니크 제약 둘이 되지는
# 않으며, spec 의 한 줄짜리 값은 줄 하나지 글자 목록이 아니다.
@case('introspect: a primary key is joined by table, not by constraint name alone')
def _(work):
    # Postgres 에서 유일한 것은 (테이블, 제약 이름) 쌍이다. PK·UNIQUE 는 같은 이름의
    # 인덱스를 만들어 저희끼리 못 겹치지만 FK·CHECK 는 인덱스를 안 만들므로 남의
    # 테이블에서 같은 이름을 그대로 쓸 수 있다:
    #     create table a (id bigint, x text, constraint foo primary key (id));
    #     create table b (x bigint, constraint foo foreign key (x) references a(id));
    # 이름만으로 이으면 b 의 FK 컬럼이 a 의 PK 목록에 딸려 와 `a.pk = ['id','x']` 가
    # 된다 — 실측(PG16)으로 정의서의 a 표에서 x 가 `● PK` 로, 굵게 실렸다.
    #
    # 서버 없이 재므로 **조회문 자체**를 본다. 이 자리를 결과로 재려면 도커가 필요한데
    # (그 케이스들은 selftest_history.py 에 있고 기본으로는 안 돈다), 그러면 평소
    # 실행에서는 아무도 이 자리를 안 지킨다. 조회가 무엇을 무엇에 맞추는지는 조회문에
    # 다 적혀 있으므로, 여기서는 그 술어가 있는지를 잰다.
    sys.path.insert(0, str(HERE))
    import introspect
    q = ' '.join(introspect.Q_PK.split())
    join = q[q.index('key_column_usage'):q.index('where')]
    for a, b in (('kcu.table_name', 'tc.table_name'),
                 ('kcu.constraint_name', 'tc.constraint_name'),
                 ('kcu.table_schema', 'tc.table_schema')):
        if f'{a}={b}' not in join and f'{b}={a}' not in join:
            raise Fail(f'Q_PK joins key_column_usage without matching {a} to {b} — '
                       f"another table's constraint of the same name comes along\n"
                       f'      join: {join}')
    # 컬럼 차례는 테이블별로 읽힌다 — 여러 테이블의 행이 섞여 와도 순서가 남게.
    order = q[q.index('order by'):]
    for piece in ('tc.table_name', 'kcu.ordinal_position'):
        has(order, piece, f'Q_PK orders by {piece}')


@case('parse: a parent name in REFERENCES is not read as this column\'s own key')
def _(work):
    # `mask` 는 큰따옴표 식별자를 일부러 안 가린다(`_IDC` 가 거기서 이름을 읽어야
    # 한다). 그래서 부모 쪽 이름이 그대로 판정에 샜다 — not_null·identity 는
    # `_decl()` 로 REFERENCES 절을 걷어낸 뒤에 보고 있었는데 PK·UNIQUE 만 규율 밖이라,
    # 같은 함수 안에서 같은 규칙이 두 벌이던 자리다.
    s = ddl(work, '''
CREATE TABLE p ("PRIMARY KEY" bigint, "UNIQUE" text);
CREATE TABLE t (
  id bigint PRIMARY KEY,
  x  int REFERENCES p ("PRIMARY KEY"),
  y  int REFERENCES p ("UNIQUE"),
  z  int NOT NULL REFERENCES p ("PRIMARY KEY")
);
''')
    eq(s['t']['pk'], ['id'], "a parent column named \"PRIMARY KEY\" is not this table's PK")
    eq(s['t']['uniques'], [], 'a parent column named "UNIQUE" does not invent a UNIQUE')
    # 걷어낸 사본으로 보면서도 관계와 NOT NULL 은 그대로 읽혀야 한다 — REFERENCES 는
    # 원본에서 읽고 나머지만 걷어낸 사본에서 본다.
    eq(sorted((f['column'], f['ref_column']) for f in s['t']['fks']),
       [('x', 'PRIMARY KEY'), ('y', 'UNIQUE'), ('z', 'PRIMARY KEY')],
       'the relationships themselves survive')
    eq([c['not_null'] for c in s['t']['columns'] if c['name'] == 'z'], [True],
       'NOT NULL after a REFERENCES clause still counts')


@case('parse: the same unique index written twice is one constraint, not two')
def _(work):
    # 최종 상태에 유니크 제약이 둘인 DB 는 없다. `parse_alter` 는 인라인 UNIQUE 를
    # 붙일 때 이미 걸러 내고 있었고 `parse_unique` 만 규율 밖이라, 마이그레이션 파일이
    # 같은 인덱스를 다시 적으면 정의서의 제약 목록에 같은 줄이 두 번 실렸다.
    s = ddl(work, '''
CREATE TABLE t (id bigint PRIMARY KEY, code text UNIQUE, a int, b int);
CREATE UNIQUE INDEX i1 ON t (code);
CREATE UNIQUE INDEX IF NOT EXISTS i1 ON t (code);
CREATE UNIQUE INDEX i2 ON t (a, b);
''')
    eq(s['t']['uniques'], [['code'], ['a', 'b']],
       'a repeated unique index is carried once, and the others keep their order')
    # 문서까지 따라간다 — 세는 자리가 하나면 접힌 제약 목록도 한 줄이다.
    run('build_erd.py', work)
    run('build_html.py', work)
    html = (work / 'T.html').read_text(encoding='utf-8')
    eq(html.count('UNIQUE (code)'), 1, 'the document lists that constraint once')


@case('spec: a doc value of the wrong shape names itself instead of splitting')
def _(work):
    # `doc` 은 'dict 이기만 하면' 통과였다. 그 안쪽은 저장소에서 사람이 가장 많이 손으로
    # 쓰는 자리인데도 아무도 안 재는 자리라, 한 줄짜리 문자열을 적으면 build_docx 의
    # `for line in DOC['scope']` 가 **글자마다** 문단을 하나씩 만들었다 — 죽지도, 알리지도
    # 않고 완성된 얼굴로 나왔다.
    write_schema(work, {'t': table('t', [col('id')])})
    for value, key in (({'scope': 'one line'}, 'doc.scope'),
                       ({'sources': 'information_schema'}, 'doc.sources'),
                       ({'scope': [['a']]}, 'doc.scope'),
                       ({'title': 123}, 'doc.title'),
                       ({'area_desc': ['x']}, 'doc.area_desc')):
        (work / 'erd.spec.json').write_text(json.dumps({'doc': value}), encoding='utf-8')
        r = run('build_erd.py', work, expect_ok=False)
        if 'Traceback' in r.stderr:
            raise Fail(f'a malformed {key} should be a message, not a traceback:\n'
                       f'{r.stderr[-400:]}')
        has(r.stdout + r.stderr, key, f'the message names {key}')
    # 성한 값은 그대로 지나간다. `mapping` 의 문자열 행은 `meta_cells` 가 한 칸으로
    # 받아 주기로 이미 정해 둔 동작이라 여기서 더 엄하게 굴지 않는다.
    (work / 'erd.spec.json').write_text(json.dumps(
        {'doc': {'title': 'T', 'scope': ['a', 'b'], 'sources': [['x', 'y']],
                 'meta': [['a', 'b', 'c', 'd']],
                 'mapping': [['a whole row as one string']]}}),
        encoding='utf-8')
    run('build_erd.py', work)
    run('build_docx.py', work)


# ── 제 말로 지은 이름·제 말로 쓴 문서 ────────────────────────────────────────
# 이 스킬은 네 말(en·ko·ja·es)로 문서를 내고 그 말로 지은 이름을 받는다. 그런데 세
# 자리가 영어 기준으로만 짜여 있었다 — 이름의 글자, 장(章)의 손잡이, 글꼴 이름.
@case('merge: a quoted non-ASCII table name inherits its polished descriptions')
def _(work):
    # 이전 판 문서에서 설명을 물려받는 기능이 있는 이유는 **사람이 다듬은 말을 다시
    # 뽑을 때 잃지 않는 것**이다. 그런데 이름을 거르는 그물이 `[A-Za-z_][A-Za-z0-9_.]*`
    # 라, 따옴표로 지은 한글 이름은 그 말을 통째로 잃었다. 게다가 조용하다 — 인계
    # 건수는 찍히니 넘어온 줄 알고, 빠진 것은 '아직 설명 없는 컬럼' 에 섞여 처음
    # 만든 컬럼처럼 보인다. 한국어로 쓰는 사람이 한국어 이름에서만 손해를 봤다.
    s = ddl(work, 'CREATE TABLE "주문" (id bigint PRIMARY KEY, "금액" numeric(10,2));\n'
                  'CREATE TABLE orders (id bigint PRIMARY KEY, amt numeric(10,2));\n')
    for t in s.values():
        for c in t['columns']:
            c['comment'] = '다듬은 설명: ' + c['name']
    write_schema(work, s)
    run('build_erd.py', work)
    run('build_html.py', work)

    for t in s.values():                     # 설명을 지우고 이전 판에서 받아 온다
        for c in t['columns']:
            c['comment'] = ''
    write_schema(work, s)
    run('merge_desc.py', work, env={'ERD_DOC_HTML': str(work / 'T.html')})

    got = json.loads((work / 'schema.json').read_text(encoding='utf-8'))
    for tkey, cname in (('주문', 'id'), ('주문', '금액'), ('orders', 'id'), ('orders', 'amt')):
        c = next(c for c in got[tkey]['columns'] if c['name'] == cname)
        eq(c['comment'], '다듬은 설명: ' + cname,
           f'{tkey}.{cname} came back from the previous edition')


@case('doc: chapter 7 takes its lead paragraph from the spec, like chapter 6')
def _(work):
    # 6장은 `doc.mapping_intro` 로 머리말을 바꿀 수 있는데 7장만 카탈로그에 못박혀
    # 있었다. 두 장이 같은 자리에서 재료를 받고 같은 규칙으로 실리는데 손잡이만
    # 한쪽에 있었고, 그 비대칭은 어느 문서에도 적혀 있지 않다.
    write_schema(work, {'t': table('t', [col('id')])})
    (work / 'erd.spec.json').write_text(json.dumps({'doc': {
        'title': 'T',
        'mapping': [['1', 'a', 'b', 'O', 'c']], 'mapping_intro': 'MY SIXTH',
        'open_items': [['P1', 'x', 'y', 'z', 'w']], 'open_intro': 'MY SEVENTH',
    }}), encoding='utf-8')
    run('build_erd.py', work)
    run('build_docx.py', work)
    run('build_html.py', work)

    from docx import Document
    said = [p.text for p in Document(str(work / 'T.docx')).paragraphs]
    for want in ('MY SIXTH', 'MY SEVENTH'):
        if want not in said:
            raise Fail(f'{want} never reached the docx — the spec wrote it')
    html = (work / 'T.html').read_text(encoding='utf-8')
    for want in ('MY SIXTH', 'MY SEVENTH'):
        has(html, want, f'{want} reached the html too')
    # 안 적으면 카탈로그 문구로 돌아간다 — 손잡이를 다는 것과 기본값을 바꾸는 것은 다르다
    (work / 'erd.spec.json').write_text(json.dumps({'doc': {
        'title': 'T', 'open_items': [['P1', 'x', 'y', 'z', 'w']]}}), encoding='utf-8')
    run('build_erd.py', work)
    run('build_docx.py', work)
    said = '\n'.join(p.text for p in Document(str(work / 'T.docx')).paragraphs)
    sys.path.insert(0, str(HERE))
    import lang.en as en
    has(said, en.M['docx.ch7_intro'], 'without the key, chapter 7 keeps the catalogue text')


@case('merge: SQLAlchemy 1.x Column comments are inherited like Mapped comments')
def _(work):
    # SQLAlchemy 1.x declarative models are still widespread.  They assign
    # Column(...) rather than annotating Mapped[...]; both forms must preserve
    # the same preceding/inline comment semantics.
    import config
    import merge_desc
    models = work / 'models'
    models.mkdir(parents=True)
    (models / 'models.py').write_text(
        "class User(Base):\n"
        "    __tablename__ = 'users'\n"
        "    # 회원 번호\n"
        "    id = Column(BigInteger, primary_key=True)\n"
        "    email = sqlalchemy.Column(Text)  # 로그인 이메일\n"
        "    name = mapped_column(String)  # 표시 이름\n", encoding='utf-8')
    old = config.MODEL_DIR
    try:
        config.MODEL_DIR = models
        got = merge_desc.parse_orm()
    finally:
        config.MODEL_DIR = old
    eq(got, {'users.id': '회원 번호', 'users.email': '로그인 이메일',
             'users.name': '표시 이름'},
       'both SQLAlchemy declarative assignment forms keep their comments')


@case('svg: every font erd calls monospace is recognised as monospace here')
def _(work):
    # `svg_canvas` 는 글꼴 **이름**으로 고정폭인지 짐작한다. 그 그물이 `'mono' in 이름`
    # + `('Menlo','D2Coding')` 뿐이던 때 `erd._MONO_CANDIDATES` 의 후보 하나가 빠져
    # 나갔다 — Windows 의 `consola.ttf`, 곧 `Consolas` 다. 그러면 컬럼 이름·타입이
    # PNG 에서는 고정폭인데 SVG 에서는 본문 글꼴로 나가, 두 그림이 같다는 이 저장소의
    # 전제가 거기서 깨진다. 리눅스·macOS 에서 도는 시험은 후보 넷 중 셋만 밟으므로
    # 파일이 있는지로 재면 이 자리는 영영 안 걸린다 — **이름으로 잰다.**
    import svg_canvas
    for fam in ('Menlo', 'Consolas', 'DejaVu Sans Mono', 'Liberation Mono', 'D2Coding'):
        if not svg_canvas.is_mono(fam):
            raise Fail(f'{fam} is a monospace face erd may pick, but svg_canvas would '
                       f'send it out with the body stack')
    for fam in ('Pretendard', 'Helvetica', 'Apple SD Gothic Neo', '', None):
        if svg_canvas.is_mono(fam):
            raise Fail(f'{fam!r} is not monospace — the column tables would lose their grid')
    # 그림에서도 그렇다: 고정폭으로 그린 글자는 mono 스택을 달고 나간다.
    write_schema(work, {'t': table('t', [col('id')])})
    run('build_erd.py', work)
    svg = (work / 'out' / 'erd_full.svg').read_text(encoding='utf-8')
    if svg_canvas.MONO_STACK not in svg:
        raise Fail('no text in the diagram carries the monospace stack')


@case('install: the run-it-yourself line names every script the quick start does')
def _(work):
    # 설치가 끝나고 마지막에 찍는 그 줄은 **사람이 그대로 복사해 치는 줄**이다. 그런데
    # 네 말 전부 `build_html.py` 가 빠져 있었다 — SKILL.md 의 빠른 시작은 다섯 단계인데
    # 안내는 넷이었다. 그 줄만 따르면 HTML 스키마 정의서가 안 나오고, SKILL.md 가
    # '한 개만 보내면 된다' 고 크게 다루는 산출물이 있는 줄도 모른다.
    #
    # 이 저장소는 문서가 약속한 것을 도구가 하는지 이미 여러 자리에서 잰다(케이스 수·
    # 도커 개수). 실행 순서도 그 부류다 — 두 자리에 적힌 같은 목록이 갈리면 조용하다.
    root = HERE.parent
    quick = {}
    for name, head in (('SKILL.md', '## Quick start'), ('SKILL.ko.md', '## 빠른 시작')):
        text = (root / name).read_text(encoding='utf-8')
        i = text.index(head)
        j = text.index('\n## ', i + len(head))
        # 그 절의 **첫** 코드블록만 본다 — 뒤에 오는 다중 DB 예시는 다른 이야기다.
        block = text[i:j].split('```')[1]
        quick[name] = re.findall(r'^python3 (\w+\.py)', block, re.M)
        if len(quick[name]) < 4:
            raise Fail(f'{name}: the quick start block reads as {quick[name]} — '
                       f'the case is measuring the wrong block')
    eq(quick['SKILL.ko.md'], quick['SKILL.md'], 'both SKILL documents run the same order')

    sh = (root / 'install.sh').read_text(encoding='utf-8')
    lines = [ln for ln in sh.splitlines() if ln.lstrip().startswith('next)')]
    eq(len(lines), 4, 'install.sh carries a closing note in four languages')
    for ln in lines:
        eq(re.findall(r'%s (\w+\.py)', ln), quick['SKILL.md'],
           f'the closing note runs what the quick start runs\n      note: {ln[:80]}…')
    # 자리표시자 수와 넘기는 인자 수가 어긋나면 printf 가 빈 칸을 찍는다.
    call = re.search(r't next "\$DEST"((?: "\$PY")+)', sh)
    if not call:
        raise Fail('install.sh no longer calls `t next` the way this case reads it')
    eq(call.group(1).count('$PY'), len(quick['SKILL.md']),
       'every script in the note gets an interpreter argument')


if __name__ == '__main__':
    # 옆에 놓인 `selftest_*.py` 도 같은 목록에 올린다 — 이 파일 하나만 부르면 전부 돈다.
    load_extras()
    sys.exit(main())
